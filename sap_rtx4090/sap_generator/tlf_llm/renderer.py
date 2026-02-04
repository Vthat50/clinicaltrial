"""Render LLM-generated TLF shell specifications to Markdown or DOCX format.

DOCX output uses real pharmaceutical TLF shell formatting:
- Monospaced text layout (Courier New) — NOT Word tables
- Columns aligned by character position
- Horizontal rules made of underscore characters
- Page header: Protocol ID + CONFIDENTIAL + Page X of Y
- Page footer: Output date placeholder
- Footnotes as bracketed numbers below the shell
"""

import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Format code → placeholder mapping (### notation for field width/precision)
# ---------------------------------------------------------------------------

_FORMAT_PLACEHOLDERS = {
    # Using ### notation: # = digit position, shows exact field width and decimal places
    "count": "###",                        # 3-digit integer
    "count_pct": "### (##.#%)",            # count with percentage (1 decimal)
    "percentage": "##.#%",                 # percentage with 1 decimal
    "mean": "###.#",                       # mean with 1 decimal
    "sd": "##.##",                         # SD with 2 decimals
    "mean_sd": "###.# (##.##)",            # mean (SD)
    "mean_se": "###.# (##.##)",            # mean (SE)
    "mean_ci": "###.# (###.#, ###.#)",     # mean with 95% CI
    "median": "###.#",                     # median with 1 decimal
    "median_ci": "###.# (###.#, ###.#)",   # median with 95% CI
    "median_range": "###.# (###-###)",     # median (min-max)
    "min": "###.#",                        # minimum
    "max": "###.#",                        # maximum
    "q1_q3": "###.#, ###.#",               # Q1, Q3
    "min_max": "###-###",                  # min-max range
    "ci_95": "(###.#, ###.#)",             # 95% CI only
    "hr_ci": "#.## (##.##, ##.##)",        # hazard ratio with CI (2 decimals)
    "hazard_ratio": "#.## (##.##, ##.##)", # hazard ratio with CI
    "diff_ci": "##.# (##.#, ##.#)",        # treatment difference with CI
    "ratio_ci": "#.## (##.##, ##.##)",     # ratio with CI
    "rate_ci": "##.#% (##.#, ##.#)",       # rate with CI
    "p_value": "#.####",                   # p-value with 4 decimals
    "rate_ratio": "#.## (##.##, ##.##)",   # rate ratio with CI
    "or_ci": "#.## (##.##, ##.##)",        # odds ratio with CI
    "events_rate": "### (##.#%)",          # events with rate
    "n_pct": "### (##.#%)",                # n with percentage
    "fixed": "###",                        # generic integer
    "text": "XXXXXX",                      # text placeholder
}

# Comparison formats: single value spanning columns, not per-arm
_COMPARISON_FORMATS = {
    "hr_ci", "hazard_ratio", "diff_ci", "ratio_ci", "p_value",
    "or_ci", "rate_ratio",
}

_SECTION_LABELS = {
    "14.1": "Section 14.1 — Disposition and Demographics",
    "14.2": "Section 14.2 — Efficacy",
    "14.3": "Section 14.3 — Safety",
    "16.2": "Section 16.2 — Data Listings",
}

# Standard abbreviations for column headers (used in listings to prevent truncation)
_HEADER_ABBREVIATIONS = {
    "Subject ID": "Subj ID",
    "Treatment Group": "Trt Grp",
    "Assessment Date": "Assess Dt",
    "Collection Date": "Coll Dt",
    "Collection Time": "Coll Time",
    "Target Lesion Response": "Tgt Resp",
    "Non-Target Lesion Response": "Non-Tgt",
    "Non-Target Response": "Non-Tgt",
    "New Lesions": "New Les",
    "Overall Response": "Ovrl Resp",
    "Best Overall Response": "BOR",
    "Review Type": "Rev Type",
    "CTCAE Grade": "Grade",
    "Change from Baseline": "CFB",
    "Baseline": "BL",
    "Baseline Value": "BL Value",
    "Study Day": "Study Dy",
    "Reference Range": "Ref Range",
    "Ref Range": "Ref Rng",
    "Abnormal Flag": "Abn Flag",
    "Clinical Significance": "Clin Sig",
    "Parameter": "Param",
    "Visit": "Visit",
    "Result": "Result",
    "Unit": "Unit",
}


def _placeholder(fmt: str) -> str:
    return _FORMAT_PLACEHOLDERS.get(fmt, "xxx")


# =========================================================================
# MARKDOWN RENDERING
# =========================================================================

def _render_table_md(table: dict) -> str:
    number = table.get("number", "")
    title = table.get("title", "")
    population = table.get("population", "")
    source = table.get("source", "")
    orientation = table.get("orientation", "PORTRAIT")
    columns = table.get("columns", [])
    rows = table.get("rows", [])
    footnotes = table.get("footnotes", [])

    lines = []

    pop_suffix = ""
    if population and population.lower() not in title.lower():
        pop_suffix = f" ({population} Population)"
    lines.append(f"### {number}: {title}{pop_suffix}")
    lines.append("")
    lines.append(
        f"**Source Dataset:** {source} | **Population:** {population} "
        f"| **Orientation:** {orientation}"
    )
    lines.append("")

    if columns:
        headers = [c.get("header", "").replace("\n", " ") for c in columns]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---" for _ in columns]) + "|")
    else:
        columns = [{"header": "Parameter"}, {"header": "Value"}]
        lines.append("| Parameter | Value |")
        lines.append("|-----------|-------|")

    # Detect multi-column label rows (visit/statistic fields)
    has_multi_label = any(row.get("visit") or row.get("statistic") for row in rows)

    if has_multi_label:
        num_label_cols = 3
        num_data_cols = max(len(columns) - num_label_cols, 1)
    else:
        num_label_cols = 1
        num_data_cols = max(len(columns) - 1, 1)

    for row in rows:
        label = row.get("label", "")
        row_type = row.get("type", "data")
        fmt = row.get("format", "")
        indent = row.get("indent", 0)
        bold = row.get("bold", False)
        visit = row.get("visit", "")
        statistic = row.get("statistic", "")

        if row_type == "spacer":
            lines.append("| " + " | ".join(["" for _ in columns]) + " |")
            continue

        if not has_multi_label and not label:
            lines.append("| " + " | ".join(["" for _ in columns]) + " |")
            continue

        display_label = ("\u00a0\u00a0" * indent) + label
        if bold:
            display_label = f"**{display_label}**"

        if has_multi_label:
            if row_type == "header":
                cells = [display_label, visit, statistic] + ["" for _ in range(num_data_cols)]
            elif fmt in _COMPARISON_FORMATS:
                ph = _placeholder(fmt)
                cells = [display_label, visit, statistic, ph] + ["" for _ in range(num_data_cols - 1)]
            else:
                ph = _placeholder(fmt) if fmt else ""
                cells = [display_label, visit, statistic] + ([ph] * num_data_cols if ph else ["" for _ in range(num_data_cols)])
        else:
            if row_type == "header":
                cells = [display_label] + ["" for _ in range(num_data_cols)]
            elif fmt in _COMPARISON_FORMATS:
                ph = _placeholder(fmt)
                cells = [display_label, ph] + ["" for _ in range(num_data_cols - 1)]
            else:
                ph = _placeholder(fmt)
                cells = [display_label] + [ph for _ in range(num_data_cols)]

        lines.append("| " + " | ".join(cells) + " |")

    lines.append("")

    if footnotes:
        for i, fn in enumerate(footnotes, 1):
            lines.append(f"[{i}] {fn}")
        lines.append("")

    programming_notes = table.get("programming_notes", "")
    if programming_notes:
        lines.append(f"**Programming Notes:** {programming_notes}")
        lines.append("")

    table_num = str(number).replace("Table ", "").replace(".", "_")
    lines.append(f"Source: {source}  |  Program: t_{table_num}.sas  |  Date: DDMONYYYY")
    lines.append("")

    return "\n".join(lines)


def _render_figure_md(fig: dict) -> str:
    number = fig.get("number", "")
    title = fig.get("title", "")
    fig_type = fig.get("type", "figure")
    population = fig.get("population", "")
    endpoint = fig.get("endpoint", "")

    lines = [f"### {number}: {title}", ""]
    if population:
        lines.append(f"Population: {population}")
    if endpoint:
        lines.append(f"Endpoint: {endpoint}")
    lines.append(f"Type: {fig_type.replace('_', ' ').title()}")
    lines.append("")

    if "km" in fig_type.lower():
        lines.append("[Kaplan-Meier survival curve — to be generated by statistical programming]")
    elif "forest" in fig_type.lower():
        lines.append("[Forest plot — to be generated by statistical programming]")
    elif "waterfall" in fig_type.lower():
        lines.append("[Waterfall plot — to be generated by statistical programming]")
    elif "swimmer" in fig_type.lower():
        lines.append("[Swimmer plot — to be generated by statistical programming]")
    else:
        lines.append("[Figure placeholder — to be generated by statistical programming]")

    lines.append("")
    return "\n".join(lines)


def _render_listing_md(listing: dict) -> str:
    number = listing.get("number", "")
    title = listing.get("title", "")
    population = listing.get("population", "Safety")
    source = listing.get("source", "")
    sort_order = listing.get("sort_order", "")
    page_break_by = listing.get("page_break_by", "")
    footnotes = listing.get("footnotes", [])
    programming_notes = listing.get("programming_notes", "")
    variables = listing.get(
        "variables", ["Subject ID", "Treatment", "Parameter", "Value"]
    )

    lines = [
        f"### {number}: {title} ({population} Population)",
        "",
        f"**Population:** {population} | **Source:** {source} | **Orientation:** LANDSCAPE",
    ]
    if sort_order:
        lines.append(f"**Sort Order:** {sort_order}")
    if page_break_by:
        lines.append(f"**Page Break By:** {page_break_by}")
    lines.append("")

    lines.append("| " + " | ".join(variables) + " |")
    lines.append("|" + "|".join(["---" for _ in variables]) + "|")
    lines.append("| " + " | ".join(["xxx" for _ in variables]) + " |")
    lines.append("")

    if footnotes:
        for i, fn in enumerate(footnotes, 1):
            lines.append(f"[{i}] {fn}")
        lines.append("")

    if programming_notes:
        lines.append(f"**Programming Notes:** {programming_notes}")
        lines.append("")

    listing_num = str(number).replace("Listing ", "").replace(".", "_")
    lines.append(f"Source: {source}  |  Program: l_{listing_num}.sas  |  Date: DDMONYYYY")
    lines.append("")

    return "\n".join(lines)


def _render_summary_md(
    tables: list[dict], figures: list[dict], listings: list[dict]
) -> str:
    lines = ["---", "", "### TLF Shell Summary", "", "| Category | Count |", "|----------|-------|"]

    section_counts: dict[str, int] = {}
    for t in tables:
        sec = t.get("section", "14.3")
        cat = {"14.1": "Disposition & Demographics", "14.2": "Efficacy", "14.3": "Safety"}.get(sec, sec)
        section_counts[cat] = section_counts.get(cat, 0) + 1

    for cat, count in sorted(section_counts.items()):
        lines.append(f"| {cat} | {count} |")

    lines.append(f"| Figures | {len(figures)} |")
    lines.append(f"| Listings | {len(listings)} |")
    total = len(tables) + len(figures) + len(listings)
    lines.append(f"| **Total TLFs** | **{total}** |")
    lines.append("")

    return "\n".join(lines)


def render_markdown(
    tables: list[dict],
    figures: list[dict],
    listings: list[dict],
    domain_facts: Optional[dict] = None,
) -> str:
    parts = ["# TLF Shell Specifications\n"]
    parts.append(
        f"**Total:** {len(tables)} tables, {len(figures)} figures, "
        f"{len(listings)} listings\n"
    )

    if tables:
        parts.append("## Tables\n")
        current_section = ""
        for t in tables:
            logger.info(f"[Renderer] Table '{t.get('title', '?')[:50]}' keys: {list(t.keys())}, "
                        f"columns={len(t.get('columns', []))}, rows={len(t.get('rows', []))}")
            section = t.get("section", "")
            if section != current_section:
                current_section = section
                section_label = _SECTION_LABELS.get(section, section)
                parts.append(f"\n### {section_label}\n")
            parts.append(_render_table_md(t))

    if figures:
        parts.append("\n## Figures\n")
        for fig in figures:
            parts.append(_render_figure_md(fig))

    if listings:
        parts.append("\n## Listings (ICH E3 Section 16.2)\n")
        for li in listings:
            parts.append(_render_listing_md(li))

    parts.append(_render_summary_md(tables, figures, listings))

    return "\n".join(parts)


# =========================================================================
# DOCX RENDERING — Monospaced Text Layout (NOT Word Tables)
# =========================================================================

_FONT = "Courier New"
_FONT_SIZE_BODY = 9
_FONT_SIZE_TITLE = 10
_FONT_SIZE_FOOTNOTE = 8
_FONT_SIZE_PAGE_HDR = 8

# Line widths in characters
_LINE_WIDTH_LANDSCAPE = 132
_LINE_WIDTH_PORTRAIT = 80


def render_docx(
    tables: list[dict],
    figures: list[dict],
    listings: list[dict],
) -> bytes:
    """Render TLF shells as a Word document using monospaced text layout.

    Key differences from Word tables:
    - Uses fixed-width Courier New text, not table cells
    - Columns aligned by character position
    - Horizontal rules made of underscore characters
    - No column reflow or wrapping issues
    - Opens identically in Word, Google Docs, LibreOffice
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX output")

    doc = Document()

    # --- Default font: Courier New ---
    style = doc.styles["Normal"]
    style.font.name = _FONT
    style.font.size = Pt(_FONT_SIZE_BODY)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # --- Helpers ---

    def _setup_section(section, landscape=False):
        """Configure page layout and header/footer."""
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Header
        hdr = section.header
        hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0]
        hp.text = ""

        run1 = hp.add_run("Sponsor Name")
        run1.font.name = _FONT
        run1.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        run1.font.color.rgb = RGBColor(100, 100, 100)

        hp.add_run("    ")

        run2 = hp.add_run("Protocol Number")
        run2.font.name = _FONT
        run2.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        run2.font.color.rgb = RGBColor(100, 100, 100)

        hp.add_run("    ")

        run3 = hp.add_run("CONFIDENTIAL")
        run3.font.name = _FONT
        run3.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        run3.font.color.rgb = RGBColor(100, 100, 100)

        hp.add_run("    ")

        run4 = hp.add_run("Page ")
        run4.font.name = _FONT
        run4.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        run4.font.color.rgb = RGBColor(100, 100, 100)

        # Dynamic PAGE field
        page_field = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "/>')
        hp._p.append(page_field)

        run5 = hp.add_run(" of ")
        run5.font.name = _FONT
        run5.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        run5.font.color.rgb = RGBColor(100, 100, 100)

        # Dynamic NUMPAGES field
        numpages_field = parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "/>')
        hp._p.append(numpages_field)

        # Footer
        ftr = section.footer
        ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0]
        fp.text = ""
        ftr_run = fp.add_run("Output generated: DDMONYYYY")
        ftr_run.font.name = _FONT
        ftr_run.font.size = Pt(_FONT_SIZE_PAGE_HDR)
        ftr_run.font.color.rgb = RGBColor(100, 100, 100)

    def _add_mono_line(text: str, bold: bool = False, size: int = None, color: RGBColor = None):
        """Add a monospaced text line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.name = _FONT
        run.font.size = Pt(size or _FONT_SIZE_BODY)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    def _add_page_break():
        """Add a page break."""
        doc.add_page_break()

    def _add_landscape_section():
        """Add a new landscape section."""
        new_sect = doc.add_section(2)  # Continuous section break
        _setup_section(new_sect, landscape=True)
        return new_sect

    def _add_portrait_section():
        """Add a new portrait section."""
        new_sect = doc.add_section(2)
        _setup_section(new_sect, landscape=False)
        return new_sect

    def _should_be_landscape(table_dict: dict) -> bool:
        """Determine if a table should be landscape."""
        orientation = table_dict.get("orientation", "").upper()
        if orientation == "LANDSCAPE":
            return True

        table_type = table_dict.get("type", "").lower()
        title_lower = table_dict.get("title", "").lower()

        if table_type in ("labs_summary", "labs_shift", "lab_listing"):
            return True
        if "chemistry" in title_lower or "hematology" in title_lower:
            return True
        if "laboratory" in title_lower and "by visit" in title_lower:
            return True
        if "system organ class" in title_lower or "preferred term" in title_lower:
            return True

        columns = table_dict.get("columns", [])
        if len(columns) >= 5:
            return True

        # Check if row labels + placeholders would squeeze columns too much in portrait
        # If longest label > 35 chars AND we have CI/comparison stats, need landscape
        rows = table_dict.get("rows", [])
        max_label_len = 0
        has_wide_placeholders = False
        for row in rows:
            label = row.get("label", "")
            indent = row.get("indent", 0)
            max_label_len = max(max_label_len, len(label) + indent * 2)
            # Check for wide format codes (CI, HR, etc.)
            fmt = row.get("format", "")
            if fmt in ("median_ci", "hr_ci", "hazard_ratio", "rate_ci", "diff_ci", "ratio_ci", "mean_ci"):
                has_wide_placeholders = True

        # If we have long labels AND wide placeholders, need landscape for proper spacing
        if max_label_len > 35 and has_wide_placeholders:
            return True

        return False

    def _calculate_column_widths(headers: list[str], num_cols: int, line_width: int) -> list[int]:
        """Calculate column widths in characters based on header content.

        Returns list of character widths for each column.
        Ensures no header gets truncated.
        """
        if num_cols == 0:
            return []

        if num_cols == 1:
            return [line_width - 2]  # Leave margin

        # Calculate minimum width needed for each header (with padding)
        min_widths = []
        for hdr in headers:
            # Handle headers with (N=xxx) - take the longer part
            if "(N=" in hdr:
                parts = hdr.split("(N=")
                hdr_len = max(len(parts[0].strip()), len("(N=" + parts[1]) if len(parts) > 1 else 0)
            else:
                hdr_len = len(hdr)
            min_widths.append(max(hdr_len + 2, 12))  # At least 12 chars per column

        # First column (Parameter/label) needs more space for indented labels
        # Make it at least 50% of line width or header length, whichever is larger
        # This prevents truncation of long labels like "Number of patients with event, n (%)"
        min_widths[0] = max(min_widths[0], int(line_width * 0.50))

        # Check total
        total_needed = sum(min_widths) + (num_cols - 1) * 2  # 2 char spacing

        if total_needed <= line_width:
            # Fits - distribute extra space to data columns
            extra = line_width - total_needed
            extra_per_col = extra // (num_cols - 1) if num_cols > 1 else extra
            widths = [min_widths[0]] + [w + extra_per_col for w in min_widths[1:]]
        else:
            # Doesn't fit - scale down data columns but protect first column
            first_col = min_widths[0]
            remaining = line_width - first_col - (num_cols - 1) * 2
            data_col_width = max(remaining // (num_cols - 1), 10) if num_cols > 1 else remaining
            widths = [first_col] + [data_col_width] * (num_cols - 1)

        return widths

    def _format_row(values: list[str], widths: list[int], indent: int = 0) -> str:
        """Format a row with fixed column positions.

        Args:
            values: List of cell values
            widths: List of column widths in characters
            indent: Number of spaces to indent the first column
        """
        parts = []
        for i, (val, width) in enumerate(zip(values, widths)):
            if i == 0:
                # First column: left-align with indent
                indented_val = " " * (indent * 2) + str(val)
                parts.append(indented_val[:width].ljust(width))
            else:
                # Data columns: center-align
                parts.append(str(val)[:width].center(width))
        return "  ".join(parts)

    def _render_table_mono(table_dict: dict, line_width: int):
        """Render a table shell as monospaced text."""
        number = table_dict.get("number", "")
        title = table_dict.get("title", "")
        population = table_dict.get("population", "")
        source = table_dict.get("source", "")
        columns = table_dict.get("columns", [])
        rows = table_dict.get("rows", [])
        footnotes = table_dict.get("footnotes", [])
        prog_notes = table_dict.get("programming_notes", "")

        # Rule line
        rule = "_" * line_width

        # Title
        pop_suffix = ""
        if population and population.lower() not in title.lower():
            pop_suffix = f" ({population} Population)"
        _add_mono_line(f"{number}: {title}{pop_suffix}", bold=True)
        _add_mono_line("")

        if not columns:
            columns = [{"header": "Parameter"}, {"header": "Value"}]

        headers = [c.get("header", "").replace("\n", " ") for c in columns]
        num_cols = len(headers)

        # Find the longest row label to ensure first column is wide enough
        max_label_len = 0
        for row in rows:
            label = row.get("label", "")
            indent = row.get("indent", 0)
            label_len = len(label) + (indent * 2)  # Account for indent spaces
            max_label_len = max(max_label_len, label_len)

        col_widths = _calculate_column_widths(headers, num_cols, line_width)

        # Ensure first column can fit the longest label (with padding)
        if max_label_len > 0:
            col_widths[0] = max(col_widths[0], max_label_len + 2)

        # Detect multi-column label (Parameter, Visit, Statistic)
        has_multi_label = any(row.get("visit") or row.get("statistic") for row in rows)

        # Top rule
        _add_mono_line(rule)

        # Header rows - may need 2 lines if headers have (N=xxx)
        header_line1 = []
        header_line2 = []
        for i, hdr in enumerate(headers):
            if "\n" in hdr or "(N=" in hdr:
                # Split into two lines
                if "(N=" in hdr:
                    parts = hdr.split("(N=")
                    header_line1.append(parts[0].strip())
                    header_line2.append("(N=" + parts[1] if len(parts) > 1 else "")
                else:
                    parts = hdr.split("\n")
                    header_line1.append(parts[0])
                    header_line2.append(parts[1] if len(parts) > 1 else "")
            else:
                header_line1.append(hdr)
                header_line2.append("")

        _add_mono_line(_format_row(header_line1, col_widths))
        if any(header_line2):
            _add_mono_line(_format_row(header_line2, col_widths))

        # Header underline rule
        _add_mono_line(rule)

        # Data rows
        num_data_cols = num_cols - 1 if not has_multi_label else num_cols - 3

        for row in rows:
            label = row.get("label", "")
            row_type = row.get("type", "data")
            fmt = row.get("format", "")
            indent = row.get("indent", 0)
            bold = row.get("bold", False)
            visit = row.get("visit", "")
            statistic = row.get("statistic", "")

            if row_type == "spacer":
                _add_mono_line("")
                continue

            is_comparison = fmt in _COMPARISON_FORMATS

            if has_multi_label:
                if row_type == "header":
                    cells = [label, visit, statistic] + [""] * max(num_data_cols, 0)
                else:
                    ph = _placeholder(fmt) if fmt else ""
                    cells = [label, visit, statistic] + ([ph] * max(num_data_cols, 0) if ph else [""] * max(num_data_cols, 0))
            else:
                if row_type == "header":
                    cells = [label] + [""] * max(num_data_cols, 0)
                elif is_comparison:
                    # Comparison stats go in a single centered position
                    ph = _placeholder(fmt)
                    cells = [label, ph] + [""] * max(num_data_cols - 1, 0)
                elif not label:
                    cells = [""] * num_cols
                else:
                    ph = _placeholder(fmt)
                    cells = [label] + [ph] * max(num_data_cols, 0)

            # Ensure cells match column count
            while len(cells) < num_cols:
                cells.append("")
            cells = cells[:num_cols]

            row_text = _format_row(cells, col_widths, indent=indent)
            _add_mono_line(row_text, bold=bold or row_type == "header")

        # Bottom rule
        _add_mono_line(rule)
        _add_mono_line("")

        # Footnotes
        if footnotes:
            for i, fn in enumerate(footnotes, 1):
                _add_mono_line(f"[{i}] {fn}", size=_FONT_SIZE_FOOTNOTE)

        # Programming notes
        if prog_notes:
            _add_mono_line("")
            _add_mono_line(f"Programming Notes: {prog_notes}", size=7, color=RGBColor(100, 100, 100))

        # Source line
        table_num = str(number).replace("Table ", "").replace(".", "_")
        _add_mono_line("")
        _add_mono_line(f"Source: {source}    Program: t_{table_num}.sas    Date: DDMONYYYY",
                       size=_FONT_SIZE_FOOTNOTE, color=RGBColor(100, 100, 100))

    def _render_listing_mono(lst: dict, line_width: int):
        """Render a listing shell as monospaced text."""
        number = lst.get("number", "")
        title = lst.get("title", "")
        population = lst.get("population", "Safety")
        source = lst.get("source", "")
        sort_order = lst.get("sort_order", "")
        page_break = lst.get("page_break_by", "")
        footnotes = lst.get("footnotes", [])
        prog_notes = lst.get("programming_notes", "")
        variables = lst.get("variables", ["Subject ID", "Treatment", "Parameter", "Value"])

        rule = "_" * line_width

        # Title
        pop_suffix = ""
        if population and population.lower() not in title.lower():
            pop_suffix = f" ({population} Population)"
        _add_mono_line(f"{number}: {title}{pop_suffix}", bold=True)
        _add_mono_line("")

        num_cols = len(variables)

        # Apply abbreviations for long headers and track which ones were abbreviated
        abbreviated_headers = []
        abbreviation_map = {}  # abbrev -> full name
        for var in variables:
            if var in _HEADER_ABBREVIATIONS:
                abbrev = _HEADER_ABBREVIATIONS[var]
                abbreviated_headers.append(abbrev)
                if abbrev != var:
                    abbreviation_map[abbrev] = var
            else:
                abbreviated_headers.append(var)

        # Calculate column widths based on actual header lengths
        # Minimum width = max(header length, 8) + 2 padding
        col_widths = []
        for hdr in abbreviated_headers:
            min_width = max(len(hdr), 6) + 2
            col_widths.append(min_width)

        # Check if total width fits; if not, scale down proportionally
        total_needed = sum(col_widths) + (num_cols - 1) * 2  # 2 char spacing
        if total_needed > line_width:
            # Scale down, but keep minimum of 8 chars per column
            scale = (line_width - (num_cols - 1) * 2) / sum(col_widths)
            col_widths = [max(int(w * scale), 8) for w in col_widths]

        # Top rule
        _add_mono_line(rule)

        # Header row (using abbreviated headers)
        _add_mono_line(_format_row(abbreviated_headers, col_widths))

        # Header underline
        _add_mono_line(rule)

        # Sample data rows (4 rows)
        for _ in range(4):
            data_row = ["xxx"] * num_cols
            _add_mono_line(_format_row(data_row, col_widths))

        # Bottom rule
        _add_mono_line(rule)
        _add_mono_line("")

        # Sort order / page break
        if sort_order:
            _add_mono_line(f"Sort Order: {sort_order}", size=_FONT_SIZE_FOOTNOTE)
        if page_break:
            _add_mono_line(f"Page Break By: {page_break}", size=_FONT_SIZE_FOOTNOTE)

        # Add existing footnotes
        footnote_idx = 1
        if footnotes:
            for fn in footnotes:
                _add_mono_line(f"[{footnote_idx}] {fn}", size=_FONT_SIZE_FOOTNOTE)
                footnote_idx += 1

        # Add abbreviation footnote if any were used
        if abbreviation_map:
            abbrev_list = ", ".join(f"{k}={v}" for k, v in sorted(abbreviation_map.items()))
            _add_mono_line(f"[{footnote_idx}] Abbreviations: {abbrev_list}", size=_FONT_SIZE_FOOTNOTE)

        # Programming notes
        if prog_notes:
            _add_mono_line("")
            _add_mono_line(f"Programming Notes: {prog_notes}", size=7, color=RGBColor(100, 100, 100))

        # Source line
        lst_num = str(number).replace("Listing ", "").replace(".", "_")
        _add_mono_line("")
        _add_mono_line(f"Source: {source}    Program: l_{lst_num}.sas    Date: DDMONYYYY",
                       size=_FONT_SIZE_FOOTNOTE, color=RGBColor(100, 100, 100))

    def _render_figure_shell(fig: dict):
        """Render a figure shell placeholder."""
        number = fig.get("number", "")
        title = fig.get("title", "")
        fig_type = fig.get("type", "figure")
        population = fig.get("population", "")
        endpoint = fig.get("endpoint", "")

        _add_mono_line(f"{number}: {title}", bold=True)
        if population:
            _add_mono_line(f"Population: {population}", size=_FONT_SIZE_FOOTNOTE)
        if endpoint:
            _add_mono_line(f"Endpoint: {endpoint}", size=_FONT_SIZE_FOOTNOTE)
        _add_mono_line("")

        # Figure type placeholder
        if "km" in fig_type.lower():
            placeholder = "[Kaplan-Meier survival curve — to be generated by statistical programming]"
        elif "forest" in fig_type.lower():
            placeholder = "[Forest plot — to be generated by statistical programming]"
        elif "waterfall" in fig_type.lower():
            placeholder = "[Waterfall plot — to be generated by statistical programming]"
        elif "swimmer" in fig_type.lower():
            placeholder = "[Swimmer plot — to be generated by statistical programming]"
        else:
            placeholder = "[Figure — to be generated by statistical programming]"

        _add_mono_line("")
        _add_mono_line(placeholder)
        _add_mono_line("")

        fig_num = str(number).replace("Figure ", "").replace(".", "_")
        _add_mono_line(f"Program: f_{fig_num}.sas    Date: DDMONYYYY",
                       size=_FONT_SIZE_FOOTNOTE, color=RGBColor(100, 100, 100))

    # =================================================================
    # BUILD DOCUMENT
    # =================================================================

    # Setup initial section (portrait)
    _setup_section(doc.sections[0], landscape=False)

    # --- TITLE PAGE ---
    _add_mono_line("")
    _add_mono_line("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TABLE, LISTING, AND FIGURE SHELLS")
    run.font.name = _FONT
    run.font.size = Pt(14)
    run.font.bold = True

    _add_mono_line("")
    total = len(tables) + len(figures) + len(listings)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{len(tables)} Tables  |  {len(figures)} Figures  |  {len(listings)} Listings  |  {total} Total")
    run2.font.name = _FONT
    run2.font.size = Pt(10)

    # --- TLF INDEX (Page 2) ---
    _add_page_break()
    _add_mono_line("TLF INDEX", bold=True, size=12)
    _add_mono_line("")

    # Simple text-based index
    _add_mono_line("_" * _LINE_WIDTH_PORTRAIT)
    _add_mono_line(f"{'Number':<20}  {'Title':<45}  {'Population':<15}")
    _add_mono_line("_" * _LINE_WIDTH_PORTRAIT)

    for t in tables:
        num = t.get("number", "")
        title = t.get("title", "")[:42]
        pop = t.get("population", "")[:15]
        _add_mono_line(f"{num:<20}  {title:<45}  {pop:<15}")

    for f in figures:
        num = f.get("number", "")
        title = f.get("title", "")[:42]
        pop = f.get("population", "")[:15]
        _add_mono_line(f"{num:<20}  {title:<45}  {pop:<15}")

    for li in listings:
        num = li.get("number", "")
        title = li.get("title", "")[:42]
        pop = li.get("population", "")[:15]
        _add_mono_line(f"{num:<20}  {title:<45}  {pop:<15}")

    _add_mono_line("_" * _LINE_WIDTH_PORTRAIT)

    # --- TABLES ---
    current_orientation = "PORTRAIT"

    for table_dict in tables:
        is_landscape = _should_be_landscape(table_dict)
        target_orientation = "LANDSCAPE" if is_landscape else "PORTRAIT"
        line_width = _LINE_WIDTH_LANDSCAPE if is_landscape else _LINE_WIDTH_PORTRAIT

        # Switch section if orientation changes
        if target_orientation != current_orientation:
            if target_orientation == "LANDSCAPE":
                _add_landscape_section()
            else:
                _add_portrait_section()
            current_orientation = target_orientation
        else:
            _add_page_break()

        _render_table_mono(table_dict, line_width)

    # --- FIGURES ---
    if figures:
        if current_orientation != "PORTRAIT":
            _add_portrait_section()
            current_orientation = "PORTRAIT"
        else:
            _add_page_break()

        _add_mono_line("FIGURE SHELLS", bold=True, size=12)
        _add_mono_line("")

        for fig in figures:
            _add_page_break()
            _render_figure_shell(fig)

    # --- LISTINGS ---
    for lst in listings:
        # Listings always landscape
        if current_orientation != "LANDSCAPE":
            _add_landscape_section()
            current_orientation = "LANDSCAPE"
        else:
            _add_page_break()

        _render_listing_mono(lst, _LINE_WIDTH_LANDSCAPE)

    # --- SUMMARY PAGE ---
    if current_orientation != "PORTRAIT":
        _add_portrait_section()
    else:
        _add_page_break()

    _add_mono_line("TLF SHELL SUMMARY", bold=True, size=12)
    _add_mono_line("")

    section_counts: dict[str, int] = {}
    for t in tables:
        sec = t.get("section", "14.3")
        cat = {"14.1": "Disposition & Demographics", "14.2": "Efficacy", "14.3": "Safety"}.get(sec, sec)
        section_counts[cat] = section_counts.get(cat, 0) + 1

    _add_mono_line("_" * 50)
    _add_mono_line(f"{'Category':<35}  {'Count':>10}")
    _add_mono_line("_" * 50)

    for cat, count in sorted(section_counts.items()):
        _add_mono_line(f"{cat:<35}  {count:>10}")

    _add_mono_line(f"{'Figures':<35}  {len(figures):>10}")
    _add_mono_line(f"{'Listings':<35}  {len(listings):>10}")
    _add_mono_line("_" * 50)
    _add_mono_line(f"{'TOTAL TLFs':<35}  {total:>10}", bold=True)
    _add_mono_line("_" * 50)

    # Save
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
