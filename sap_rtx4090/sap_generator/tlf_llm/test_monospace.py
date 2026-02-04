"""Test monospaced TLF shell rendering - verify alignment in Word/Google Docs."""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

def create_test_monospace_doc(output_path: str):
    """Create a test DOCX with monospaced text shells (no tables)."""

    doc = Document()

    # Set up page - landscape, narrow margins for 132-char width
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Helper to add monospace paragraph
    def add_mono_line(text: str, bold: bool = False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.bold = bold
        return p

    # Line width for landscape (132 chars standard)
    LINE_WIDTH = 132

    # =========================================================================
    # SHELL 1: PFS Summary Table
    # =========================================================================

    rule = "_" * LINE_WIDTH

    add_mono_line("Table 14.2.1.1: Summary of Progression-Free Survival (PFS) - Intent-to-Treat Population", bold=True)
    add_mono_line("")
    add_mono_line(rule)

    # Column positions (character-based)
    # Col 1: Parameter (0-50)
    # Col 2: Treatment A (52-75)
    # Col 3: Treatment B (77-100)
    # Col 4: Total (102-125)

    def format_row(col1: str, col2: str = "", col3: str = "", col4: str = "", indent: int = 0):
        """Format a row with fixed column positions."""
        col1_width = 50
        col2_width = 24
        col3_width = 24
        col4_width = 24

        # Apply indent
        col1 = " " * (indent * 2) + col1

        # Pad/truncate each column
        c1 = col1[:col1_width].ljust(col1_width)
        c2 = col2[:col2_width].center(col2_width) if col2 else " " * col2_width
        c3 = col3[:col3_width].center(col3_width) if col3 else " " * col3_width
        c4 = col4[:col4_width].center(col4_width) if col4 else " " * col4_width

        return f"{c1}  {c2}  {c3}  {c4}"

    # Header rows
    add_mono_line(format_row("", "CT-P16", "EU-Avastin", "Total"))
    add_mono_line(format_row("Parameter", "(N=###)", "(N=###)", "(N=###)"))
    add_mono_line(rule)

    # Data rows
    add_mono_line(format_row("Progression-Free Survival", "", "", ""), bold=True)
    add_mono_line(format_row("Number of patients with event, n (%)", "### (##.#%)", "### (##.#%)", "### (##.#%)"))
    add_mono_line(format_row("Progressive disease", "### (##.#%)", "### (##.#%)", "### (##.#%)", indent=1))
    add_mono_line(format_row("Death", "### (##.#%)", "### (##.#%)", "### (##.#%)", indent=1))
    add_mono_line(format_row("Number of patients censored, n (%)", "### (##.#%)", "### (##.#%)", "### (##.#%)"))
    add_mono_line(format_row(""))
    add_mono_line(format_row("Kaplan-Meier Estimates", "", "", ""), bold=True)
    add_mono_line(format_row("Median PFS, months (95% CI)", "##.# (##.#, ##.#)", "##.# (##.#, ##.#)", ""))
    add_mono_line(format_row(""))
    add_mono_line(format_row("Treatment Comparison: CT-P16 vs EU-Avastin", "", "", ""), bold=True)
    add_mono_line(format_row("Hazard Ratio (95% CI)", "#.## (##.##, ##.##)", "", ""))
    add_mono_line(format_row("P-value (log-rank)", "#.####", "", ""))

    add_mono_line(rule)
    add_mono_line("")
    add_mono_line("Source: ADTTE")
    add_mono_line("Population: Intent-to-Treat Population - all randomized patients.")
    add_mono_line("Hazard ratio from Cox proportional hazards model stratified by region.")
    add_mono_line("P-value from stratified log-rank test.")
    add_mono_line("")
    add_mono_line("")

    # =========================================================================
    # SHELL 2: Listing with many columns
    # =========================================================================

    add_mono_line("Listing 16.2.1: Individual Tumor Assessment Data - Intent-to-Treat Population", bold=True)
    add_mono_line("")
    add_mono_line(rule)

    # For listings, use abbreviated headers
    def format_listing_row(subj, trt, visit, date, tgt, ntgt, new, ovrl, bor, rev):
        """Format listing row with abbreviated columns."""
        # Fixed widths for each column
        return (f"{subj:<10}  {trt:<12}  {visit:<12}  {date:<11}  "
                f"{tgt:<8}  {ntgt:<8}  {new:<6}  {ovrl:<8}  {bor:<8}  {rev:<10}")

    # Header with abbreviations
    add_mono_line(format_listing_row("Subject", "Treatment", "Visit", "Date", "Target", "Non-Tgt", "New", "Overall", "Best OR", "Review"))
    add_mono_line(format_listing_row("ID", "Group", "", "(DDMONYY)", "Response", "Response", "Lesion", "Response", "", "Type"))
    add_mono_line(rule)

    # Data placeholder rows
    add_mono_line(format_listing_row("###-###", "xxxxxxxxxxx", "xxxxxxxxxxx", "DDMONYYYY", "xx", "xx", "xx", "xx", "xx", "xxxxxxxxx"))
    add_mono_line(format_listing_row("###-###", "xxxxxxxxxxx", "xxxxxxxxxxx", "DDMONYYYY", "xx", "xx", "xx", "xx", "xx", "xxxxxxxxx"))
    add_mono_line(format_listing_row("###-###", "xxxxxxxxxxx", "xxxxxxxxxxx", "DDMONYYYY", "xx", "xx", "xx", "xx", "xx", "xxxxxxxxx"))
    add_mono_line(format_listing_row("###-###", "xxxxxxxxxxx", "xxxxxxxxxxx", "DDMONYYYY", "xx", "xx", "xx", "xx", "xx", "xxxxxxxxx"))

    add_mono_line(rule)
    add_mono_line("")
    add_mono_line("Source: ADRS")
    add_mono_line("Population: Intent-to-Treat Population")
    add_mono_line("Sorted by Subject ID, Visit Date")
    add_mono_line("Abbreviations: Tgt=Target, Non-Tgt=Non-Target, OR=Overall Response")

    # Save
    doc.save(output_path)
    print(f"Saved test document to: {output_path}")
    return output_path


if __name__ == "__main__":
    output = "/tmp/test_monospace_tlf.docx"
    create_test_monospace_doc(output)
    print(f"\nOpen this file in Word and Google Docs to verify alignment:")
    print(f"  {output}")
