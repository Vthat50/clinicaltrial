"""Demo TLF generator — 4 targeted items via 2 Claude API calls.

Inputs:
  1. Protocol file (PDF/DOCX/TXT) — detailed clinical content
  2. Generated SAP (.docx) — TLF index + statistical methodology

The generator parses the TLF index from the SAP to find:
  - First table in the index → generates full shell
  - Clinical Chemistry CFB table → generates full shell
  - First listing in the index → generates full shell
  - Laboratory Values: Clinical Chemistry listing → generates full shell

Both protocol and SAP text are passed to Claude for maximum context.
Reuses tlf_llm infrastructure (Claude caller, JSON parser, numbering, rendering).
"""

import logging
import re
from typing import Any

from docx import Document as DocxDocument

from tlf_llm.generator import _call_claude  # Fix 2.1: Removed unused _MAX_PROTOCOL_CHARS
from tlf_llm.numbering import assign_numbers
from tlf_llm.reference_library import _parse_json_from_response
from tlf_llm.renderer import render_docx, render_markdown

from .prompts import DEMO_SYSTEM_PROMPT, FIRST_TLF_PROMPT, LAB_CHEMISTRY_PROMPT

logger = logging.getLogger(__name__)

# Budget: split the context window between protocol and SAP
_MAX_PROTOCOL_CHARS_DEMO = 60_000
_MAX_SAP_CHARS_DEMO = 40_000


# ---------------------------------------------------------------------------
# SAP parsing helpers
# ---------------------------------------------------------------------------

def _extract_text_from_docx(content: bytes) -> str:
    """Extract full text from a .docx file, including table content.

    Fix 2.4: Now extracts both paragraph text AND table content so Claude
    can see the TLF index which lives in tables.
    """
    from io import BytesIO
    doc = DocxDocument(BytesIO(content))

    parts = []

    # Extract paragraph text
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Extract table content (Fix 2.4: TLF index is in tables!)
    for tbl in doc.tables:
        table_rows = []
        for row in tbl.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):  # Skip empty rows
                table_rows.append("\t".join(row_cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    return "\n".join(parts)


def _truncate_at_boundary(text: str, max_chars: int) -> str:
    """Truncate text at paragraph boundary instead of mid-sentence.

    Fix 2.5: Avoids cutting mid-sentence or mid-word which can lose critical
    protocol content like endpoint definitions or visit schedules.
    """
    if len(text) <= max_chars:
        return text

    # Find the last paragraph break before the limit
    truncated = text[:max_chars]

    # Try to find a paragraph break (double newline)
    last_para = truncated.rfind('\n\n')
    if last_para > max_chars * 0.8:  # Only use if we keep at least 80%
        return truncated[:last_para].strip()

    # Fall back to single newline
    last_newline = truncated.rfind('\n')
    if last_newline > max_chars * 0.9:  # Only use if we keep at least 90%
        return truncated[:last_newline].strip()

    # Fall back to sentence boundary
    last_period = truncated.rfind('. ')
    if last_period > max_chars * 0.95:
        return truncated[:last_period + 1].strip()

    # Last resort: hard truncate
    return truncated.strip()


def _parse_tlf_index(content: bytes) -> dict:
    """Parse the TLF index tables from a generated SAP .docx.

    Looks for tables with columns [Number, Title, Population] and
    classifies them as tables/figures/listings by the Number prefix.

    Returns:
        {
            "tables": [{"number": "14.2.1", "title": "...", "population": "..."}, ...],
            "figures": [...],
            "listings": [...],
        }
    """
    from io import BytesIO
    doc = DocxDocument(BytesIO(content))

    index = {"tables": [], "figures": [], "listings": []}

    for tbl in doc.tables:
        if len(tbl.rows) < 2 or len(tbl.columns) < 3:
            continue

        # Check if header row matches [Number, Title, Population] (flexible matching)
        hdr = [tbl.cell(0, c).text.strip().lower() for c in range(min(len(tbl.columns), 3))]
        if not ("number" in hdr[0] or "no" == hdr[0] or "no." == hdr[0]):
            continue
        if "title" not in hdr[1]:
            continue
        if "population" not in hdr[2] and "analysis" not in hdr[2]:
            continue

        for r in range(1, len(tbl.rows)):
            number = tbl.cell(r, 0).text.strip()
            title = tbl.cell(r, 1).text.strip()
            population = tbl.cell(r, 2).text.strip()

            if not number or not title:
                continue

            entry = {"number": number, "title": title, "population": population}

            # Fix 2.8: More robust classification logic
            # ICH E3: 16.2.x are listings, 14.x.Fx are figures, 14.x.x are tables
            number_upper = number.upper()

            # Check for figures: contains .F followed by digit, or "FIGURE" in title
            if re.search(r'\.F\d', number_upper) or 'FIGURE' in title.upper():
                index["figures"].append(entry)
            # Check for listings: starts with 16.2 (ICH E3 listing section)
            elif number.startswith('16.2') or re.search(r'\.L\d', number_upper) or 'LISTING' in title.upper():
                index["listings"].append(entry)
            else:
                index["tables"].append(entry)

    logger.info(
        f"[TLF Demo] Parsed TLF index: {len(index['tables'])} tables, "
        f"{len(index['figures'])} figures, {len(index['listings'])} listings"
    )
    return index


def _find_chemistry_entry(entries: list[dict]) -> dict | None:
    """Find the Clinical Chemistry entry from a list of index entries.

    Fix 2.2: Removed unused 'kind' parameter.
    """
    chemistry_patterns = [
        r"clinical\s+chemistry",
        r"chemistry\s+parameters",
        r"chemistry.*change\s+from\s+baseline",
        r"chemistry.*actual\s+values",
        r"serum\s+chemistry",
        r"blood\s+chemistry",
        r"chemistry\s+results",
        r"chemistry\s+values",
        r"laboratory.*chemistry",
    ]

    for entry in entries:
        title_lower = entry["title"].lower()
        for pat in chemistry_patterns:
            if re.search(pat, title_lower):
                return entry

    return None


def _validate_tlf_item(item: dict, item_type: str) -> bool:
    """Validate that a TLF item has required fields.

    Fix 2.7: Validates JSON response structure before using.
    """
    required_common = ["title"]
    required_table = ["columns", "rows"]
    required_listing = ["variables"]

    # Check common fields
    for field in required_common:
        if field not in item or not item[field]:
            logger.warning(f"[TLF Demo] {item_type} missing required field: {field}")
            return False

    # Check type-specific fields
    if item_type == "table":
        for field in required_table:
            if field not in item:
                logger.warning(f"[TLF Demo] Table missing required field: {field}")
                return False
    elif item_type == "listing":
        for field in required_listing:
            if field not in item:
                logger.warning(f"[TLF Demo] Listing missing required field: {field}")
                return False

    return True


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

async def generate_demo_shells(
    protocol_text: str,
    sap_content: bytes,
    output_format: str = "markdown",
) -> Any:
    """Generate 4 demo TLF shells from a protocol + generated SAP.

    Args:
        protocol_text: Extracted text from the protocol document.
        sap_content: Raw bytes of the generated SAP .docx file.
        output_format: "markdown", "docx", or "json".

    Returns:
        Markdown string, DOCX bytes, or dict depending on output_format.

    Raises:
        RuntimeError: If all Claude API calls fail and no TLF shells are generated.
    """
    logger.info("[TLF Demo] Starting demo generation (4 items, 2 calls)")

    # Step 1: Parse the TLF index from the SAP
    tlf_index = _parse_tlf_index(sap_content)

    if not tlf_index["tables"]:
        raise ValueError("No tables found in the SAP's TLF index. "
                         "Please upload a generated SAP with a populated TLF index.")
    if not tlf_index["listings"]:
        raise ValueError("No listings found in the SAP's TLF index. "
                         "Please upload a generated SAP with a populated TLF index.")

    # Step 2: Identify the 4 target items from the index
    first_table = tlf_index["tables"][0]
    first_listing = tlf_index["listings"][0]
    chem_table = _find_chemistry_entry(tlf_index["tables"])  # Fix 2.2: removed kind param
    chem_listing = _find_chemistry_entry(tlf_index["listings"])  # Fix 2.2: removed kind param

    # Check for chemistry entries — may not exist for all protocols
    has_chem_table = chem_table is not None
    has_chem_listing = chem_listing is not None

    # Avoid duplicates: if chemistry entry IS the first entry, skip it
    if has_chem_table and chem_table["number"] == first_table["number"]:
        has_chem_table = False
        logger.info("[TLF Demo] Chemistry table is same as first table — skipping duplicate")
    if has_chem_listing and chem_listing["number"] == first_listing["number"]:
        has_chem_listing = False
        logger.info("[TLF Demo] Chemistry listing is same as first listing — skipping duplicate")

    logger.info("[TLF Demo] Target items:")
    logger.info(f"  Table 1: {first_table['number']} - {first_table['title']}")
    if has_chem_table:
        logger.info(f"  Table 2: {chem_table['number']} - {chem_table['title']}")
    else:
        logger.info("  Table 2: No Clinical Chemistry table found in index — skipping")
    logger.info(f"  Listing 1: {first_listing['number']} - {first_listing['title']}")
    if has_chem_listing:
        logger.info(f"  Listing 2: {chem_listing['number']} - {chem_listing['title']}")
    else:
        logger.info("  Listing 2: No Clinical Chemistry listing found in index — skipping")

    # Step 3: Prepare text excerpts for prompts
    # Fix 2.5: Use boundary-aware truncation
    protocol_excerpt = _truncate_at_boundary(protocol_text, _MAX_PROTOCOL_CHARS_DEMO)
    sap_text = _extract_text_from_docx(sap_content)
    sap_excerpt = _truncate_at_boundary(sap_text, _MAX_SAP_CHARS_DEMO)

    all_tables: list[dict] = []
    all_listings: list[dict] = []
    call_1_failed = False
    call_2_failed = False

    # Step 4: Call 1 — First table + first listing from the index
    logger.info("[TLF Demo] Call 1: First table + first listing")
    try:
        user_msg_1 = FIRST_TLF_PROMPT.format(
            protocol_text=protocol_excerpt,
            sap_text=sap_excerpt,
            first_table_title=first_table["title"],
            first_table_population=first_table["population"],
            first_table_number=first_table["number"],
            first_listing_title=first_listing["title"],
            first_listing_population=first_listing["population"],
            first_listing_number=first_listing["number"],
        )
        response_1 = await _call_claude(DEMO_SYSTEM_PROMPT, user_msg_1)
        result_1 = _parse_json_from_response(response_1)

        # Fix 2.7: Validate before adding
        tables_1 = result_1.get("tables", [])
        listings_1 = result_1.get("listings", [])

        # Fix v23: Explicitly set numbers from SAP index (Claude may not return them)
        for t in tables_1:
            if _validate_tlf_item(t, "table"):
                t["number"] = f"Table {first_table['number']}"
                all_tables.append(t)
        for l in listings_1:
            if _validate_tlf_item(l, "listing"):
                l["number"] = f"Listing {first_listing['number']}"
                all_listings.append(l)

        logger.info(
            f"[TLF Demo] Call 1 returned: {len(tables_1)} tables, "
            f"{len(listings_1)} listings"
        )
    except Exception as e:
        logger.error(f"[TLF Demo] Call 1 failed: {e}")
        call_1_failed = True

    # Step 5: Call 2 — Clinical Chemistry table + listing (if either found in index)
    if has_chem_table or has_chem_listing:
        logger.info("[TLF Demo] Call 2: Clinical Chemistry table + listing")
        try:
            # Fix 2.6: Use complete section numbers for fallbacks
            ct = chem_table if has_chem_table else {
                "title": "Summary of Clinical Chemistry Parameters by Visit",
                "population": "Safety",
                "number": "14.3.2.1"  # Fix 2.6: Complete number, not just prefix
            }
            cl = chem_listing if has_chem_listing else {
                "title": "Laboratory Values: Clinical Chemistry",
                "population": "Safety",
                "number": "16.2.4.1"  # Fix 2.6: Complete number, not just prefix
            }

            user_msg_2 = LAB_CHEMISTRY_PROMPT.format(
                protocol_text=protocol_excerpt,
                sap_text=sap_excerpt,
                chem_table_title=ct["title"],
                chem_table_population=ct["population"],
                chem_table_number=ct["number"],
                chem_listing_title=cl["title"],
                chem_listing_population=cl["population"],
                chem_listing_number=cl["number"],
            )
            response_2 = await _call_claude(DEMO_SYSTEM_PROMPT, user_msg_2)
            result_2 = _parse_json_from_response(response_2)

            # Fix 2.7: Validate before adding
            tables_2 = result_2.get("tables", [])
            listings_2 = result_2.get("listings", [])

            # Fix v23: Explicitly set numbers from SAP index (Claude may not return them)
            for t in tables_2:
                if _validate_tlf_item(t, "table"):
                    t["number"] = f"Table {ct['number']}"
                    all_tables.append(t)
            for l in listings_2:
                if _validate_tlf_item(l, "listing"):
                    l["number"] = f"Listing {cl['number']}"
                    all_listings.append(l)

            logger.info(
                f"[TLF Demo] Call 2 returned: {len(tables_2)} tables, "
                f"{len(listings_2)} listings"
            )
        except Exception as e:
            logger.error(f"[TLF Demo] Call 2 failed: {e}")
            call_2_failed = True
    else:
        logger.info("[TLF Demo] Skipping Call 2 — no Clinical Chemistry entries in SAP index")

    # Fix 2.3: Check for empty results and raise error
    if not all_tables and not all_listings:
        error_msg = "All Claude API calls failed; no TLF shells generated."
        if call_1_failed and call_2_failed:
            error_msg = "Both Claude API calls failed. Check API credentials and connectivity."
        elif call_1_failed:
            error_msg = "Claude API call 1 failed and call 2 produced no results."
        elif call_2_failed:
            error_msg = "Claude API call 2 failed and call 1 produced no results."
        raise RuntimeError(error_msg)

    # Step 6: Preserve SAP index numbers (do NOT overwrite with assign_numbers)
    # Claude uses the exact numbers from the SAP index as instructed in the prompts.
    # Previously this called assign_numbers() which overwrote them with 14.2.1.1, 16.2.x, etc.
    all_figures: list[dict] = []

    # Just ensure section field is set for rendering (based on existing number)
    for t in all_tables:
        num = t.get("number", "")
        if "14.1" in num:
            t["section"] = "14.1"
        elif "14.2" in num:
            t["section"] = "14.2"
        elif "14.3" in num:
            t["section"] = "14.3"
        else:
            t["section"] = "14.3"  # Default to safety

    for li in all_listings:
        li["section"] = "16.2"

    total = len(all_tables) + len(all_listings)
    logger.info(
        f"[TLF Demo] Complete: {len(all_tables)} tables, "
        f"{len(all_listings)} listings ({total} total)"
    )

    # Step 7: Render output
    if output_format == "json":
        return {
            "tables": all_tables,
            "figures": [],
            "listings": all_listings,
            "index_parsed": {
                "first_table": first_table,
                "chem_table": chem_table if has_chem_table else None,
                "first_listing": first_listing,
                "chem_listing": chem_listing if has_chem_listing else None,
            },
        }
    elif output_format == "docx":
        return render_docx(all_tables, all_figures, all_listings)
    else:
        return render_markdown(all_tables, all_figures, all_listings)
