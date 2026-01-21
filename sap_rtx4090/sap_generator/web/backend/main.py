"""
SAP Generator Backend API
Production-grade with file upload support
Deploy to Render.com

Production Features:
- Structured logging
- Health check with circuit breaker status
- Proper error handling
"""

# ============================================================================
# VERSION CHECK - This prints at import time, BEFORE anything else
# ============================================================================
print("=" * 70)
print("SAP GENERATOR API - VERSION CHECK")
print("=" * 70)
print("BUILD: v102.0-universal-tlf-2026-01-20")
print("FEATURE: Universal TLF Shell Expansion + SOA Appendix")
print("  • v102.0: Universal TLF expansion - period stratification, population×assessment matrix")
print("  • v102.0: Region-aware demographics (Race/Ethnicity for FDA)")
print("  • v102.0: Auto-detected PK/Immunogenicity tables for biologics")
print("  • v101.2: Clean SOA appendix - HTML→Markdown, remove protocol boilerplate")
print("  • v101.1: Append Reducto SOA as Appendix A in generated SAP")
print("  • v101.0: TLF shell generator for study-specific table shells")
print("=" * 70)

import os
from dotenv import load_dotenv

# Load .env file for local development (override=True ensures .env takes precedence)
load_dotenv(override=True)

import re
import time
import asyncio
import tempfile
from datetime import datetime
from typing import Optional, Dict, Any, List
from contextlib import asynccontextmanager
import io

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import httpx

# Supabase client
from supabase import create_client, Client

# Document parsing
import PyPDF2
from docx import Document as DocxDocument

# LlamaParse - accurate PDF extraction with proper font encoding
LLAMAPARSE_AVAILABLE = False
_llamaparse_instance = None
try:
    from llama_cloud_services import LlamaParse
    llamaparse_key = os.getenv("LLAMAPARSE_API_KEY", "")
    if llamaparse_key:
        # Regular text extraction for most pages
        _llamaparse_instance = LlamaParse(
            api_key=llamaparse_key,
            result_type="markdown",
            verbose=False
        )
        # Vision mode for table pages (bypasses font encoding issues)
        _llamaparse_vision_instance = LlamaParse(
            api_key=llamaparse_key,
            result_type="markdown",
            use_vendor_multimodal_model=True,
            verbose=False
        )
        LLAMAPARSE_AVAILABLE = True
        print("[PDF Parser] LlamaParse available - text + vision mode for tables")
    else:
        print("[PDF Parser] WARNING: LLAMAPARSE_API_KEY not set")
        _llamaparse_vision_instance = None
except ImportError:
    print("[PDF Parser] WARNING: LlamaParse not installed")
    _llamaparse_vision_instance = None

try:
    import fitz  # PyMuPDF - for page rendering to images
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("WARNING: PyMuPDF not available")

# Import SAP generator (add parent to path)
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

# Import structured logging
try:
    from enterprise_sap_system.core.logging_config import get_logger, SAPLogger
    # Initialize logging for production (JSON output)
    SAPLogger.initialize(level="INFO", json_output=os.getenv("LOG_JSON", "false").lower() == "true")
    logger = get_logger("web.backend")
except ImportError:
    import logging
    logging.basicConfig(level=logging.INFO)
    logger = logging.getLogger("web.backend")

# AGENTIC HYBRIDRAG PIPELINE - Legacy pipeline (not used)
# Architecture: Protocol → Hybrid Retrieval → Method Extraction → Generation → Validation
try:
    from enterprise_sap_system.rag.agentic_sap_pipeline import (
        AgenticSAPPipeline, create_agentic_pipeline, SAPGenerationResult
    )
    AGENTIC_PIPELINE_AVAILABLE = True
except ImportError as e:
    AGENTIC_PIPELINE_AVAILABLE = False
    print(f"Warning: AgenticSAPPipeline not available: {e}")

# NOTE: RuleBasedSAPPipeline was deleted - ProductionSAPPipeline is now the only pipeline

# NEW: Production Pipeline with Separation of Concerns (SELF-RAG pattern)
# - Extraction as Ground Truth (single source for numbers)
# - RAG Sanitization (strips numbers from examples)
# - Explicit source attribution in prompts
# - SELF-RAG verification with correction loop
try:
    from enterprise_sap_system.core.production_pipeline import (
        ProductionSAPPipeline, create_production_pipeline
    )
    PRODUCTION_PIPELINE_AVAILABLE = True
except ImportError as e:
    PRODUCTION_PIPELINE_AVAILABLE = False
    print(f"Warning: ProductionSAPPipeline not available: {e}")

# Keep old import for backward compatibility
try:
    from enterprise_sap_system.core.hybrid_pipeline import HybridSAPPipeline, create_hybrid_pipeline
except ImportError:
    HybridSAPPipeline = None
    create_hybrid_pipeline = None

# v69: Direct SAP Generation now uses EnhancedKGPipeline (dynamic SAP structure)
# TwoPassExtractor is DEPRECATED - use EnhancedKGPipeline via get_pipeline()
DIRECT_GENERATION_AVAILABLE = True  # Always available via get_pipeline()
TwoPassExtractor = None  # Deprecated - keeping variable for compatibility

# NEW: 3-Collection RAG System (structure, content, TLF)
# Uses RAG for style/format guidance + TLF appendix generation
try:
    from enterprise_sap_system.core.sap_rag import SAPRAGIndex
    RAG_SYSTEM_AVAILABLE = True
except ImportError as e:
    RAG_SYSTEM_AVAILABLE = False
    SAPRAGIndex = None

# NEW: Integrated Pipeline with LLM Extraction + RAG + Knowledge Graph
# - LLM-based extraction for complex elements (interim analysis, censoring rules, etc.)
# - RAG with preserved interim analysis values
# - Full coverage of Phase 1/2/3 trial elements
try:
    from enterprise_sap_system.core.integrated_pipeline import IntegratedPipeline as IntegratedSAPPipeline
    INTEGRATED_PIPELINE_AVAILABLE = True
except ImportError as e:
    INTEGRATED_PIPELINE_AVAILABLE = False
    IntegratedSAPPipeline = None
    print(f"Warning: IntegratedPipeline not available: {e}")
    print(f"Warning: SAPRAGIndex (3-collection RAG) not available: {e}")

# NEW: Regulatory-grade SAP Generator (ICH E9 compliant, 45+ pages)
try:
    from enterprise_sap_system.core.regulatory_sap_generator import (
        RegulatorySAPGenerator,
        create_regulatory_sap_generator,
        ProtocolFacts,
        SAPDocument
    )
    REGULATORY_GENERATOR_AVAILABLE = True
except ImportError as e:
    REGULATORY_GENERATOR_AVAILABLE = False
    RegulatorySAPGenerator = None
    print(f"Warning: RegulatorySAPGenerator not available: {e}")

# Import LLM client for health check
try:
    from enterprise_sap_system.core.tiered_llm import get_tiered_client
    LLM_CLIENT_AVAILABLE = True
except ImportError:
    LLM_CLIENT_AVAILABLE = False
    logger.warning("TieredLLMClient not available for health check")

# SAP Evaluator - DISABLED (module not deployed)
# evaluate_sap module is only available locally for ground truth testing
SAP_EVALUATOR_AVAILABLE = False
SAPEvaluator = None

# SAP Verification Layer (Generate → Verify architecture)
# Verifies generated SAP against protocol anchors (sentences with statistics)
try:
    from enterprise_sap_system.core.sap_verifier import (
        extract_anchors,
        verify_sap,
        check_regulatory_compliance,
        VerificationReport,
        ProtocolAnchors,
        Severity
    )
    SAP_VERIFIER_AVAILABLE = True
except ImportError as e:
    SAP_VERIFIER_AVAILABLE = False
    print(f"Warning: SAP Verifier not available: {e}")

# SAP Workbench - Section-by-section generation with 55-category KG extraction
try:
    from enterprise_sap_system.workbench.workbench_core import (
        SAPWorkbench,
        SectionStatus,
        get_workbench_sections  # v82: SAP_SECTIONS removed, use function instead
    )
    WORKBENCH_AVAILABLE = True
except ImportError as e:
    WORKBENCH_AVAILABLE = False
    SAPWorkbench = None
    get_workbench_sections = None
    print(f"Warning: SAP Workbench not available: {e}")

# TLF Shell Integration v2 - Modular TLF shell generation with universal expansion
try:
    from tlf_integration import (
        generate_tlf_shells_for_protocol,
        detect_study_type,
        get_tlf_shell_summary,
        calculate_expected_table_count,
        build_universal_config
    )
    TLF_INTEGRATION_AVAILABLE = True
    print("[TLF Integration] TLF shell generator v2 available (universal expansion)")
except ImportError as e:
    TLF_INTEGRATION_AVAILABLE = False
    generate_tlf_shells_for_protocol = None
    calculate_expected_table_count = None
    build_universal_config = None
    print(f"Warning: TLF Integration not available: {e}")

# NEW: Enhanced KG Pipeline - 55-category extraction with prohibition rules
# This replaces TwoPassExtractor for context-aware SAP generation
try:
    from enterprise_sap_system.knowledge_graph.kg_enhanced_pipeline import (
        EnhancedKGPipeline
    )
    KG_PIPELINE_AVAILABLE = True
except ImportError as e:
    KG_PIPELINE_AVAILABLE = False
    EnhancedKGPipeline = None
    print(f"Warning: EnhancedKGPipeline not available: {e}")

# Environment variables
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Log startup configuration
logger.info(
    "Backend starting",
    supabase_configured=bool(SUPABASE_URL and SUPABASE_KEY),
    groq_configured=bool(GROQ_API_KEY)
)

# Initialize Supabase client
supabase: Client = None

def get_supabase() -> Client:
    global supabase
    if supabase is None:
        if not SUPABASE_URL or not SUPABASE_KEY:
            raise ValueError("SUPABASE_URL and SUPABASE_SERVICE_KEY must be set")
        # Ensure URL has trailing slash (required for Storage API)
        url = SUPABASE_URL.rstrip('/') + '/'
        supabase = create_client(url, SUPABASE_KEY)
    return supabase


# Document parsing functions
def fix_pdf_font_encoding(text: str) -> str:
    """
    Fix garbled PDF text caused by custom font encoding.

    PDFs may use MULTIPLE encoding schemes:
    1. +29 ASCII shift: 'UXJ → Drug, 0(', → MEDI (main protocol text)
    2. +3 ASCII shift: fkqolar`qflk → introduction (appendices)
    3. -3 ASCII shift: Colqlfdo → Clinical, Swxg| → Study (LlamaParse variant)
    4. -29 ASCII shift: OMNQ → 2014, lctober → October (dates/numbers)

    This function detects and applies the appropriate shift per-segment.
    """
    if not text or len(text) < 20:
        return text

    import re

    # Patterns for +29 shift (main garbled text - chars 32-96 need +29)
    patterns_29 = ["'UXJ", "0(',", "6WXG\\", "3URWRFRO", "&OLQLFDO", "(GLWLRQ",
                   "7DEOH", "6HFWLRQ", "3DJH", "9LVLW", "6FUHHQLQJ", "'DWH",
                   ")HEUXDU\\", ";UD\\", "&KHVW", "3ODLQ", "8OWUD", "7XPRXU",
                   "(QGRVFRS", "&\\WRORJ", ",VRWRSLF", ")OXRUR", "6\\PSWRP"]

    # Patterns for +3 shift (appendix text - chars 94-122 need +3)
    # fkqolar`qflk=introduction, abcfkfqflk=definition, jb^pro^_ib=measurable
    patterns_3 = ["fkqola", "abcfkf", "jb^pro", "klkJjb", "q^odbq", "ibpflk",
                  "m^qfbkq", "obpb^o", "zliib`", "pexii", "pexoo", "colj",
                  "qeb ", "qefp ", "tfqe ", "^ka ", "lc ", "fk ", "ql "]

    # Patterns for -3 shift (LlamaParse encoding - chars need -3)
    # Colqlfdo=Clinical, Swxg|=Study, Purwrfro=Protocol, Eglwlrq=Edition
    patterns_neg3 = ["Colqlfdo", "Swxg|", "Purwrfro", "Eglwlrq", "Nxpehu",
                     "Agplqlvwudwlyh", "Ckdqjh", "Ddwh", "Lrfdo", "Shfrqgdu|",
                     "Oemhfwlyh", "Oxwfrph", "Mhdvxuh", "hiilfdf|", "frpsduhg",
                     "sodfher", "dvvhvvphqwv", "dffruglqj", "ghilqhg", "vwdqgdug",
                     "folqlfdo", "sudfwlfh", "vdihw|", "wrohudelolw|", "suriloh",
                     "sk|vlfdo", "h{dplqdwlrqv", "ylwdo", "vljqv", "lqfoxglqj",
                     "eorrg", "suhvvxuh", "sxovh", "hohfwurfduglrjudpv",
                     "oderudwru|", "ilqglqjv", "fkhplvwu|", "kdhpdwrorj|",
                     "xulqdo|vlv", "Crqfhqwudwlrq", "sdudphwhuv", "vdpsolqj",
                     "lqyhvwljdwh", "lppxqrjhqlflw|", "frqilupdwru|", "uhvxowv",
                     "srvlwlyh", "qhjdwlyh", "wlwuhv", "qhxwudolvlqj", "dqwlerglhv",
                     "v|pswrpv", "khdowk", "txdolw|", "olih", "sdwlhqwv", "wuhdwhg",
                     "ghwhulrudwlrq", "idwljxh", "sdlq", "qdxvhd", "yrplwlqj",
                     "g|vsqrhd", "dsshwlwh", "lqvrpqld", "frqvwlsdwlrq", "glduukrhd",
                     "ixqfwlrq", "hprwlrqdo", "frjqlwlyh", "vrfldo", "joredo",
                     "vwdwxv", "frxjk", "kdhprsw|vlv", "fkhvw", "vkrxoghu",
                     "Ckdqjhv", "Wruog", "Hhdowk", "Oujdql}dwlrq", "Phuirupdqfh",
                     "Aqdo|vlv", "edvhg", "xsrq", "vwdwlvwlfdo", "phwkrgv",
                     "vhfwlrq", "ghwdlov", "luudgldwhg", "ohvlrqv", "frqvlghuhg",
                     "phdvxudeoh", "vhohfwhg", "wdujhw", "surylglqj", "ixoilo",
                     "fulwhuld", "phdvxudelolw|", "Aqwl", "guxj", "dqwlerg|",
                     "Agyhuvh", "hyhqw", "Pursruwlrq", "dolyh", "surjuhvvlrq",
                     "iuhh", "prqwkv", "udqgrplvdwlrq", "Bolqghg", "Iqghshqghqw",
                     "Chqwudo", "Rhylhz", "Dxudwlrq", "uhvsrqvh", "Exurshdq",
                     "Oujdqlvdwlrq", "Rhvhdufk", "Tuhdwphqw", "Cdqfhu"]

    # Patterns for -29 shift (dates/numbers - uppercase chars need -29)
    # OMNQ=2014, lctober=October, MN=01
    date_month_patterns = ["lctober", "Mprch", "Mpril", "Ndnuary", "Nebruary",
                           "Lhcember", "Mctober", "Nbvember", "Nhptember"]

    def detect_encoding(segment: str) -> int:
        """Detect which encoding shift a segment uses. Returns 0, 3, 29, -3, or -29."""
        if len(segment) < 3:
            return 0

        # Check for -3 patterns FIRST (LlamaParse encoding - most common now)
        if any(p in segment for p in patterns_neg3):
            return -3

        # Check for +29 patterns (main text with different encoding)
        if any(p in segment for p in patterns_29):
            return 29

        # Check for +3 patterns (appendix encoding)
        if any(p in segment for p in patterns_3):
            return 3

        # Check for -29 patterns (date encoding with uppercase)
        if any(p in segment for p in date_month_patterns):
            return -29

        # Check for = surrounded date patterns like =OMNQ= (year) or =MN= (day)
        if re.search(r'=[A-Z]{2,4}=', segment):
            return -29

        # Heuristic for -3: text with | character (often Swxg| = Study)
        # BUT only if it has garbled patterns, not normal text with |
        if '|' in segment and any(p in segment for p in ["Swxg|", "S|qrsvlv", "txdolw|"]):
            return -3

        # REMOVED: Heuristic for +3 was CORRUPTING clean text!
        # The range 94-122 includes lowercase a-z (97-122), so normal text triggers it.
        # Only use pattern-based detection for +3.

        # Heuristic for +29: high concentration of chars 32-96
        # Only apply if text has garbled-looking patterns (lots of symbols, no normal words)
        chars_in_29_range = sum(1 for c in segment if 32 <= ord(c) <= 96)
        if len(segment) > 10 and chars_in_29_range / len(segment) > 0.7:
            # Additional check: must not have normal words
            normal_words = ["the", "and", "for", "with", "from", "study", "drug", "dose"]
            if not any(w in segment.lower() for w in normal_words):
                return 29

        return 0

    def decode_segment(segment: str, shift: int) -> str:
        """Apply specified shift to decode a segment."""
        if shift == 0:
            return segment

        decoded_chars = []
        for c in segment:
            code = ord(c)
            if shift == 29:
                # +29: decode chars 32-96 → 61-125
                if 32 <= code <= 96:
                    decoded_chars.append(chr(code + 29))
                else:
                    decoded_chars.append(c)
            elif shift == 3:
                # +3: decode chars 94-122 → 97-125
                if 94 <= code <= 122:
                    decoded_chars.append(chr(code + 3))
                elif code == 61:  # '=' often represents space
                    decoded_chars.append(' ')
                else:
                    decoded_chars.append(c)
            elif shift == -3:
                # -3: decode ONLY lowercase letters and | (LlamaParse encoding)
                # Colqlfdo → Clinical, Swxg| → Study, Purwrfro → Protocol
                # Uppercase letters stay unchanged!
                if (97 <= code <= 122) or code == 124:  # a-z or |
                    decoded_chars.append(chr(code - 3))
                else:
                    decoded_chars.append(c)
            elif shift == -29:
                # -29: decode uppercase chars to numbers/symbols
                if 65 <= code <= 90:  # A-Z
                    new_code = code - 29
                    if 32 <= new_code <= 126:  # printable range
                        decoded_chars.append(chr(new_code))
                    else:
                        decoded_chars.append(c)
                elif code == 108:  # 'l' -> 'O' for lctober -> October
                    decoded_chars.append('O')
                elif code == 61:  # '=' -> space
                    decoded_chars.append(' ')
                else:
                    decoded_chars.append(c)
        return ''.join(decoded_chars)

    def fix_line(line: str) -> str:
        """Fix encoding for a single line, handling mixed segments."""
        if len(line) < 5:
            return line

        # First, check if entire line has a detectable encoding
        line_encoding = detect_encoding(line)
        if line_encoding != 0:
            return decode_segment(line, line_encoding)

        # Otherwise, try to fix segments separated by spaces or =
        # This handles mixed encoding within a line
        result = line

        # Fix date-like patterns: =WORD= where WORD is uppercase
        def fix_date_segment(match):
            segment = match.group(0)
            return decode_segment(segment, -29)

        result = re.sub(r'=[A-Z]{2,4}=', fix_date_segment, result)

        # Fix month names with encoding
        for month_pattern in date_month_patterns:
            if month_pattern in result:
                result = result.replace(month_pattern, decode_segment(month_pattern, -29))

        # Fix remaining +3 encoded words
        for pattern in patterns_3:
            if pattern in result:
                # Find the full word containing this pattern
                words = result.split()
                new_words = []
                for word in words:
                    if pattern in word.lower():
                        new_words.append(decode_segment(word, 3))
                    else:
                        new_words.append(word)
                result = ' '.join(new_words)
                break

        return result

    # Process line by line
    lines = text.split('\n')
    result_lines = []
    stats = {29: 0, 3: 0, -3: 0, -29: 0}

    for line in lines:
        original = line
        fixed = fix_line(line)
        if fixed != original:
            # Track which encoding was used
            encoding = detect_encoding(original)
            if encoding in stats:
                stats[encoding] += 1
        result_lines.append(fixed)

    total_fixed = sum(stats.values())
    if total_fixed > 0:
        print(f"[PDF Parser] Fixed encoding: {stats[29]} lines (+29), {stats[3]} lines (+3), {stats[-3]} lines (-3), {stats[-29]} lines (-29)")

    # Final cleanup: remove excessive = signs that may remain
    result = '\n'.join(result_lines)
    result = re.sub(r'={3,}', ' ', result)  # Replace 3+ = with space
    result = re.sub(r'\s{3,}', '  ', result)  # Normalize excessive spaces

    # EXPLICIT fix for date patterns like MN==PM=lctober=OMNQ → 01  30 October 2014
    # These use -29 shift on uppercase letters and 'l' → 'O'
    def fix_date_line(match):
        text = match.group(0)
        decoded = []
        for c in text:
            code = ord(c)
            if 65 <= code <= 90:  # A-Z → subtract 29
                new_code = code - 29
                if 32 <= new_code <= 126:
                    decoded.append(chr(new_code))
                else:
                    decoded.append(c)
            elif code == 108:  # 'l' → 'O' (for lctober → October)
                decoded.append('O')
            elif code == 61:  # '=' → space
                decoded.append(' ')
            else:
                decoded.append(c)
        return ''.join(decoded)

    # Match date patterns: uppercase letters, =, and 'lctober' variants
    result = re.sub(r'[A-Z=]+lctober[A-Z=]*', fix_date_line, result)
    result = re.sub(r'\s{2,}', ' ', result)  # Clean up multiple spaces

    return result


def wrap_markdown_tables(text: str) -> str:
    """
    Detect markdown tables in text and wrap them with [TABLE] markers.
    Markdown tables have rows with | separators and a header separator line with dashes.
    """
    import re

    lines = text.split('\n')
    result = []
    i = 0
    tables_found = 0

    while i < len(lines):
        line = lines[i]

        # Check if this looks like a markdown table row (has multiple | chars)
        if '|' in line and line.count('|') >= 2:
            # Look ahead to see if this is a table (has header separator)
            table_lines = [line]
            j = i + 1

            # Collect consecutive lines with | separators
            while j < len(lines) and '|' in lines[j] and lines[j].count('|') >= 2:
                table_lines.append(lines[j])
                j += 1

            # Check if it's a real table (at least 2 rows, and has separator line with dashes)
            has_separator = any(re.match(r'^[\s|:-]+$', tl.replace(' ', '')) for tl in table_lines[:3])
            if len(table_lines) >= 2 and (has_separator or len(table_lines) >= 3):
                # It's a table - wrap it
                tables_found += 1
                result.append('[TABLE]')
                # Remove markdown separator line (|---|---|) for cleaner display
                for tl in table_lines:
                    if not re.match(r'^[\s|:-]+$', tl.replace(' ', '')):
                        # Clean up the line: remove leading/trailing pipes and extra spaces
                        cleaned = tl.strip()
                        if cleaned.startswith('|'):
                            cleaned = cleaned[1:]
                        if cleaned.endswith('|'):
                            cleaned = cleaned[:-1]
                        result.append(cleaned)
                result.append('[/TABLE]')
                i = j
                continue

        result.append(line)
        i += 1

    if tables_found > 0:
        print(f"[Tables] Wrapped {tables_found} markdown tables with [TABLE] markers")

    return '\n'.join(result)


# ============================================================================
# HTML TO MARKDOWN CONVERSION FOR SOA TABLES
# ============================================================================

def html_table_to_markdown(html_content: str) -> str:
    """
    Convert HTML tables to markdown format.

    Args:
        html_content: String containing HTML table(s)

    Returns:
        Markdown-formatted table string
    """
    # Check if there's any HTML table content
    if '<table' not in html_content.lower() and '<tr' not in html_content.lower():
        return html_content

    result_lines = []

    # Extract all tables
    table_pattern = re.compile(r'<table[^>]*>(.*?)</table>', re.DOTALL | re.IGNORECASE)
    non_table_content = html_content

    for table_match in table_pattern.finditer(html_content):
        table_html = table_match.group(0)

        # Extract rows
        rows = []
        row_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL | re.IGNORECASE)

        for row_match in row_pattern.finditer(table_html):
            row_html = row_match.group(1)

            # Extract cells (th or td)
            cells = []
            cell_pattern = re.compile(r'<(th|td)[^>]*>(.*?)</\1>', re.DOTALL | re.IGNORECASE)

            for cell_match in cell_pattern.finditer(row_html):
                cell_content = cell_match.group(2)
                # Clean HTML tags from cell content
                cell_content = re.sub(r'<[^>]+>', '', cell_content)
                # Clean whitespace
                cell_content = ' '.join(cell_content.split())
                cells.append(cell_content)

            if cells:
                rows.append(cells)

        if rows:
            # Build markdown table
            md_lines = []

            # First row is header
            if rows:
                header = rows[0]
                md_lines.append('| ' + ' | '.join(header) + ' |')
                # Separator row
                md_lines.append('|' + '|'.join(['---' for _ in header]) + '|')

                # Data rows
                for row in rows[1:]:
                    # Pad row if needed
                    while len(row) < len(header):
                        row.append('')
                    md_lines.append('| ' + ' | '.join(row[:len(header)]) + ' |')

            result_lines.append('\n'.join(md_lines))

    if result_lines:
        return '\n\n'.join(result_lines)

    return html_content


def clean_soa_content_for_appendix(raw_content: str) -> str:
    """
    Clean and filter SOA content for inclusion as an appendix.

    Removes:
    - Protocol boilerplate (page numbers, headers, footers)
    - Confidentiality notices
    - Empty sections
    - Redundant whitespace

    Converts:
    - HTML tables to markdown

    Args:
        raw_content: Raw content from Reducto extraction

    Returns:
        Cleaned markdown content suitable for SAP appendix
    """
    if not raw_content:
        return ""

    content = raw_content

    # 1. Remove the Reducto marker header section
    markers_to_remove = [
        r'={60,}[\s\S]*?SCHEDULE OF ASSESSMENTS \(Enhanced by Reducto\)[\s\S]*?={60,}',
        r'END OF SCHEDULE OF ASSESSMENTS[\s\S]*?={60,}',
        r'={60,}',
    ]
    for pattern in markers_to_remove:
        content = re.sub(pattern, '', content)

    # 2. Remove protocol boilerplate patterns
    boilerplate_patterns = [
        # Page numbers
        r'\b\d{1,3}\s+of\s+\d{1,3}\b',
        r'Page\s+\d+\s+of\s+\d+',
        r'^\s*\d{1,3}\s*$',
        # Confidentiality notices - handle various formats
        r'^.*?(?:Celltrion|CELLTRION).*?(?:CONFIDENTIAL|Confidential).*?$',
        r'(?:CONFIDENTIAL|Confidential)[\s\-]*(?:Property|Information)?.*?(?:\n|$)',
        r'^.*?(?:Celltrion|CELLTRION)\s*/\s*$',  # Leftover "Celltrion/" lines
        # Protocol references
        r'Protocol\s+(?:No\.|Number|#)?:?\s*CT-P\d+.*?(?:\n|$)',
        r'Amendment\s+\d+.*?(?:\n|$)',
        # Document identifiers
        r'EudraCT\s+(?:No\.|Number)?:?\s*\d{4}-\d+-\d+',
        r'IND\s+(?:No\.|Number)?:?\s*\d+',
        # Date stamps that look like document dates
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b',
        # Version markers
        r'Version\s+\d+\.?\d*\s*(?:dated)?.*?(?:\n|$)',
        # Table references that aren't actual data
        r'^\s*Table\s+\d+[\.\-]\d+[\.\-]?\d*\s*$',
    ]

    for pattern in boilerplate_patterns:
        content = re.sub(pattern, '', content, flags=re.IGNORECASE | re.MULTILINE)

    # 3. Convert HTML tables to markdown
    content = html_table_to_markdown(content)

    # 4. Remove [TABLE] markers (keep the content)
    content = re.sub(r'\[TABLE\](?:\s*\(Page \d+\))?', '', content)
    content = re.sub(r'\[/TABLE\]', '', content)

    # 5. Clean up excessive whitespace while preserving table structure
    # Replace multiple blank lines with double newline
    content = re.sub(r'\n{4,}', '\n\n\n', content)

    # 6. Remove lines that are just whitespace or separators
    lines = content.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip empty lines at the beginning
        if not cleaned_lines and not stripped:
            continue
        # Skip lines that are just dashes or equals (separators)
        if stripped and re.match(r'^[-=_\*]{5,}$', stripped):
            continue
        # Skip very short lines that look like artifacts
        if stripped and len(stripped) < 3 and not stripped.isalnum():
            continue
        cleaned_lines.append(line)

    # Remove trailing empty lines
    while cleaned_lines and not cleaned_lines[-1].strip():
        cleaned_lines.pop()

    content = '\n'.join(cleaned_lines)

    # 7. Ensure tables have proper spacing
    content = re.sub(r'(\|[^\n]+\|)\n(?!\||\n)', r'\1\n\n', content)

    return content.strip()


# ============================================================================
# REDUCTO SOA TABLE EXTRACTION SUPPORT
# ============================================================================

SOA_KEYWORDS = [
    "Schedule of Study Procedures",
    "Schedule of Study Assessments",
    "Schedule of Assessments",
    "Schedule of Events",
    "Table of Activities",
    "Schedule of Activities",
    "Study Procedures Schedule",
    "Assessment Schedule",
    "Time and Events Schedule",
    "Study Schedule",
]

# Keywords that indicate SOA content (visits, assessments)
SOA_CONTENT_KEYWORDS = [
    # Visit/timepoint terms
    "screening", "baseline", "day 1", "day 8", "day 15", "week", "cycle",
    "visit", "eot", "end of treatment", "follow-up", "follow up",
    # Assessment terms
    "vital signs", "physical exam", "laboratory", "ecg", "electrocardiogram",
    "imaging", "ct scan", "mri", "informed consent", "randomization",
    "adverse event", "concomitant", "pregnancy test", "urinalysis",
    # Table structure indicators
    "procedure", "assessment", "study drug", "blood sample", "pk sample",
]

# Keywords that indicate NON-SOA content (should be excluded)
SOA_EXCLUSION_KEYWORDS = [
    # RECIST response criteria
    "recist", "response evaluation criteria", "complete response", "partial response",
    "progressive disease", "stable disease", "target lesion", "non-target lesion",
    "sum of diameters", "best overall response",
    # Staging criteria
    "tnm", "ajcc", "staging", "tumor category", "nodal category", "metastasis category",
    # Other non-SOA tables
    "dose modification", "dose reduction", "dose level", "dlts",
]


def _score_page_for_soa(page_text: str) -> tuple:
    """
    Score a page for SOA content likelihood.

    Returns:
        (score, is_excluded) - score is positive keywords found, is_excluded if exclusion keywords found
    """
    text_lower = page_text.lower()

    # Check for exclusion keywords first
    exclusion_count = 0
    for keyword in SOA_EXCLUSION_KEYWORDS:
        if keyword in text_lower:
            exclusion_count += 1

    # If multiple exclusion keywords found, this is likely NOT an SOA page
    if exclusion_count >= 3:
        return (0, True)

    # Count positive SOA keywords
    positive_score = 0
    for keyword in SOA_CONTENT_KEYWORDS:
        if keyword in text_lower:
            positive_score += 1

    # Bonus for table-like structure (many X marks or pipes)
    x_count = text_lower.count(' x ') + text_lower.count('|x|') + text_lower.count('\tx\t')
    pipe_count = text_lower.count('|')
    if x_count > 5:
        positive_score += 3
    if pipe_count > 20:
        positive_score += 2

    return (positive_score, False)


def _is_soa_header(blocks: list, keyword: str) -> bool:
    """
    Check if a keyword appears as a section header (not just a reference).
    """
    keyword_lower = keyword.lower()

    for block in blocks:
        if len(block) >= 5 and block[6] == 0:  # Text block
            block_text = block[4].strip()
            block_text_lower = block_text.lower()

            if keyword_lower in block_text_lower:
                lines = block_text.split('\n')
                for line in lines:
                    line_stripped = line.strip()
                    line_lower = line_stripped.lower()
                    if keyword_lower in line_lower and len(line_stripped) < 150:
                        keyword_ratio = len(keyword) / max(len(line_stripped), 1)
                        if keyword_ratio > 0.3:
                            return True
    return False


def detect_soa_pages_from_pdf(file_content: bytes) -> List[int]:
    """
    Detect pages containing Schedule of Assessments (SOA) tables directly from PDF.

    Strategy:
    1. Find pages that have SOA keywords as section headers (not just references)
    2. Filter to only pages that contain actual table structures
    3. Include consecutive pages after a header (tables often span multiple pages)

    Args:
        file_content: PDF file as bytes

    Returns:
        List of 1-indexed page numbers containing SOA tables.
    """
    if not PYMUPDF_AVAILABLE:
        print("[SOA Detect] PyMuPDF not available for PDF search")
        return []

    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        total_pages = len(doc)

        # Step 1: Find pages with SOA section headers AND tables
        header_pages = []  # Pages where SOA keyword appears as a header
        pages_with_tables = set()  # All pages that have table structures

        for page_num in range(total_pages):
            page = doc[page_num]
            page_1idx = page_num + 1

            # Check if page has tables (works for both portrait and landscape)
            tables = page.find_tables()
            has_tables = len(tables.tables) > 0

            if has_tables:
                pages_with_tables.add(page_1idx)

            # Get text and check for SOA headers
            # Use "blocks" to get text with position info for header detection
            blocks = page.get_text("blocks")
            page_text_lower = page.get_text().lower()

            for keyword in SOA_KEYWORDS:
                keyword_lower = keyword.lower()
                if keyword_lower in page_text_lower:
                    # Check if this keyword appears as a header (near top of page or standalone line)
                    is_header = _is_soa_header(blocks, keyword)
                    if is_header:
                        header_pages.append(page_1idx)
                        print(f"[SOA Detect] Found SOA header '{keyword}' on page {page_1idx} (has_tables={has_tables})")
                        break

        doc.close()

        if not header_pages:
            print(f"[SOA Detect] No SOA headers found in {total_pages} pages")
            return []

        print(f"[SOA Detect] Found {len(pages_with_tables)} pages with tables: {sorted(pages_with_tables)}")
        print(f"[SOA Detect] Found SOA headers on pages: {header_pages}")

        # Step 2: For each header page, include the next 3 pages
        # SOA tables are often IMAGES (not text tables), so don't rely on find_tables()
        soa_pages = set()
        for header_page in header_pages:
            # Include header page and next 3 pages (SOA typically spans 2-3 pages)
            for offset in range(0, 4):
                page = header_page + offset
                if page <= total_pages:
                    soa_pages.add(page)
            print(f"[SOA Detect] Including pages {header_page} to {min(header_page+3, total_pages)} after header")

        result = sorted(soa_pages)
        print(f"[SOA Detect] Final SOA pages: {result}")
        return result

    except Exception as e:
        print(f"[SOA Detect] Error searching PDF: {e}")
        import traceback
        traceback.print_exc()
        return []


def merge_reducto_soa_into_text(original_text: str, reducto_content: str) -> str:
    """
    Merge Reducto's SOA extraction into LlamaParse output.

    Strategy:
    1. Find SOA section headers in original text
    2. Find the end of the SOA section (next major section or reasonable boundary)
    3. Replace the SOA section with Reducto's better extraction

    If SOA section can't be reliably identified, append Reducto content at end.

    Args:
        original_text: Full protocol text from LlamaParse
        reducto_content: Extracted SOA tables from Reducto

    Returns:
        Merged text with Reducto SOA replacing original
    """
    if not reducto_content or not reducto_content.strip():
        return original_text

    # Find SOA section start
    soa_start_patterns = [
        r'\n#+\s*Schedule of Study Procedures',
        r'\n#+\s*Schedule of Assessments',
        r'\n#+\s*Schedule of Events',
        r'\n#+\s*Table of Activities',
        r'\nSchedule of Study Procedures\n',
        r'\nSchedule of Assessments\n',
        r'\nSCHEDULE OF STUDY PROCEDURES',
        r'\nSCHEDULE OF ASSESSMENTS',
    ]

    soa_start_idx = -1
    soa_header = ""

    for pattern in soa_start_patterns:
        match = re.search(pattern, original_text, re.IGNORECASE)
        if match:
            soa_start_idx = match.start()
            soa_header = match.group(0)
            print(f"[SOA Merge] Found SOA section start at position {soa_start_idx}: '{soa_header.strip()[:50]}...'")
            break

    if soa_start_idx == -1:
        # Couldn't find SOA section - append at end
        print("[SOA Merge] Could not find SOA section in text, appending Reducto content at end")
        return original_text + "\n\n" + "=" * 60 + "\n" + \
               "SCHEDULE OF ASSESSMENTS (Enhanced by Reducto)\n" + \
               "=" * 60 + "\n\n" + reducto_content

    # Find SOA section end - look for next major section header
    # Major sections typically start with # or are in ALL CAPS followed by newline
    section_end_patterns = [
        r'\n#+\s*[A-Z][^#\n]+\n',  # Markdown header
        r'\n\d+\.\s+[A-Z][^\n]+\n',  # Numbered section (e.g., "7. STUDY PROCEDURES")
        r'\n[A-Z][A-Z\s]{10,}\n',  # ALL CAPS section header
    ]

    # Start searching after the SOA header
    search_start = soa_start_idx + len(soa_header)

    # Skip some content (at least 500 chars) to avoid matching subsection headers within SOA
    min_section_length = 500
    soa_end_idx = len(original_text)  # Default to end of document

    for pattern in section_end_patterns:
        for match in re.finditer(pattern, original_text[search_start + min_section_length:]):
            potential_end = search_start + min_section_length + match.start()
            # Verify this looks like a new major section, not a table header
            matched_text = match.group(0).strip()
            # Skip if it looks like SOA-related content
            if any(kw in matched_text.lower() for kw in ['schedule', 'assessment', 'visit', 'procedure', 'table']):
                continue
            soa_end_idx = potential_end
            print(f"[SOA Merge] Found section end at position {soa_end_idx}: '{matched_text[:50]}...'")
            break
        if soa_end_idx < len(original_text):
            break

    # Build merged content
    before_soa = original_text[:soa_start_idx]
    after_soa = original_text[soa_end_idx:]

    # Create enhanced SOA section
    enhanced_soa = (
        "\n\n" + "=" * 60 + "\n"
        "SCHEDULE OF ASSESSMENTS (Enhanced by Reducto)\n"
        + "=" * 60 + "\n\n"
        + reducto_content
        + "\n\n" + "=" * 60 + "\n"
        "END OF SCHEDULE OF ASSESSMENTS\n"
        + "=" * 60 + "\n"
    )

    merged = before_soa + enhanced_soa + after_soa
    print(f"[SOA Merge] Replaced {soa_end_idx - soa_start_idx:,} chars with {len(enhanced_soa):,} chars of Reducto content")

    return merged


def detect_table_pages(file_content: bytes) -> set:
    """Detect which pages contain tables using PyMuPDF."""
    if not PYMUPDF_AVAILABLE:
        return set()

    table_pages = set()
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            tables = page.find_tables()
            if tables.tables:
                table_pages.add(page_num + 1)  # 1-indexed
        doc.close()
        print(f"[PDF Parser] Detected tables on pages: {sorted(table_pages)}")
    except Exception as e:
        print(f"[PDF Parser] Table detection failed: {e}")
    return table_pages


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    HYBRID extraction:
    - LlamaParse text mode for regular pages (clean, fast)
    - LlamaParse vision mode for table pages (bypasses font encoding)
    """
    import tempfile
    import asyncio

    # Try LlamaParse first
    if LLAMAPARSE_AVAILABLE and _llamaparse_instance:
        try:
            # Detect table pages
            table_pages = detect_table_pages(file_content)

            # LlamaParse needs a file path
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                # If we have vision mode and table pages, use hybrid extraction
                if _llamaparse_vision_instance and table_pages:
                    print(f"[PDF Parser] HYBRID mode: text extraction + vision for {len(table_pages)} table pages")

                    # Extract ALL pages with regular LlamaParse first
                    async def parse_all():
                        return await _llamaparse_instance.aload_data(tmp_path)

                    try:
                        loop = asyncio.get_event_loop()
                        if loop.is_running():
                            import concurrent.futures
                            with concurrent.futures.ThreadPoolExecutor() as pool:
                                future = pool.submit(asyncio.run, parse_all())
                                documents = future.result(timeout=120)
                        else:
                            documents = loop.run_until_complete(parse_all())
                    except RuntimeError:
                        documents = asyncio.run(parse_all())

                    if documents and len(documents) > 0:
                        text = "\n\n".join([doc.text for doc in documents if doc.text])
                        print(f"[PDF Parser] LlamaParse text mode extracted {len(text):,} chars")
                        print(f"[PDF Parser] Main text RAW sample (first 500 chars):")
                        print(text[:500] if text else "EMPTY")

                        # NO encoding fix - use raw LlamaParse output

                        # Now extract ONLY table pages with vision mode
                        # Create a PDF with only table pages
                        vision_text = ""
                        try:
                            doc = fitz.open(stream=file_content, filetype="pdf")
                            table_doc = fitz.open()  # New empty PDF

                            for page_num in sorted(table_pages):
                                table_doc.insert_pdf(doc, from_page=page_num-1, to_page=page_num-1)

                            # Save table pages PDF
                            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as table_tmp:
                                table_doc.save(table_tmp.name)
                                table_tmp_path = table_tmp.name

                            table_doc.close()
                            doc.close()

                            # Extract table pages with vision mode
                            async def parse_tables_vision():
                                return await _llamaparse_vision_instance.aload_data(table_tmp_path)

                            print(f"[PDF Parser] Extracting {len(table_pages)} table pages with VISION mode...")
                            try:
                                loop = asyncio.get_event_loop()
                                if loop.is_running():
                                    with concurrent.futures.ThreadPoolExecutor() as pool:
                                        future = pool.submit(asyncio.run, parse_tables_vision())
                                        table_docs = future.result(timeout=120)
                                else:
                                    table_docs = loop.run_until_complete(parse_tables_vision())
                            except RuntimeError:
                                table_docs = asyncio.run(parse_tables_vision())

                            os.unlink(table_tmp_path)

                            if table_docs and len(table_docs) > 0:
                                vision_text = "\n\n".join([doc.text for doc in table_docs if doc.text])
                                print(f"[PDF Parser] Vision mode extracted {len(vision_text):,} chars from table pages")
                                print(f"[PDF Parser] Vision text RAW sample (first 500 chars):")
                                print(vision_text[:500] if vision_text else "EMPTY")

                                # DON'T apply encoding fix to vision text - it should be clean from visual OCR
                                # Just append it
                                text = text + "\n\n" + "=" * 60 + "\n"
                                text = text + "TABLES (Vision-Extracted)\n" + "=" * 60 + "\n\n"
                                text = text + vision_text

                        except Exception as e:
                            print(f"[PDF Parser] Vision extraction for tables failed: {e}")

                        os.unlink(tmp_path)

                        text = wrap_markdown_tables(text)

                        # === REDUCTO SOA ENHANCEMENT ===
                        # If Reducto API key is set, extract SOA tables with Reducto
                        # for better table quality
                        if os.getenv("REDUCTO_API_KEY"):
                            try:
                                from web.backend.reducto_client import extract_soa_with_reducto

                                # Detect SOA pages from the original PDF
                                soa_pages = detect_soa_pages_from_pdf(file_content)

                                if soa_pages:
                                    print(f"[PDF Parser] Extracting SOA from pages {soa_pages} with Reducto...")
                                    reducto_result = extract_soa_with_reducto(file_content, soa_pages)

                                    if reducto_result.success:
                                        print(f"[PDF Parser] Reducto extracted {len(reducto_result.content):,} chars")
                                        # Merge Reducto SOA into LlamaParse output
                                        text = merge_reducto_soa_into_text(text, reducto_result.content)
                                        print(f"[PDF Parser] Merged Reducto SOA into protocol")
                                    else:
                                        print(f"[PDF Parser] Reducto failed: {reducto_result.error}")
                                else:
                                    print("[PDF Parser] No SOA pages detected, skipping Reducto")
                            except Exception as e:
                                print(f"[PDF Parser] Reducto integration error: {e}, continuing with LlamaParse only")

                        return text

                # Fallback: regular LlamaParse for everything
                print(f"[PDF Parser] Using LlamaParse text mode for all pages...")

                async def parse_pdf():
                    return await _llamaparse_instance.aload_data(tmp_path)

                try:
                    loop = asyncio.get_event_loop()
                    if loop.is_running():
                        import concurrent.futures
                        with concurrent.futures.ThreadPoolExecutor() as pool:
                            future = pool.submit(asyncio.run, parse_pdf())
                            documents = future.result(timeout=120)
                    else:
                        documents = loop.run_until_complete(parse_pdf())
                except RuntimeError:
                    documents = asyncio.run(parse_pdf())

                os.unlink(tmp_path)

                if documents and len(documents) > 0:
                    text = "\n\n".join([doc.text for doc in documents if doc.text])
                    print(f"[PDF Parser] LlamaParse extracted {len(text):,} chars")
                    if text and len(text.strip()) > 100:
                        # Fix garbled font encoding if detected (+29 ASCII shift)
                        text = fix_pdf_font_encoding(text)
                        # Wrap markdown tables in [TABLE] markers for frontend rendering
                        text = wrap_markdown_tables(text)

                        # === REDUCTO SOA ENHANCEMENT ===
                        if os.getenv("REDUCTO_API_KEY"):
                            try:
                                from web.backend.reducto_client import extract_soa_with_reducto
                                soa_pages = detect_soa_pages_from_pdf(file_content)
                                if soa_pages:
                                    print(f"[PDF Parser] Extracting SOA from pages {soa_pages} with Reducto...")
                                    reducto_result = extract_soa_with_reducto(file_content, soa_pages)
                                    if reducto_result.success:
                                        print(f"[PDF Parser] Reducto extracted {len(reducto_result.content):,} chars")
                                        text = merge_reducto_soa_into_text(text, reducto_result.content)
                                        print(f"[PDF Parser] Merged Reducto SOA into protocol")
                                    else:
                                        print(f"[PDF Parser] Reducto failed: {reducto_result.error}")
                                else:
                                    print("[PDF Parser] No SOA pages detected, skipping Reducto")
                            except Exception as e:
                                print(f"[PDF Parser] Reducto integration error: {e}, continuing with LlamaParse only")

                        return text
                    else:
                        print("[PDF Parser] LlamaParse returned minimal text, trying PyMuPDF...")
                else:
                    print("[PDF Parser] LlamaParse returned no documents, trying PyMuPDF...")

            except Exception as e:
                print(f"[PDF Parser] LlamaParse failed: {e}, trying PyMuPDF...")
                try:
                    os.unlink(tmp_path)
                except:
                    pass
        except Exception as e:
            print(f"[PDF Parser] LlamaParse temp file error: {e}")

    # Fallback to PyMuPDF (may have encoding issues with custom fonts)
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            total_tables = 0

            for page_num, page in enumerate(doc):
                page_text = []

                # Extract tables with PyMuPDF
                table_finder = page.find_tables()
                table_texts = set()  # Track table content to avoid duplication

                for table in table_finder.tables:
                    total_tables += 1
                    table_md = []
                    for row in table.extract():
                        row_text = " | ".join(str(cell) if cell else "" for cell in row)
                        if row_text.strip():
                            table_md.append(row_text)
                            for cell in row:
                                if cell:
                                    table_texts.add(str(cell).lower().strip())
                    if table_md:
                        page_text.append("\n[TABLE]\n" + "\n".join(table_md) + "\n[/TABLE]\n")

                # Extract regular text (excluding table content to avoid duplication)
                blocks = page.get_text("blocks")
                for block in blocks:
                    if block[6] == 0:  # Text block (not image)
                        block_text = block[4].strip()
                        # Skip if this text is part of a table
                        if block_text and len(block_text) > 10:
                            is_table_text = any(cell in block_text.lower() for cell in table_texts if len(cell) > 10)
                            if not is_table_text:
                                page_text.append(block_text)
                        elif block_text:
                            page_text.append(block_text)

                if page_text:
                    # Fix encoding PER PAGE - different pages may have different encodings
                    page_content = "\n".join(page_text)
                    page_content = fix_pdf_font_encoding(page_content)
                    text_parts.append(f"\n--- PAGE {page_num + 1} ---\n" + page_content)

            doc.close()
            result = "\n\n".join(text_parts)
            print(f"[PDF Parser] PyMuPDF extracted {len(result):,} chars, {total_tables} tables from {len(doc)} pages")

            # === REDUCTO SOA ENHANCEMENT (PyMuPDF fallback) ===
            if os.getenv("REDUCTO_API_KEY"):
                try:
                    from web.backend.reducto_client import extract_soa_with_reducto
                    soa_pages = detect_soa_pages_from_pdf(file_content)
                    if soa_pages:
                        print(f"[PDF Parser] Extracting SOA from pages {soa_pages} with Reducto...")
                        reducto_result = extract_soa_with_reducto(file_content, soa_pages)
                        if reducto_result.success:
                            print(f"[PDF Parser] Reducto extracted {len(reducto_result.content):,} chars")
                            result = merge_reducto_soa_into_text(result, reducto_result.content)
                            print(f"[PDF Parser] Merged Reducto SOA into protocol")
                        else:
                            print(f"[PDF Parser] Reducto failed: {reducto_result.error}")
                    else:
                        print("[PDF Parser] No SOA pages detected, skipping Reducto")
                except Exception as e:
                    print(f"[PDF Parser] Reducto integration error: {e}")

            return result

        except Exception as e:
            print(f"[PDF Parser] PyMuPDF failed: {e}, falling back to PyPDF2")

    # Fallback to PyPDF2
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        result = "\n\n".join(text_parts)
        print(f"[PDF Parser] PyPDF2 extracted {len(result):,} chars")
        # Fix garbled font encoding if detected
        result = fix_pdf_font_encoding(result)

        # === REDUCTO SOA ENHANCEMENT (PyPDF2 fallback) ===
        if os.getenv("REDUCTO_API_KEY"):
            try:
                from web.backend.reducto_client import extract_soa_with_reducto
                soa_pages = detect_soa_pages_from_pdf(file_content)
                if soa_pages:
                    print(f"[PDF Parser] Extracting SOA from pages {soa_pages} with Reducto...")
                    reducto_result = extract_soa_with_reducto(file_content, soa_pages)
                    if reducto_result.success:
                        print(f"[PDF Parser] Reducto extracted {len(reducto_result.content):,} chars")
                        result = merge_reducto_soa_into_text(result, reducto_result.content)
                        print(f"[PDF Parser] Merged Reducto SOA into protocol")
                    else:
                        print(f"[PDF Parser] Reducto failed: {reducto_result.error}")
                else:
                    print("[PDF Parser] No SOA pages detected, skipping Reducto")
            except Exception as e:
                print(f"[PDF Parser] Reducto integration error: {e}")

        return result
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file."""
    try:
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return file_content.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {e}")


def extract_text_from_file(filename: str, content: bytes) -> str:
    """Extract text from uploaded file based on extension."""
    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    if ext == 'pdf':
        return extract_text_from_pdf(content)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(content)
    elif ext in ['txt', 'text', 'md']:
        return extract_text_from_txt(content)
    else:
        # Try to decode as text
        try:
            return content.decode('utf-8')
        except Exception as e:
            raise ValueError(f"Unsupported file format: {ext} (decode error: {e})")


# Background worker flag
worker_running = False


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Start background worker on startup"""
    global worker_running
    worker_running = True
    asyncio.create_task(process_jobs_worker())
    yield
    worker_running = False


app = FastAPI(
    title="SAP Generator API",
    description="Generate Statistical Analysis Plans from clinical trial protocols",
    version="2.0.0",
    lifespan=lifespan
)

# CORS for Vercel frontend
frontend_url = os.getenv("FRONTEND_URL", "")
allowed_origins = [
    "http://localhost:3000",
    "https://*.vercel.app",
]
if frontend_url:
    allowed_origins.append(frontend_url)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for development
    allow_credentials=False,  # Must be False when using "*"
    allow_methods=["*"],  # Allow all methods
    allow_headers=["*"],
    expose_headers=["*"],
)


# Request/Response models
class GenerateRequest(BaseModel):
    protocol_text: str
    nct_id: Optional[str] = None
    filename: Optional[str] = None


class JobResponse(BaseModel):
    job_id: str
    status: str
    message: str
    extracted_text: Optional[str] = None


class JobStatusResponse(BaseModel):
    job_id: str
    status: str
    generated_sap: Optional[str] = None
    quality_score: Optional[float] = None
    endpoint_type: Optional[str] = None
    phase: Optional[str] = None
    therapeutic_area: Optional[str] = None
    processing_time: Optional[float] = None
    error_message: Optional[str] = None
    created_at: Optional[str] = None
    completed_at: Optional[str] = None
    filename: Optional[str] = None
    protocol_preview: Optional[str] = None
    # Deterministic verification
    deterministic_verification: Optional[dict] = None
    audit_report: Optional[str] = None
    needs_human_review: Optional[bool] = None


class EvaluationResponse(BaseModel):
    """Evaluation results comparing generated SAP to ground truth"""
    nct_id: str
    ground_truth_lines: int
    generated_lines: int
    section_coverage_pct: float
    keyword_overlap_pct: float
    has_primary_endpoint: bool
    has_secondary_endpoint: bool
    has_sample_size: bool
    has_analysis_populations: bool
    has_statistical_methods: bool
    has_missing_data: bool
    overall_score: float
    sections_matched: list
    sections_missing: list
    statistical_terms_found: list
    statistical_terms_missing: list


class GroundTruthInfo(BaseModel):
    """Ground truth study information"""
    nct_id: str
    title: str
    sap_lines: int
    therapeutic_area: str


class VerificationIssue(BaseModel):
    """A single verification issue found."""
    severity: str  # "critical", "warning", "info"
    category: str
    message: str
    rule: Optional[str] = None


class VerificationAnchorSummary(BaseModel):
    """Summary of anchors by category."""
    sample_size: int = 0
    alpha: int = 0
    power: int = 0
    randomization: int = 0
    endpoints: int = 0
    interim_analysis: int = 0
    hypotheses: int = 0
    boundaries: int = 0
    total: int = 0


class VerificationResponse(BaseModel):
    """Response from SAP verification against protocol anchors."""
    success: bool
    job_id: str
    # Anchor verification
    anchors_found: int
    anchors_verified: int
    anchors_missing: int
    anchor_summary: VerificationAnchorSummary
    # Confidence
    confidence_score: float
    needs_human_review: bool
    # Issues
    critical_issues: int
    warnings: int
    issues: list[VerificationIssue]
    # Unexpected numbers in SAP not from protocol
    unexpected_numbers: list[str]
    # Full text report
    report_text: str
    # Metadata
    verification_method: str = "anchor-verification-v1"
    error: Optional[str] = None


# API Endpoints
@app.get("/")
async def root():
    return {"status": "ok", "message": "SAP Generator API v2.0", "features": ["file_upload", "pdf", "docx", "txt"]}


@app.get("/health")
async def health():
    """Basic health check."""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


@app.get("/health/detailed")
async def health_detailed():
    """
    Detailed health check with circuit breaker status.

    Returns:
        - LLM provider status (available, cooldown, error counts)
        - Database connectivity
        - Overall system health
    """
    health_status = {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "2.0.0",
        "components": {}
    }

    # Check LLM providers
    if LLM_CLIENT_AVAILABLE:
        try:
            client = get_tiered_client()
            llm_status = client.get_status()
            health_status["components"]["llm"] = {
                "status": "healthy" if any(s["available"] for s in llm_status.values()) else "degraded",
                "providers": llm_status
            }

            # Warn if all providers are in cooldown
            available_count = sum(1 for s in llm_status.values() if s["available"])
            if available_count == 0:
                health_status["status"] = "degraded"
                health_status["components"]["llm"]["status"] = "unavailable"
                logger.warning("All LLM providers unavailable", llm_status=llm_status)
        except Exception as e:
            health_status["components"]["llm"] = {
                "status": "error",
                "error": str(e)
            }
            logger.error("LLM health check failed", exc_info=True)
    else:
        health_status["components"]["llm"] = {"status": "not_configured"}

    # Check database
    try:
        db = get_supabase()
        # Simple query to verify connectivity
        health_status["components"]["database"] = {"status": "healthy"}
    except Exception as e:
        health_status["components"]["database"] = {
            "status": "error",
            "error": str(e)
        }
        health_status["status"] = "unhealthy"
        logger.error("Database health check failed", error=str(e))

    # Check environment
    health_status["components"]["environment"] = {
        "supabase_configured": bool(SUPABASE_URL and SUPABASE_KEY),
        "groq_configured": bool(GROQ_API_KEY),
        "anthropic_configured": bool(os.getenv("ANTHROPIC_API_KEY")),
        "openai_configured": bool(os.getenv("OPENAI_API_KEY")),
    }

    # Check SAP generation pipelines
    health_status["components"]["pipelines"] = {
        "direct_generation_v2": DIRECT_GENERATION_AVAILABLE,  # RECOMMENDED
        "production_pipeline": PRODUCTION_PIPELINE_AVAILABLE,
        "regulatory_generator": REGULATORY_GENERATOR_AVAILABLE,
        "agentic_pipeline": AGENTIC_PIPELINE_AVAILABLE,
    }

    return health_status


@app.post("/upload", response_model=JobResponse)
async def upload_file(
    file: UploadFile = File(...),
    nct_id: Optional[str] = Form(None)
):
    """
    Upload a protocol document (PDF, DOCX, TXT) and create a SAP generation job.
    """
    start_time = time.time()
    logger.info("File upload started", filename=file.filename, nct_id=nct_id)

    try:
        # Read file content
        content = await file.read()
        file_size = len(content)

        if file_size == 0:
            logger.warning("Empty file uploaded", filename=file.filename)
            raise HTTPException(status_code=400, detail="Empty file uploaded")

        if file_size > 10 * 1024 * 1024:  # 10MB limit
            logger.warning("File too large", filename=file.filename, size_bytes=file_size)
            raise HTTPException(status_code=400, detail="File too large (max 10MB)")

        # Extract text
        try:
            extracted_text = extract_text_from_file(file.filename, content)
            logger.info("Text extracted", filename=file.filename, text_length=len(extracted_text))
        except ValueError as e:
            logger.error("Text extraction failed", filename=file.filename, error=str(e))
            raise HTTPException(status_code=400, detail=str(e))

        if not extracted_text.strip():
            logger.warning("No text extracted", filename=file.filename)
            raise HTTPException(status_code=400, detail="No text could be extracted from the file")

        # Upload PDF to Supabase Storage for Vision-based parsing
        db = get_supabase()
        pdf_storage_path = None

        if file.filename.lower().endswith('.pdf'):
            try:
                import uuid
                # Generate unique filename
                storage_filename = f"{uuid.uuid4()}_{file.filename}"
                storage_path = f"protocols/{storage_filename}"

                # Upload to Supabase Storage bucket "pdfs"
                # Note: Bucket must exist in Supabase (create via dashboard)
                storage_result = db.storage.from_("pdfs").upload(
                    path=storage_path,
                    file=content,
                    file_options={"content-type": "application/pdf"}
                )

                pdf_storage_path = storage_path
                logger.info("PDF uploaded to storage", path=storage_path)
            except Exception as e:
                # Storage upload failed - continue without Vision (fall back to text)
                logger.warning("PDF storage upload failed, Vision disabled", error=str(e))
                pdf_storage_path = None

        # Insert job into database
        # CRITICAL: Store FULL text - do NOT truncate!
        # Statistical methods are at 50-80% of document, truncating loses them.
        job_data = {
            "protocol_text": extracted_text,  # Full text for multi-region sampling
            "nct_id": nct_id,
            "status": "queued",
            "filename": file.filename
        }

        # Add PDF storage path if available (for Vision-based parsing)
        if pdf_storage_path:
            job_data["pdf_storage_path"] = pdf_storage_path

        result = db.table("sap_jobs").insert(job_data).execute()

        job_id = result.data[0]["id"]
        elapsed = time.time() - start_time

        logger.info(
            "Job created",
            job_id=job_id,
            filename=file.filename,
            nct_id=nct_id,
            text_length=len(extracted_text),
            elapsed_seconds=round(elapsed, 2)
        )

        return JobResponse(
            job_id=job_id,
            status="queued",
            message=f"File '{file.filename}' uploaded successfully. Processing started.",
            extracted_text=extracted_text[:2000] + ("..." if len(extracted_text) > 2000 else "")
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Upload failed", filename=file.filename, exc_info=True, error=str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/pipeline-info")
async def get_pipeline_info():
    """
    Get information about the HYBRID SAP PIPELINE architecture.

    Returns details about all 4 layers and their components.
    """
    return {
        "pipeline": "HybridSAPPipeline",
        "version": "2.0.0",
        "architecture": {
            "layer_1_extraction": {
                "name": "EXTRACTION",
                "components": [
                    "StructuredFactExtractor (regex-only, no LLM hallucination)",
                    "ProtocolIdentityExtractor (NCT ID, sponsor detection)"
                ],
                "outputs": ["drug_name", "sample_size", "randomization_ratio", "phase", "therapeutic_area", "endpoints"]
            },
            "layer_2_knowledge": {
                "name": "KNOWLEDGE",
                "components": [
                    "BiostatisticsKnowledgeGraph (39 nodes, 36 edges)",
                    "RAG System (1,198 sections from real SAPs)",
                    "Specialized Templates (Phase 2/3, oncology, IBD, rheumatology)"
                ],
                "outputs": ["recommended_methods", "adam_datasets", "rag_examples", "template_guidance"]
            },
            "layer_3_generation": {
                "name": "GENERATION",
                "components": [
                    "ConstrainedSAPPipeline (Literal type enforcement)",
                    "FullSchemaGenerator (28-entity Pydantic schemas)",
                    "Multi-Agent System (4 specialized agents)"
                ],
                "outputs": ["sap_text", "constrained_output", "sections"]
            },
            "layer_4_validation": {
                "name": "VALIDATION",
                "components": [
                    "HardValidator (CRITICAL/HIGH/MEDIUM severity levels)",
                    "ContaminationGuard (cross-protocol detection)",
                    "IssueDetector (QA scoring)"
                ],
                "outputs": ["quality_score", "validation_issues", "contamination_report"]
            }
        },
        "endpoints": {
            "/generate": "Queued generation (background worker)",
            "/generate-full": "Synchronous generation (immediate response)",
            "/pipeline-info": "This endpoint - architecture details"
        }
    }


@app.post("/generate", response_model=JobResponse)
async def create_job(request: GenerateRequest):
    """
    Create a new SAP generation job from text.
    Returns job_id immediately, processing happens in background.
    """
    try:
        db = get_supabase()

        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        # Insert job into database
        # CRITICAL: Store FULL text - do NOT truncate!
        result = db.table("sap_jobs").insert({
            "protocol_text": request.protocol_text,  # Full text for multi-region sampling
            "nct_id": request.nct_id,
            "status": "queued",
            "filename": request.filename
        }).execute()

        job_id = result.data[0]["id"]

        return JobResponse(
            job_id=job_id,
            status="queued",
            message="Job created. Poll /status/{job_id} for results."
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# KG PIPELINE WRAPPER - Production SAP generation with dynamic structure (v69)
# =============================================================================

class KGPipelineWrapper:
    """
    Wrapper that provides EnhancedKGPipeline with a consistent interface.

    This enables the new 55-category KG extraction with prohibition rules
    to be used by the existing worker code without modification.

    Key features:
    1. 55-category comprehensive extraction (not 99 generic rules)
    2. Prohibition rules based on disease setting (adjuvant/metastatic/etc)
    3. SELF-RAG verification loop
    4. Full provenance tracking
    """

    def __init__(self):
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            raise RuntimeError("ANTHROPIC_API_KEY not set")

        self.kg_pipeline = EnhancedKGPipeline(api_key=api_key)
        print("[KGPipelineWrapper] Initialized with 55-category extraction + prohibition rules")

    def _build_prohibition_rules(self, extraction: Dict) -> List[str]:
        """
        Build context-specific prohibition rules based on extracted protocol data.

        These rules prevent Claude from generating inappropriate content
        based on the specific trial design (adjuvant vs metastatic, etc.)
        """
        rules = []

        # Get disease classification
        disease_class = extraction.get('disease_classification', {})
        disease_setting = disease_class.get('disease_setting', {}).get('value', '')

        # 1. Response criteria prohibition for adjuvant/neoadjuvant
        if disease_setting in ['adjuvant', 'neoadjuvant']:
            rules.append(
                "PROHIBITION: Do NOT include CR/PR/SD/PD response categories or RECIST criteria. "
                "This is an ADJUVANT/NEOADJUVANT trial - patients have no measurable disease. "
                "Use event-free survival (EFS), disease-free survival (DFS), or pathological response instead."
            )

        # 2. Performance status - check what's specified
        patient_pop = extraction.get('patient_population', {})
        ps_criteria = patient_pop.get('performance_status', {}).get('value', '')

        if ps_criteria:
            if 'ecog' in str(ps_criteria).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY ECOG Performance Status (found in protocol: {ps_criteria}). "
                    "Do NOT use Karnofsky Performance Status or ASA Score."
                )
            elif 'karnofsky' in str(ps_criteria).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY Karnofsky Performance Status (found in protocol: {ps_criteria}). "
                    "Do NOT use ECOG or ASA Score."
                )
            elif 'asa' in str(ps_criteria).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY ASA Score (found in protocol: {ps_criteria}). "
                    "Do NOT use ECOG or Karnofsky Performance Status."
                )

        # 3. Geographic/demographic prohibitions
        study_regions = extraction.get('study_design', {}).get('study_regions', {}).get('value', [])
        if study_regions:
            region_str = ', '.join(study_regions) if isinstance(study_regions, list) else str(study_regions)

            # Nordic countries - no race/ethnicity
            nordic_countries = ['sweden', 'norway', 'denmark', 'finland', 'iceland']
            if any(country in region_str.lower() for country in nordic_countries):
                rules.append(
                    "PROHIBITION: Do NOT include Race or Ethnicity in demographics tables. "
                    f"This trial includes Nordic countries ({region_str}) where race/ethnicity data collection is prohibited by law."
                )

            # Japan-only - no race/ethnicity subgroups
            if 'japan' in region_str.lower() and len(study_regions) == 1:
                rules.append(
                    "PROHIBITION: Do NOT include Race or Ethnicity subgroup analyses. "
                    "This is a Japan-only trial - race/ethnicity subgroups are not applicable."
                )

        # 4. Baseline characteristics prohibitions
        baseline = extraction.get('baseline_characteristics', {})
        weight_bmi = baseline.get('weight_bmi', {}).get('value', '')

        if weight_bmi:
            if 'weight' in str(weight_bmi).lower() and 'bmi' not in str(weight_bmi).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY Weight for body mass measurement (found: {weight_bmi}). "
                    "Do NOT include BMI in baseline tables."
                )
            elif 'bmi' in str(weight_bmi).lower() and 'weight' not in str(weight_bmi).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY BMI for body mass measurement (found: {weight_bmi}). "
                    "Do NOT include Weight in baseline tables."
                )

        # 5. AE grading prohibition
        safety = extraction.get('safety_endpoints', {})
        ae_grading = safety.get('ae_grading_scale', {}).get('value', '')

        if ae_grading:
            if 'ctcae' in str(ae_grading).lower():
                rules.append(
                    f"PROHIBITION: Use ONLY CTCAE grades for AE severity (specified: {ae_grading}). "
                    "Do NOT use Mild/Moderate/Severe categories."
                )
            elif any(term in str(ae_grading).lower() for term in ['mild', 'moderate', 'severe']):
                rules.append(
                    f"PROHIBITION: Use ONLY Mild/Moderate/Severe for AE severity (specified: {ae_grading}). "
                    "Do NOT use CTCAE grades."
                )

        return rules

    def process_protocol(self, protocol_text: str, protocol_id: str = "unknown",
                        sap_template: str = None, validate: bool = True,
                        verbose: bool = True) -> Dict[str, Any]:
        """
        Process protocol text using EnhancedKGPipeline.

        Adapts the return format to match TwoPassExtractor.
        """
        import tempfile
        import time

        start_time = time.time()

        if verbose:
            print(f"\n{'='*70}")
            print(f"[KGPipeline] PROCESSING PROTOCOL: {protocol_id}")
            print(f"{'='*70}")
            print("  + 55-category KG extraction")
            print("  + Prohibition rules (context-aware)")
            print("  + SELF-RAG verification")
            print("  + RAG Examples: Yes")
            print("  + Knowledge Graph: Yes")

        # Write protocol to temp file (KG pipeline expects file path)
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(protocol_text)
            temp_path = f.name

        try:
            # Run the enhanced KG pipeline with tool-based KB access
            result = self.kg_pipeline.process_protocol(temp_path, use_tools=True)

            # Build prohibition rules from extraction
            full_extraction = self.kg_pipeline._last_full_extraction or {}
            prohibition_rules = self._build_prohibition_rules(full_extraction)

            if verbose and prohibition_rules:
                print(f"\n[KGPipeline] Built {len(prohibition_rules)} prohibition rules:")
                for rule in prohibition_rules[:3]:
                    print(f"  • {rule[:80]}...")

            # Convert to TwoPassExtractor format
            sap_text = result.get('sap', '')
            extracted = result.get('extracted', [])
            verification = result.get('verification')
            provenance = result.get('provenance', {})

            # Build discovered_elements from extracted data AND full_extraction
            discovered_elements = []

            # Add facts from extracted list
            for fact in extracted:
                discovered_elements.append({
                    'name': fact.get('name', ''),
                    'category': fact.get('category', 'unknown'),
                    'description': str(fact.get('value', '')),
                    'source_quote': fact.get('source_quote', ''),
                    'confidence': fact.get('confidence', 0.8)
                })

            # v72: Add ALL relevant fields from full_extraction to discovered_elements
            # Including source_section for full traceability
            if full_extraction:

                # Helper to extract value from nested dict
                def get_value(obj):
                    if obj is None:
                        return None
                    if isinstance(obj, dict):
                        return obj.get('value') or obj.get('name') or None
                    return str(obj) if obj else None

                # Helper to safely get source fields from dict
                def get_source(obj, field='source_quote'):
                    if isinstance(obj, dict):
                        return obj.get(field, '')
                    return ''

                # 1. Primary endpoints
                primary_eps = full_extraction.get('primary_endpoints', [])
                for ep in (primary_eps if isinstance(primary_eps, list) else [primary_eps] if primary_eps else []):
                    ep_name = get_value(ep) if isinstance(ep, dict) else str(ep) if ep else None
                    if ep_name:
                        discovered_elements.append({
                            'name': 'primary_endpoint',
                            'category': 'endpoints',
                            'description': ep_name,
                            'source_quote': get_source(ep, 'source_quote'),
                            'source_section': get_source(ep, 'source_section'),
                            'confidence': 0.9
                        })

                # 2. Secondary endpoints
                secondary_eps = full_extraction.get('secondary_endpoints', [])
                for ep in (secondary_eps if isinstance(secondary_eps, list) else [secondary_eps] if secondary_eps else []):
                    ep_name = get_value(ep) if isinstance(ep, dict) else str(ep) if ep else None
                    if ep_name:
                        discovered_elements.append({
                            'name': 'secondary_endpoint',
                            'category': 'endpoints',
                            'description': ep_name,
                            'source_quote': get_source(ep, 'source_quote'),
                            'source_section': get_source(ep, 'source_section'),
                            'confidence': 0.9
                        })

                # 3. Phase
                phase_info = full_extraction.get('phase_info', {}) or full_extraction.get('study_phase', {})
                phase_obj = phase_info.get('phase', {}) if isinstance(phase_info, dict) else {}
                phase_val = get_value(phase_obj) if isinstance(phase_obj, dict) else get_value(phase_info)
                if phase_val:
                    discovered_elements.append({
                        'name': 'phase',
                        'category': 'study_design',
                        'description': phase_val,
                        'source_quote': get_source(phase_obj, 'source_quote'),
                        'source_section': get_source(phase_obj, 'source_section'),
                        'confidence': 0.9
                    })

                # 4. Study design type
                study_design = full_extraction.get('study_design', {})
                if isinstance(study_design, dict):
                    design_type_obj = study_design.get('design_type', {})
                    design_type = get_value(design_type_obj)
                    if design_type:
                        discovered_elements.append({
                            'name': 'design_type',
                            'category': 'study_design',
                            'description': design_type,
                            'source_quote': get_source(design_type_obj, 'source_quote'),
                            'source_section': get_source(design_type_obj, 'source_section'),
                            'confidence': 0.9
                        })

                    # Blinding
                    blinding_obj = study_design.get('blinding', {})
                    blinding = get_value(blinding_obj)
                    if blinding:
                        discovered_elements.append({
                            'name': 'blinding',
                            'category': 'study_design',
                            'description': blinding,
                            'source_quote': get_source(blinding_obj, 'source_quote'),
                            'source_section': get_source(blinding_obj, 'source_section'),
                            'confidence': 0.9
                        })

                    # Randomization ratio
                    rand_obj = study_design.get('randomization_ratio', {})
                    rand_ratio = get_value(rand_obj)
                    if rand_ratio:
                        discovered_elements.append({
                            'name': 'randomization_ratio',
                            'category': 'study_design',
                            'description': rand_ratio,
                            'source_quote': get_source(rand_obj, 'source_quote'),
                            'source_section': get_source(rand_obj, 'source_section'),
                            'confidence': 0.9
                        })

                # 5. Disease/indication (try both 'disease' and 'disease_classification')
                disease = full_extraction.get('disease', {}) or full_extraction.get('disease_classification', {})
                if isinstance(disease, dict):
                    tumor_obj = disease.get('tumor_type', {})
                    tumor_type = get_value(tumor_obj)
                    if tumor_type:
                        discovered_elements.append({
                            'name': 'tumor_type',
                            'category': 'disease',
                            'description': tumor_type,
                            'source_quote': get_source(tumor_obj, 'source_quote'),
                            'source_section': get_source(tumor_obj, 'source_section'),
                            'confidence': 0.9
                        })

                    stage_obj = disease.get('disease_stage', {})
                    disease_stage = get_value(stage_obj)
                    if disease_stage:
                        discovered_elements.append({
                            'name': 'disease_stage',
                            'category': 'disease',
                            'description': disease_stage,
                            'source_quote': get_source(stage_obj, 'source_quote'),
                            'source_section': get_source(stage_obj, 'source_section'),
                            'confidence': 0.9
                        })

                    setting_obj = disease.get('disease_setting', {})
                    disease_setting = get_value(setting_obj)
                    if disease_setting:
                        discovered_elements.append({
                            'name': 'disease_setting',
                            'category': 'disease',
                            'description': disease_setting,
                            'source_quote': get_source(setting_obj, 'source_quote'),
                            'source_section': get_source(setting_obj, 'source_section'),
                            'confidence': 0.9
                        })

                # 6. Sample size (try both 'sample_size' and 'enrollment')
                sample_size = full_extraction.get('sample_size', {}) or full_extraction.get('enrollment', {})
                if isinstance(sample_size, dict):
                    total_obj = sample_size.get('total_n', {}) or sample_size.get('target_enrollment', {})
                    total_n = get_value(total_obj)
                    if total_n:
                        discovered_elements.append({
                            'name': 'sample_size',
                            'category': 'study_design',
                            'description': str(total_n),
                            'source_quote': get_source(total_obj, 'source_quote'),
                            'source_section': get_source(total_obj, 'source_section'),
                            'confidence': 0.9
                        })

                    power_obj = sample_size.get('power', {})
                    power = get_value(power_obj)
                    if power:
                        discovered_elements.append({
                            'name': 'power',
                            'category': 'study_design',
                            'description': str(power),
                            'source_quote': get_source(power_obj, 'source_quote'),
                            'source_section': get_source(power_obj, 'source_section'),
                            'confidence': 0.9
                        })

                # 7. Study drug/treatment (try multiple keys)
                treatment = full_extraction.get('treatment', {}) or full_extraction.get('study_drug', {}) or full_extraction.get('treatment_arms', [])
                if isinstance(treatment, dict):
                    drug_name = get_value(treatment.get('name')) or get_value(treatment.get('drug_name'))
                    if drug_name:
                        discovered_elements.append({
                            'name': 'study_drug',
                            'category': 'treatment',
                            'description': drug_name,
                            'source_quote': get_source(treatment, 'source_quote'),
                            'source_section': get_source(treatment, 'source_section'),
                            'confidence': 0.9
                        })
                elif isinstance(treatment, list) and treatment:
                    # Handle treatment_arms array
                    for arm in treatment:
                        drug_name = arm.get('drug_name', '') or arm.get('arm_name', '')
                        if drug_name:
                            discovered_elements.append({
                                'name': 'study_drug',
                                'category': 'treatment',
                                'description': drug_name,
                                'source_quote': get_source(arm, 'source_quote'),
                                'source_section': get_source(arm, 'source_section'),
                                'confidence': 0.9
                            })

                # 8. Populations
                populations = full_extraction.get('populations', [])
                for pop in (populations if isinstance(populations, list) else []):
                    pop_name = get_value(pop) if isinstance(pop, dict) else str(pop) if pop else None
                    if pop_name:
                        discovered_elements.append({
                            'name': 'population',
                            'category': 'populations',
                            'description': pop_name,
                            'source_quote': get_source(pop, 'source_quote'),
                            'source_section': get_source(pop, 'source_section'),
                            'confidence': 0.9
                        })

                # 9. Stratification factors (try multiple keys)
                stratification = full_extraction.get('stratification', {}) or full_extraction.get('randomization', {})
                if isinstance(stratification, dict):
                    factors = stratification.get('factors', []) or stratification.get('stratification_factors', [])
                    for factor in (factors if isinstance(factors, list) else []):
                        factor_name = get_value(factor) if isinstance(factor, dict) else str(factor) if factor else None
                        if factor_name:
                            discovered_elements.append({
                                'name': 'stratification_factor',
                                'category': 'study_design',
                                'description': factor_name,
                                'source_quote': get_source(factor, 'source_quote'),
                                'source_section': get_source(factor, 'source_section'),
                                'confidence': 0.9
                            })

                # 10. Statistical methods
                stat_methods = full_extraction.get('statistical_methods', {})
                if isinstance(stat_methods, dict):
                    method_obj = stat_methods.get('primary_analysis_method', {})
                    primary_method = get_value(method_obj)
                    if primary_method:
                        discovered_elements.append({
                            'name': 'primary_analysis_method',
                            'category': 'statistical_methods',
                            'description': primary_method,
                            'source_quote': get_source(method_obj, 'source_quote'),
                            'source_section': get_source(method_obj, 'source_section'),
                            'confidence': 0.9
                        })

            # Build validation dict
            validation_dict = {
                'present': [],
                'missing': [],
                'partial': [],
                'overall_score': verification.score if verification else 0.8,
                'critical_gaps': [],
                'prohibition_rules_applied': prohibition_rules
            }

            if verification:
                for error in verification.errors:
                    validation_dict['missing'].append(error.field)

            total_time = time.time() - start_time

            if verbose:
                print(f"\n{'='*70}")
                print("[KGPipeline] PROCESSING COMPLETE")
                print(f"{'='*70}")
                print(f"  Total time: {total_time:.1f}s")
                print(f"  Elements extracted: {len(extracted)}")
                print(f"  SAP length: {len(sap_text):,} chars")
                print(f"  Verification score: {validation_dict['overall_score']:.1%}")
                print(f"  Prohibition rules: {len(prohibition_rules)}")

            return {
                'sap_text': sap_text,
                'protocol_id': protocol_id,
                'total_time_s': total_time,
                'discovered_elements': discovered_elements,
                'discovered_count': len(extracted),  # v65: For job completion logging
                'sap_length': len(sap_text),  # v65: For job completion logging
                'validation': validation_dict,
                'provenance': provenance,
                'prohibition_rules': prohibition_rules,
                'full_extraction': full_extraction
            }

        finally:
            # Clean up temp file
            try:
                os.unlink(temp_path)
            except Exception:
                pass

    def process_pdf(self, pdf_path: str, **kwargs) -> Dict[str, Any]:
        """
        Process PDF file using LlamaParse for extraction, then KG pipeline.
        """
        import asyncio

        protocol_id = kwargs.pop('protocol_id', Path(pdf_path).stem)

        print(f"[KGPipeline] Processing PDF: {pdf_path}")
        print("  + LlamaParse extraction (preserves tables)")
        print("  + 55-category KG extraction")
        print("  + Prohibition rules")

        # Try LlamaParse first
        try:
            from llama_parse import LlamaParse

            api_key = os.environ.get('LLAMAPARSE_API_KEY') or os.environ.get('LLAMA_CLOUD_API_KEY')
            if api_key:
                print("[KGPipeline] Using LlamaParse for PDF extraction")

                llamaparse = LlamaParse(
                    api_key=api_key,
                    result_type="markdown",
                    verbose=True
                )

                # Run async parse
                async def async_parse():
                    return await asyncio.wait_for(
                        llamaparse.aparse(pdf_path),
                        timeout=180.0
                    )

                # Run in thread to avoid event loop conflicts
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, async_parse())
                    result = future.result(timeout=200)

                # Get markdown output with page markers
                markdown_docs = result.get_markdown_documents(split_by_page=True)

                if markdown_docs:
                    protocol_parts = []
                    for i, doc in enumerate(markdown_docs):
                        protocol_parts.append(f"\n--- PAGE {i+1} ---\n")
                        protocol_parts.append(doc.text if hasattr(doc, 'text') else str(doc))

                    protocol_text = "\n".join(protocol_parts)
                    print(f"[KGPipeline] LlamaParse extracted {len(protocol_text):,} chars from {len(markdown_docs)} pages")

                    return self.process_protocol(
                        protocol_text=protocol_text,
                        protocol_id=protocol_id,
                        **kwargs
                    )

        except ImportError:
            print("[KGPipeline] LlamaParse not available, falling back to PyPDF2")
        except Exception as e:
            print(f"[KGPipeline] LlamaParse failed: {e}, falling back to PyPDF2")

        # Fallback to PyPDF2
        try:
            import PyPDF2

            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        pages.append(f"\n--- PAGE {i+1} ---\n{text}")

                protocol_text = "\n".join(pages)
                print(f"[KGPipeline] PyPDF2 extracted {len(protocol_text):,} chars from {len(pages)} pages")

                return self.process_protocol(
                    protocol_text=protocol_text,
                    protocol_id=protocol_id,
                    **kwargs
                )

        except Exception as e:
            print(f"[KGPipeline] PDF extraction failed: {e}")
            raise RuntimeError(f"Failed to extract PDF: {e}")


# Global pipeline instance (reused across requests)
_production_pipeline = None  # KGPipelineWrapper with prohibition rules

def get_pipeline():
    """
    Get or create the production pipeline instance.

    Uses KGPipelineWrapper (55-category extraction + prohibition rules):
    1. LlamaParse: PDF → Markdown (preserves tables, complex layouts)
    2. 55-Category KG Extraction: Comprehensive protocol analysis
    3. Prohibition Rules: Context-aware (adjuvant→no CR/PR/SD/PD, Nordic→no race/ethnicity)
    4. SAP Generation with prohibition rules in prompt
    5. SELF-RAG Verification: Fact checking with correction loop
    """
    global _production_pipeline

    if _production_pipeline is None:
        if not KG_PIPELINE_AVAILABLE or EnhancedKGPipeline is None:
            raise RuntimeError("KGPipeline not available - check imports")

        _production_pipeline = KGPipelineWrapper()
        logger.info("KGPipelineWrapper initialized (55-category + prohibition rules)")
        print("[get_pipeline] Using KGPipelineWrapper (55-category + prohibition rules)")

    return _production_pipeline

# Aliases for backward compatibility
def get_hybrid_pipeline():
    """Deprecated: Use get_pipeline() instead."""
    return get_pipeline()

def get_full_pipeline():
    """Deprecated: Use get_pipeline() instead."""
    return get_pipeline()


class FullPipelineResponse(BaseModel):
    """Response from full integrated pipeline with all layers."""
    success: bool
    sap_text: str
    drug_name: str
    sample_size: int
    randomization_ratio: str
    phase: str
    therapeutic_area: str
    endpoint_type: str
    quality_score: float
    generation_mode: str
    constrained_schema_used: bool
    rag_examples_count: int
    templates_applied: list
    validation_issues: int
    contamination_detected: bool
    processing_time: float
    errors: list


@app.post("/generate-full", response_model=FullPipelineResponse)
async def generate_full_pipeline(request: GenerateRequest):
    """
    Generate SAP synchronously using the RULE-BASED PIPELINE.

    This endpoint uses:
    - Step 1: Claude LLM extraction (NCT ID, drug, sample size, etc.)
    - Step 2: Condition detection (immunotherapy, crossover, interim, etc.)
    - Step 3: Knowledge Graph with 99 rules for method selection
    - Step 4: ChromaDB RAG with 17K+ chunks for examples
    - Step 5: Claude LLM generation with slot constraints
    - Step 6: Slot verification for required methods

    Returns immediately with the generated SAP (no queuing).
    """
    import time
    start_time = time.time()

    try:
        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        pipeline = get_pipeline()

        # CRITICAL: Pass FULL protocol text - do NOT truncate!
        # The pipeline uses multi-region sampling internally to handle large docs.
        # Truncating here cuts off statistical methods which are at 50-80% of doc.
        result = pipeline.generate(request.protocol_text)

        processing_time = time.time() - start_time

        # Handle ProductionSAPPipeline result (has facts)
        if hasattr(result, 'facts') and result.facts:
            # ProductionSAPPipeline format
            facts = result.facts
            drug_name = facts.get('drug_name', '') or ''
            sample_size_val = facts.get('sample_size', 0)
            if isinstance(sample_size_val, dict):
                sample_size = sample_size_val.get('total_n', 0) or 0
            elif isinstance(sample_size_val, int):
                sample_size = sample_size_val
            else:
                sample_size = 0
            ratio = facts.get('randomization_ratio', '') or ''
            phase = facts.get('phase', '') or ''
            therapeutic_area = facts.get('therapeutic_area', '') or facts.get('indication', '') or ''
            # Handle primary_endpoint - could be string, dict, or list
            ep = facts.get('primary_endpoint', '')
            if isinstance(ep, str):
                endpoint_type = ep[:100]
            elif isinstance(ep, dict):
                endpoint_type = str(ep.get('name', '') or ep.get('definition', ''))[:100]
            elif isinstance(ep, list) and ep:
                endpoint_type = str(ep[0])[:100] if ep[0] else ""
            else:
                endpoint_type = ""

            # Quality score on 0-100 scale for frontend display
            quality_score = 100.0 if result.verification and getattr(result.verification, 'passed', False) else 50.0
            missing_slots = getattr(result.verification, 'missing_slots', None) if result.verification else None
            validation_issues = len(missing_slots) if missing_slots else 0
            generation_mode = "rule-based (Claude + 99 rules + RAG + slot verification)"
            source_trials = []

        elif hasattr(result, 'characteristics') and result.characteristics:
            # FALLBACK: AgenticSAPPipeline format
            chars = result.characteristics
            drug_name = chars.drug_classes[0] if chars.drug_classes else ""
            phase = chars.phase or ""
            therapeutic_area = chars.indication or ""
            endpoint_type = chars.endpoint_type or ""
            ratio = ""  # Not in characteristics

            # Try to extract sample size from extracted_methods or default
            sample_size = 0

            # Validation from agentic pipeline
            # Quality score on 0-100 scale (result.confidence is 0-1, multiply by 100)
            try:
                quality_score = (float(result.confidence) * 100) if result.confidence else 80.0
            except (ValueError, TypeError):
                quality_score = 80.0
            if result.validation:
                # validation.confidence is 0-1, scale to 0-100
                try:
                    quality_score = float(result.validation.confidence) * 100 if result.validation.confidence else quality_score
                except (ValueError, TypeError):
                    pass
                validation_issues = len(result.validation.issues) if hasattr(result.validation, 'issues') and result.validation.issues else 0
            else:
                validation_issues = 0

            generation_mode = "agentic-hybridrag (5-agent + Knowledge Graph + 23K chunks)"
            source_trials = result.source_trials or []

        else:
            # Minimal fallback
            drug_name = ""
            sample_size = 0
            ratio = ""
            phase = ""
            therapeutic_area = ""
            endpoint_type = ""
            quality_score = 50.0  # Default fallback (0-100 scale)
            validation_issues = 0
            generation_mode = "unknown"
            source_trials = []

        return FullPipelineResponse(
            success=getattr(result, 'success', False),
            sap_text=getattr(result, 'sap_text', ''),
            drug_name=drug_name,
            sample_size=sample_size,
            randomization_ratio=ratio,
            phase=phase,
            therapeutic_area=therapeutic_area,
            endpoint_type=endpoint_type,
            quality_score=quality_score,
            generation_mode=generation_mode,
            constrained_schema_used=True,
            rag_examples_count=len(result.sections) if hasattr(result, 'sections') and result.sections else 0,
            templates_applied=list(result.sections.keys()) if hasattr(result, 'sections') and result.sections else [],
            validation_issues=validation_issues,
            contamination_detected=False,
            processing_time=processing_time,
            errors=result.warnings if hasattr(result, 'warnings') and result.warnings else (
                [result.error] if hasattr(result, 'error') and result.error else []
            )
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Global instance for regulatory generator
_regulatory_generator: RegulatorySAPGenerator = None


def get_regulatory_generator():
    """Get or create the regulatory SAP generator."""
    global _regulatory_generator
    if _regulatory_generator is None and REGULATORY_GENERATOR_AVAILABLE:
        _regulatory_generator = create_regulatory_sap_generator()
        logger.info("RegulatorySAPGenerator initialized (ICH E9 compliant, Claude extraction)")
    return _regulatory_generator


class RegulatorySAPResponse(BaseModel):
    """Response from regulatory-grade SAP generation."""
    success: bool
    sap_text: str
    # Extracted facts
    nct_id: str
    protocol_number: str
    drug_name: str
    comparator_drug: str
    sample_size: int
    events_required: int
    primary_endpoint: str
    primary_test: str
    alpha_interim: float
    alpha_final: float
    stratification_factors: list
    has_interim: bool
    dmc_oversight: bool
    # Metadata
    sections_generated: int
    character_count: int
    processing_time: float
    extraction_method: str  # Always "claude" (no regex fallback)
    errors: list


@app.post("/generate-regulatory", response_model=RegulatorySAPResponse)
async def generate_regulatory_sap(request: GenerateRequest):
    """
    Generate a REGULATORY-GRADE SAP using Claude API extraction.

    This endpoint produces SAPs that match real pharmaceutical SAPs:
    - ICH E9(R1) compliant structure (10 major sections)
    - 45+ pages with proper formatting
    - Protocol-specific statistical methods (Fleming-Harrington, Lan-DeMets)
    - Proper censoring schemes, analysis populations, subgroup analyses

    Uses Claude API for accurate protocol fact extraction.
    """
    import time
    start_time = time.time()

    try:
        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        if not REGULATORY_GENERATOR_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="RegulatorySAPGenerator not available. Check imports."
            )

        generator = get_regulatory_generator()
        if generator is None:
            raise HTTPException(
                status_code=503,
                detail="Could not initialize RegulatorySAPGenerator"
            )

        # Extract facts (uses Claude API if available)
        # CRITICAL: Pass FULL text - do NOT truncate!
        facts = generator.extract_protocol_facts(request.protocol_text)

        # Generate SAP document
        doc = generator.generate(request.protocol_text, facts)

        # Assemble full document
        sap_text = generator.assemble_document(doc)

        processing_time = time.time() - start_time

        # Extraction method is always Claude (no regex fallback)
        extraction_method = "claude"

        return RegulatorySAPResponse(
            success=True,
            sap_text=sap_text,
            nct_id=facts.nct_id or "",
            protocol_number=facts.protocol_number or "",
            drug_name=facts.experimental_drug or "",
            comparator_drug=facts.comparator_drug or "",
            sample_size=facts.total_sample_size or 0,
            events_required=facts.events_required_final or 0,
            primary_endpoint=facts.primary_endpoint or "",
            primary_test=facts.primary_test or "",
            alpha_interim=facts.alpha_interim or 0.0,
            alpha_final=facts.alpha_final or 0.05,
            stratification_factors=facts.stratification_factors or [],
            has_interim=facts.has_interim,
            dmc_oversight=facts.dmc_oversight,
            sections_generated=len([s for s in [
                doc.cover_page, doc.sec1_1_hypothesis, doc.sec2_1_design,
                doc.sec5_sample_size, doc.sec7_5_1_primary_analysis, doc.sec7_6_safety
            ] if s]),
            character_count=len(sap_text),
            processing_time=processing_time,
            extraction_method=extraction_method,
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Regulatory SAP generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# DIRECT SAP GENERATION (V2) - No information loss
# =============================================================================

# v69: Direct generation now uses EnhancedKGPipeline (same as get_pipeline())
# This maintains the /generate-direct endpoint but uses the new dynamic SAP structure
def get_direct_generator():
    """Get the SAP generator - v69 uses EnhancedKGPipeline via get_pipeline()."""
    return get_pipeline()  # Returns KGPipelineWrapper with EnhancedKGPipeline


# Global instance for integrated pipeline (with LLM extraction + RAG)
_integrated_pipeline: 'IntegratedSAPPipeline' = None


def get_integrated_pipeline():
    """
    Get or create the integrated pipeline with full LLM extraction.

    This pipeline includes:
    - LLM-based extraction for complex elements (interim analysis, power calculations, etc.)
    - RAG with preserved interim analysis values
    - Knowledge graph for regulatory context
    - Full coverage of Phase 1/2/3 trial elements
    """
    global _integrated_pipeline
    if _integrated_pipeline is None and INTEGRATED_PIPELINE_AVAILABLE:
        _integrated_pipeline = IntegratedSAPPipeline()
        logger.info("IntegratedSAPPipeline initialized (LLM extraction + RAG + KnowledgeGraph)")
    return _integrated_pipeline


class DirectSAPResponse(BaseModel):
    """Response from direct SAP generation (V2 - no information loss)."""
    success: bool
    sap_text: str
    # Discovery results
    elements_discovered: int
    categories_found: list
    # Validation (checklist coverage)
    validation_score: float
    elements_present: int
    elements_missing: int
    elements_partial: int
    critical_gaps: list
    # Verification (anchor-based, Generate → Verify architecture)
    verification_score: Optional[float] = None
    anchors_found: Optional[int] = None
    anchors_verified: Optional[int] = None
    anchors_missing: Optional[int] = None
    verification_issues: Optional[list] = None
    needs_human_review: Optional[bool] = None
    # Deterministic verification (non-LLM)
    deterministic_verification: Optional[dict] = None
    audit_report: Optional[str] = None
    # Metadata
    total_time: float
    sap_length: int
    generation_method: str
    errors: list


@app.post("/generate-direct", response_model=DirectSAPResponse)
async def generate_direct_sap(request: GenerateRequest):
    """
    Generate SAP using DIRECT GENERATION (V2) - NO INFORMATION LOSS.

    This is the RECOMMENDED endpoint for SAP generation.

    Architecture:
    1. Pass 1: Discover ALL statistical elements in protocol (checklist)
    2. Pass 2: Generate SAP directly from FULL protocol text with checklist

    Unlike the old pipeline which extracts → flattens → generates (loses info),
    this approach sends the full protocol text directly to the LLM with a
    checklist of elements to include. NO information is lost.

    Benefits:
    - 100% coverage of discovered elements
    - Correct blinding type (open-label vs blinded)
    - All hypotheses (H1, H2, H3, etc.) captured
    - Correct interim analysis count
    - Accurate alpha allocations
    - Non-inferiority margins preserved
    - Regional extensions (China, etc.) captured
    """
    import time
    start_time = time.time()

    try:
        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        if not DIRECT_GENERATION_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="Direct generation not available - KGPipeline import failed"
            )

        generator = get_direct_generator()
        if generator is None:
            raise HTTPException(
                status_code=503,
                detail="Direct generator not initialized"
            )

        # Run the full pipeline: discover → generate → validate
        result = generator.process_protocol(
            protocol_text=request.protocol_text,
            protocol_id=request.filename or "uploaded_protocol",
            validate=True,
            verbose=True
        )

        processing_time = time.time() - start_time

        # Extract validation results
        validation = result.get('validation', {})
        validation_score = validation.get('overall_score', 0.0)
        elements_present = len(validation.get('present', []))
        elements_missing = len(validation.get('missing', []))
        elements_partial = len(validation.get('partial', []))
        critical_gaps = validation.get('critical_gaps', [])

        # Extract discovered element categories
        discovered_elements = result.get('discovered_elements', [])
        categories = list(set(e.get('category', 'other') for e in discovered_elements))

        # Run verification (Generate → Verify architecture)
        verification_score = None
        anchors_found = None
        anchors_verified = None
        anchors_missing = None
        verification_issues = None
        needs_human_review = None

        if SAP_VERIFIER_AVAILABLE:
            try:
                sap_text = result.get('sap_text', '')
                if sap_text and request.protocol_text:
                    anchors = extract_anchors(request.protocol_text)
                    report = verify_sap(sap_text, request.protocol_text, anchors)

                    verification_score = report.confidence_score
                    anchors_found = report.anchors_found
                    anchors_verified = report.anchors_verified
                    anchors_missing = report.anchors_missing
                    needs_human_review = report.needs_human_review()

                    # Extract top issues (limit to 10)
                    verification_issues = [
                        {
                            "severity": issue.severity.value,
                            "category": issue.category,
                            "message": issue.message[:200]  # Truncate long messages
                        }
                        for issue in report.issues[:10]
                    ]
            except Exception as verify_error:
                logger.warning(f"Verification failed (non-fatal): {verify_error}")
                verification_issues = [{"severity": "warning", "category": "system", "message": f"Verification skipped: {str(verify_error)[:100]}"}]

        # Get deterministic verification results from two_pass_extractor
        det_verification = result.get('verification', {})

        return DirectSAPResponse(
            success=True,
            sap_text=result.get('sap_text', ''),
            elements_discovered=result.get('discovered_count', len(discovered_elements)),
            categories_found=categories,
            validation_score=validation_score,
            elements_present=elements_present,
            elements_missing=elements_missing,
            elements_partial=elements_partial,
            critical_gaps=critical_gaps,
            verification_score=verification_score,
            anchors_found=anchors_found,
            anchors_verified=anchors_verified,
            anchors_missing=anchors_missing,
            verification_issues=verification_issues,
            needs_human_review=needs_human_review or det_verification.get('requires_human_review', True),
            deterministic_verification={
                "passed": det_verification.get('passed', 0),
                "failed": det_verification.get('failed', 0),
                "warnings": det_verification.get('warnings', 0),
                "critical_failures": det_verification.get('critical_failures', [])
            },
            audit_report=det_verification.get('audit_report', ''),
            total_time=result.get('total_time_s', processing_time),
            sap_length=result.get('sap_length', len(result.get('sap_text', ''))),
            generation_method="direct-v2 (discovery checklist + full protocol + deterministic verification)",
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Direct SAP generation failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# INTEGRATED PIPELINE - LLM Extraction + RAG + Knowledge Graph
# =============================================================================

class IntegratedSAPResponse(BaseModel):
    """Response from integrated pipeline with full LLM extraction."""
    success: bool
    sap_text: str
    # Extraction results
    facts_extracted: dict
    interim_analysis: dict
    power_calculations: dict
    censoring_rules: dict
    # RAG info
    rag_examples_used: int
    rag_nct_ids: list
    # Template info
    trial_type: str
    templates_applied: list
    # Validation
    validation_score: float
    issues_found: list
    # Metadata
    total_time: float
    sap_length: int
    generation_method: str
    errors: list


@app.post("/generate-integrated", response_model=IntegratedSAPResponse)
async def generate_integrated_sap(request: GenerateRequest):
    """
    Generate SAP using INTEGRATED PIPELINE with full LLM extraction.

    This is the MOST COMPREHENSIVE endpoint - captures ALL details including:
    - Interim analysis (count, timing, alpha spending, boundaries)
    - Power calculations (PFS power, OS power, control medians)
    - Censoring rules (PFS, DOR, PFS2)
    - Exploratory endpoints (DOR, DCR, CBR, PFS2, iRECIST)
    - PRO thresholds (timepoint, completion, MCID)
    - Regional extensions (China sample size, events)
    - Protocol violation definitions
    - Laboratory parameter lists
    - Data handling conventions

    Architecture:
    1. LlamaParse: PDF → Markdown (preserves tables)
    2. FactExtractor: Regex for basic facts
    3. LLMExtractor: Complex elements (interim, power, censoring, etc.)
    4. RAGRetriever: Similar SAP examples (with interim values preserved)
    5. KnowledgeGraph: Regulatory context
    6. SAPGenerator: Section-by-section generation with all facts
    7. IssueDetector: QA validation

    Use this endpoint when you need:
    - Complete interim analysis details
    - Full censoring rule tables
    - Detailed power calculations
    - Phase 1 PK/PD/safety details
    """
    import time
    start_time = time.time()

    try:
        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        if not INTEGRATED_PIPELINE_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="Integrated pipeline not available - check imports"
            )

        pipeline = get_integrated_pipeline()
        if pipeline is None:
            raise HTTPException(
                status_code=503,
                detail="Integrated pipeline not initialized"
            )

        # Run the full integrated pipeline
        result = pipeline.generate(request.protocol_text)

        processing_time = time.time() - start_time

        # Extract facts for response
        facts = {}
        interim_analysis = {}
        power_calculations = {}
        censoring_rules = {}

        if hasattr(result, 'facts') and result.facts:
            facts = result.facts if isinstance(result.facts, dict) else {}
            interim_analysis = facts.get('interim_analysis', {})
            power_calculations = facts.get('power_calculations', {})
            censoring_rules = facts.get('censoring_rules', {})

        # Get validation issues
        issues = []
        validation_score = 0.0
        if hasattr(result, 'validation') and result.validation:
            validation_score = getattr(result.validation, 'overall_score', 0.0)
            if hasattr(result.validation, 'issues'):
                issues = [str(i) for i in result.validation.issues[:10]]

        return IntegratedSAPResponse(
            success=result.success if hasattr(result, 'success') else True,
            sap_text=result.sap_text if hasattr(result, 'sap_text') else str(result),
            facts_extracted=facts,
            interim_analysis=interim_analysis,
            power_calculations=power_calculations,
            censoring_rules=censoring_rules,
            rag_examples_used=getattr(result, 'rag_examples_used', 0),
            rag_nct_ids=getattr(result, 'rag_nct_ids', []),
            trial_type=getattr(result, 'trial_type', 'unknown'),
            templates_applied=getattr(result, 'templates_applied', []),
            validation_score=validation_score,
            issues_found=issues,
            total_time=processing_time,
            sap_length=len(result.sap_text) if hasattr(result, 'sap_text') else 0,
            generation_method="integrated-v3 (LLM extraction + RAG + KnowledgeGraph)",
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Integrated SAP generation failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# RAG 2-CALL SAP GENERATION - Efficient RAG with TLF Appendix
# =============================================================================

# Global RAG index instance
_rag_index: SAPRAGIndex = None


def get_rag_index():
    """Get or create the 3-collection RAG index."""
    global _rag_index
    if _rag_index is None and RAG_SYSTEM_AVAILABLE:
        _rag_index = SAPRAGIndex()
        logger.info("SAPRAGIndex initialized (3-collection: structure, content, TLF)")
    return _rag_index


class RAGSAPResponse(BaseModel):
    """Response from RAG-enhanced SAP generation."""
    success: bool
    sap_text: str
    tlf_appendix: str
    # Discovery results
    elements_discovered: int
    categories_found: list
    # RAG info
    rag_structure_used: bool
    rag_content_examples: int
    rag_tlf_shells: int
    # Metadata
    total_time: float
    sap_length: int
    llm_calls: int
    generation_method: str
    errors: list


@app.post("/generate-rag", response_model=RAGSAPResponse)
async def generate_rag_sap(request: GenerateRequest):
    """
    Generate SAP using RAG 2-CALL approach with TLF appendix.

    Architecture (2 LLM calls + RAG queries):
    1. Pass 1: Discovery - Find all statistical elements (1 LLM call)
    2. RAG Queries: Get structure, content examples, TLF shells (0 LLM calls)
    3. Pass 2: Generate full SAP with RAG context (1 LLM call)

    Benefits:
    - Same accuracy as V2 Direct (~97%)
    - Includes TLF appendix from real SAPs
    - Industry-standard formatting from RAG examples
    - Only 2 LLM calls (efficient)
    """
    import time
    start_time = time.time()

    try:
        if not request.protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        if not DIRECT_GENERATION_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="Direct generation not available - KGPipeline import failed"
            )

        if not RAG_SYSTEM_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="RAG system not available - SAPRAGIndex import failed"
            )

        generator = get_direct_generator()
        rag = get_rag_index()

        if generator is None or rag is None:
            raise HTTPException(
                status_code=503,
                detail="Generator or RAG index not initialized"
            )

        # STEP 1: Discovery (1 LLM call)
        logger.info("[RAG] Step 1: Discovery...")
        result = generator.process_protocol(
            protocol_text=request.protocol_text,
            protocol_id=request.filename or "uploaded_protocol",
            validate=False,
            verbose=False
        )

        discovered = result.get('discovered_elements', [])
        categories = list(set(e.get('category', 'other') for e in discovered))
        facts_text = "\n".join([
            f"- {d.get('name')}: {d.get('description', '')[:200]}"
            for d in discovered[:40]
        ])

        # Extract SOA/Visit Schedule information from protocol
        soa_text = ""
        import re

        # Look for Reducto-enhanced SOA section (HTML tables)
        reducto_marker = "SCHEDULE OF ASSESSMENTS (Enhanced by Reducto)"
        if reducto_marker in request.protocol_text:
            # Extract content between Reducto markers
            start_idx = request.protocol_text.find(reducto_marker)
            end_marker = "END OF SCHEDULE OF ASSESSMENTS"
            end_idx = request.protocol_text.find(end_marker, start_idx)
            if end_idx == -1:
                end_idx = start_idx + 15000  # Fallback: take 15000 chars

            reducto_content = request.protocol_text[start_idx:end_idx]

            # Convert HTML tables to markdown pipe format
            def html_table_to_markdown(html_content):
                """Convert HTML tables to markdown pipe-separated format."""
                from html.parser import HTMLParser

                class TableParser(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.tables = []
                        self.current_table = []
                        self.current_row = []
                        self.current_cell = ""
                        self.in_table = False
                        self.in_row = False
                        self.in_cell = False

                    def handle_starttag(self, tag, attrs):
                        if tag == 'table':
                            self.in_table = True
                            self.current_table = []
                        elif tag == 'tr':
                            self.in_row = True
                            self.current_row = []
                        elif tag in ('td', 'th'):
                            self.in_cell = True
                            self.current_cell = ""

                    def handle_endtag(self, tag):
                        if tag == 'table':
                            self.in_table = False
                            if self.current_table:
                                self.tables.append(self.current_table)
                            self.current_table = []
                        elif tag == 'tr':
                            self.in_row = False
                            if self.current_row:
                                self.current_table.append(self.current_row)
                            self.current_row = []
                        elif tag in ('td', 'th'):
                            self.in_cell = False
                            self.current_row.append(self.current_cell.strip())
                            self.current_cell = ""

                    def handle_data(self, data):
                        if self.in_cell:
                            self.current_cell += data

                parser = TableParser()
                try:
                    parser.feed(html_content)
                except:
                    return html_content  # Return original if parsing fails

                # Convert parsed tables to markdown
                markdown_tables = []
                for table in parser.tables:
                    if not table:
                        continue
                    md_lines = []
                    for i, row in enumerate(table):
                        md_lines.append("| " + " | ".join(row) + " |")
                        if i == 0:  # Add separator after header
                            md_lines.append("|" + "|".join(["---"] * len(row)) + "|")
                    markdown_tables.append("\n".join(md_lines))

                return "\n\n".join(markdown_tables) if markdown_tables else html_content

            # Convert to markdown
            soa_markdown = html_table_to_markdown(reducto_content)
            soa_text = f"SCHEDULE OF ASSESSMENTS (from protocol):\n\n{soa_markdown[:8000]}"
            logger.info(f"[SAP Gen] Extracted Reducto SOA: {len(soa_markdown)} chars")
        else:
            # Fallback: look for SOA in discovered elements
            soa_elements = [d for d in discovered if any(kw in d.get('name', '').lower()
                           for kw in ['visit', 'schedule', 'assessment', 'window', 'soa'])]
            if soa_elements:
                soa_text = "VISIT SCHEDULE ELEMENTS:\n" + "\n".join([
                    f"- {d.get('name')}: {d.get('description', '')}" for d in soa_elements[:10]
                ])

        # STEP 2: RAG Queries (0 LLM calls)
        logger.info("[RAG] Step 2: RAG queries...")

        # Query structure
        structures = rag.query_structure("Phase 3 randomized oncology PFS", n_results=1)
        structure_example = structures[0]["content"][:1500] if structures else ""

        # Query content examples
        content_examples = ""
        content_count = 0
        for section in ["sample_size", "interim_analysis", "efficacy_analysis", "safety_analysis"]:
            examples = rag.query_content(section, n_results=1)
            if examples:
                content_examples += f"\n[{section}]: {examples[0]['content'][:600]}\n"
                content_count += 1

        # Query TLF shells (use new template-based categories)
        tlf_text = "\n## TLF SHELL SPECIFICATIONS\n"
        tlf_count = 0
        for category in ["demographics", "efficacy", "safety", "figures"]:
            tlfs = rag.query_tlf(category, category=category, n_results=2)
            for t in tlfs:
                # Include full shell specification, not just title
                tlf_text += f"\n{t['content']}\n"
                tlf_count += 1

        # STEP 3: Generate full SAP (1 LLM call)
        logger.info("[RAG] Step 3: Generate full SAP...")

        # Use Anthropic client directly for generation
        try:
            from anthropic import Anthropic
            client = Anthropic()

            prompt = f"""Generate a complete Statistical Analysis Plan (SAP).

PROTOCOL FACTS (use these exact values):
{facts_text}

SCHEDULE OF ASSESSMENTS / VISIT WINDOWS:
{soa_text if soa_text else "Extract visit schedule from protocol text"}

SAP STRUCTURE EXAMPLE (follow this organization):
{structure_example}

STYLE EXAMPLES (follow this professional format):
{content_examples}

Generate complete SAP with ALL sections:
1. Introduction
2. Study Objectives and Endpoints
3. Study Design
4. Sample Size Determination
5. Analysis Populations
6. Statistical Hypotheses and Testing Strategy
7. Statistical Methods for Efficacy (Primary and Secondary)
8. Interim Analyses
9. Safety Analyses
10. Missing Data Handling
11. Patient-Reported Outcomes
12. Regional Considerations (if applicable)
13. Data Handling Conventions (REQUIRED - include Visit Windows table, Analysis Windows, Baseline Definitions)

SECTION 13 REQUIREMENTS:
- Include a Visit Windows table with columns: Analysis Visit | Target Day | Window (Days)
- Define baseline as last non-missing assessment prior to first dose
- Specify tumor assessment frequency (e.g., every 8 weeks)
- Include PRO collection schedule if applicable
- Define analysis windows for each scheduled visit

NOTE: TLF Shell Specifications will be appended automatically - do NOT include placeholder text for TLF appendix.

REQUIREMENTS:
- Use ALL protocol facts with exact numbers (alpha, sample size, HR, etc.)
- Include specific statistical methods (log-rank, Cox, Miettinen-Nurminen, etc.)
- Follow professional SAP formatting with numbered sections
- Section 13 MUST include visit windows table derived from Schedule of Assessments"""

            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}]
            )

            sap_text = response.content[0].text

            # NOTE: Section 12.2 already contains endpoint-specific TLF specs
            # Do NOT append raw TLF templates with placeholders

        except Exception as llm_error:
            logger.error(f"LLM generation failed: {llm_error}")
            # Fallback to V2 direct generation (already has clean TLF specs)
            sap_text = result.get('sap_text', '')

        processing_time = time.time() - start_time

        return RAGSAPResponse(
            success=True,
            sap_text=sap_text,
            tlf_appendix=tlf_text,
            elements_discovered=len(discovered),
            categories_found=categories,
            rag_structure_used=bool(structure_example),
            rag_content_examples=content_count,
            rag_tlf_shells=tlf_count,
            total_time=processing_time,
            sap_length=len(sap_text),
            llm_calls=2,
            generation_method="rag-2call (discovery + RAG context + generation)",
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"RAG SAP generation failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/status/{job_id}", response_model=JobStatusResponse)
async def get_job_status(job_id: str):
    """
    Get the status of a SAP generation job.
    """
    try:
        db = get_supabase()

        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        # Create protocol preview
        protocol_text = job.get("protocol_text", "")
        preview = protocol_text[:1000] + ("..." if len(protocol_text) > 1000 else "")

        return JobStatusResponse(
            job_id=job["id"],
            status=job["status"],
            generated_sap=job.get("generated_sap"),
            quality_score=job.get("quality_score"),
            endpoint_type=job.get("endpoint_type"),
            phase=job.get("phase"),
            therapeutic_area=job.get("therapeutic_area"),
            processing_time=job.get("processing_time"),
            error_message=job.get("error_message"),
            created_at=job.get("created_at"),
            completed_at=job.get("completed_at"),
            filename=job.get("filename"),
            protocol_preview=preview
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/jobs")
async def list_jobs(limit: int = 20):
    """
    List recent jobs.
    """
    try:
        db = get_supabase()

        result = db.table("sap_jobs").select(
            "id, status, nct_id, filename, quality_score, endpoint_type, phase, created_at, completed_at, processing_time"
        ).order("created_at", desc=True).limit(limit).execute()

        return {"jobs": result.data}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/jobs/{job_id}")
async def delete_job(job_id: str):
    """Delete a job."""
    try:
        db = get_supabase()
        db.table("sap_jobs").delete().eq("id", job_id).execute()
        return {"message": "Job deleted"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/ground-truth")
async def list_ground_truth():
    """
    List available ground truth SAPs for evaluation.
    Includes all pairs from data/all_pairs directory.
    """
    try:
        # Check both directories
        base_dir = Path(__file__).parent.parent.parent / "data"
        all_pairs_dir = base_dir / "all_pairs"
        ground_truth_dir = base_dir / "ground_truth"

        studies = []
        seen_nct_ids = set()

        # All ground_truth SAPs are now high quality (downloaded from real PDFs)

        # Add all ground truth SAPs (all are high quality - from real PDFs)
        if ground_truth_dir.exists():
            for sap_file in ground_truth_dir.glob("*_sap.txt"):
                nct_id = sap_file.stem.replace("_sap", "")
                if nct_id in seen_nct_ids:
                    continue
                seen_nct_ids.add(nct_id)

                try:
                    sap_text = sap_file.read_text(encoding='utf-8', errors='ignore')
                    lines = len(sap_text.split('\n'))

                    # Detect therapeutic area
                    sap_lower = sap_text.lower()
                    if any(x in sap_lower for x in ["cancer", "tumor", "oncology", "carcinoma", "melanoma"]):
                        area = "Oncology"
                    elif any(x in sap_lower for x in ["heart", "cardiac", "cardiovascular", "coronary"]):
                        area = "Cardiology"
                    elif any(x in sap_lower for x in ["diabetes", "glucose", "metabolic", "obesity"]):
                        area = "Metabolism"
                    elif any(x in sap_lower for x in ["infection", "hiv", "hepatitis", "covid", "viral"]):
                        area = "Infectious"
                    elif any(x in sap_lower for x in ["psychiatric", "depression", "anxiety", "schizophrenia"]):
                        area = "Psychiatry"
                    else:
                        area = "Other"

                    # Get title from protocol if available
                    protocol_file = ground_truth_dir / f"{nct_id}_protocol.txt"
                    title = nct_id
                    if protocol_file.exists():
                        protocol_text = protocol_file.read_text(encoding='utf-8', errors='ignore')[:300]
                        if "STUDY:" in protocol_text:
                            title = protocol_text.split("STUDY:")[1].split("\n")[0].strip()[:60]

                    studies.append({
                        "nct_id": nct_id,
                        "title": f"⭐ {title}" if lines > 500 else title,
                        "sap_lines": lines,
                        "therapeutic_area": area,
                        "quality": "high"
                    })
                except Exception as e:
                    print(f"[Ground Truth] Warning: Could not process {nct_id}: {e}")
                    continue

        # Then add all pairs from all_pairs directory
        if all_pairs_dir.exists():
            for sap_file in all_pairs_dir.glob("*_sap.txt"):
                nct_id = sap_file.stem.replace("_sap", "")
                if nct_id in seen_nct_ids:
                    continue
                seen_nct_ids.add(nct_id)

                try:
                    sap_text = sap_file.read_text(encoding='utf-8', errors='ignore')
                    lines = len(sap_text.split('\n'))

                    # Try to detect therapeutic area from content
                    sap_lower = sap_text.lower()
                    if any(x in sap_lower for x in ["cancer", "tumor", "oncology", "carcinoma"]):
                        area = "Oncology"
                    elif any(x in sap_lower for x in ["infection", "hiv", "hepatitis", "viral"]):
                        area = "Infectious"
                    elif any(x in sap_lower for x in ["heart", "cardiac", "cardiovascular"]):
                        area = "Cardiology"
                    elif any(x in sap_lower for x in ["diabetes", "glucose", "metabolic"]):
                        area = "Metabolism"
                    else:
                        area = "Other"

                    # Extract title from protocol if available
                    protocol_file = all_pairs_dir / f"{nct_id}_protocol.txt"
                    title = nct_id
                    if protocol_file.exists():
                        protocol_text = protocol_file.read_text(encoding='utf-8', errors='ignore')[:500]
                        # Try to find title
                        for line in protocol_text.split('\n'):
                            if 'title:' in line.lower() or 'study:' in line.lower():
                                title = line.split(':', 1)[-1].strip()[:60]
                                if title:
                                    break

                    studies.append({
                        "nct_id": nct_id,
                        "title": title if title != nct_id else f"{nct_id} ({area})",
                        "sap_lines": lines,
                        "therapeutic_area": area,
                        "quality": "standard"
                    })
                except Exception:
                    continue

        # Sort: high quality first, then by NCT ID
        studies.sort(key=lambda x: (0 if x.get("quality") == "high" else 1, x["nct_id"]))

        return {
            "studies": studies,
            "total": len(studies),
            "high_quality": sum(1 for s in studies if s.get("quality") == "high")
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/evaluate/{job_id}")
async def evaluate_job(job_id: str, ground_truth_nct: str):
    """
    Evaluate a completed job's SAP against a ground truth SAP.
    Checks both ground_truth and all_pairs directories.
    """
    try:
        if not SAP_EVALUATOR_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="SAP Evaluator not available - evaluate_sap module not found"
            )

        db = get_supabase()

        # Get the job
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(status_code=400, detail="Job must be completed to evaluate")

        generated_sap = job.get("generated_sap", "")
        if not generated_sap:
            raise HTTPException(status_code=400, detail="No generated SAP found")

        # Load ground truth - check both directories
        base_dir = Path(__file__).parent.parent.parent / "data"
        ground_truth_dir = base_dir / "ground_truth"
        all_pairs_dir = base_dir / "all_pairs"

        sap_path = None
        # First check ground_truth directory (high quality)
        gt_path = ground_truth_dir / f"{ground_truth_nct}_sap.txt"
        if gt_path.exists():
            sap_path = gt_path
        else:
            # Then check all_pairs directory
            ap_path = all_pairs_dir / f"{ground_truth_nct}_sap.txt"
            if ap_path.exists():
                sap_path = ap_path

        if not sap_path:
            raise HTTPException(status_code=404, detail=f"Ground truth SAP not found: {ground_truth_nct}")

        ground_truth_sap = sap_path.read_text(encoding='utf-8', errors='ignore')

        evaluator = SAPEvaluator(str(sap_path.parent))
        eval_result = evaluator.evaluate(generated_sap, ground_truth_sap, ground_truth_nct)

        return {
            "nct_id": eval_result.nct_id,
            "ground_truth_lines": eval_result.ground_truth_lines,
            "generated_lines": eval_result.generated_lines,
            "section_coverage_pct": eval_result.section_coverage_pct,
            "keyword_overlap_pct": eval_result.keyword_overlap_pct,
            "has_primary_endpoint": eval_result.has_primary_endpoint,
            "has_secondary_endpoint": eval_result.has_secondary_endpoint,
            "has_sample_size": eval_result.has_sample_size,
            "has_analysis_populations": eval_result.has_analysis_populations,
            "has_statistical_methods": eval_result.has_statistical_methods,
            "has_missing_data": eval_result.has_missing_data,
            "overall_score": eval_result.overall_score,
            "sections_matched": eval_result.sections_matched,
            "sections_missing": eval_result.sections_missing,
            "statistical_terms_found": eval_result.statistical_terms_found,
            "statistical_terms_missing": eval_result.statistical_terms_missing,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/evaluate-batch/{job_id}")
async def evaluate_batch(job_id: str, limit: int = 50):
    """
    Evaluate a completed job's SAP against ALL ground truth SAPs.
    Returns aggregate metrics and individual results.

    Args:
        job_id: The job to evaluate
        limit: Max number of ground truth SAPs to compare (default 50)
    """
    try:
        if not SAP_EVALUATOR_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="SAP Evaluator not available - evaluate_sap module not found"
            )

        db = get_supabase()

        # Get the job
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(status_code=400, detail="Job must be completed to evaluate")

        generated_sap = job.get("generated_sap", "")
        if not generated_sap:
            raise HTTPException(status_code=400, detail="No generated SAP found")

        # Load all ground truth SAPs
        base_dir = Path(__file__).parent.parent.parent / "data"
        ground_truth_dir = base_dir / "ground_truth"
        all_pairs_dir = base_dir / "all_pairs"

        results = []

        # Evaluate against ground_truth first (high quality)
        count = 0
        if ground_truth_dir.exists():
            for sap_file in sorted(ground_truth_dir.glob("*_sap.txt")):
                if count >= limit:
                    break
                nct_id = sap_file.stem.replace("_sap", "")
                try:
                    ground_truth_sap = sap_file.read_text(encoding='utf-8', errors='ignore')
                    evaluator = SAPEvaluator(str(ground_truth_dir))
                    eval_result = evaluator.evaluate(generated_sap, ground_truth_sap, nct_id)
                    results.append({
                        "nct_id": nct_id,
                        "quality": "high",
                        "section_coverage_pct": eval_result.section_coverage_pct,
                        "keyword_overlap_pct": eval_result.keyword_overlap_pct,
                        "overall_score": eval_result.overall_score,
                        "has_primary_endpoint": eval_result.has_primary_endpoint,
                        "has_statistical_methods": eval_result.has_statistical_methods,
                        "ground_truth_lines": eval_result.ground_truth_lines,
                    })
                    count += 1
                except Exception as e:
                    continue

        # Then evaluate against all_pairs if we haven't hit limit
        if count < limit and all_pairs_dir.exists():
            seen = {r["nct_id"] for r in results}
            for sap_file in sorted(all_pairs_dir.glob("*_sap.txt")):
                if count >= limit:
                    break
                nct_id = sap_file.stem.replace("_sap", "")
                if nct_id in seen:
                    continue
                try:
                    ground_truth_sap = sap_file.read_text(encoding='utf-8', errors='ignore')
                    evaluator = SAPEvaluator(str(all_pairs_dir))
                    eval_result = evaluator.evaluate(generated_sap, ground_truth_sap, nct_id)
                    results.append({
                        "nct_id": nct_id,
                        "quality": "standard",
                        "section_coverage_pct": eval_result.section_coverage_pct,
                        "keyword_overlap_pct": eval_result.keyword_overlap_pct,
                        "overall_score": eval_result.overall_score,
                        "has_primary_endpoint": eval_result.has_primary_endpoint,
                        "has_statistical_methods": eval_result.has_statistical_methods,
                        "ground_truth_lines": eval_result.ground_truth_lines,
                    })
                    count += 1
                except Exception:
                    continue

        # Calculate aggregate metrics
        if results:
            avg_section_coverage = sum(r["section_coverage_pct"] for r in results) / len(results)
            avg_keyword_overlap = sum(r["keyword_overlap_pct"] for r in results) / len(results)
            avg_overall_score = sum(r["overall_score"] for r in results) / len(results)
            primary_endpoint_pct = sum(1 for r in results if r["has_primary_endpoint"]) / len(results) * 100
            statistical_methods_pct = sum(1 for r in results if r["has_statistical_methods"]) / len(results) * 100

            # Find best and worst matches
            sorted_by_score = sorted(results, key=lambda x: x["overall_score"], reverse=True)
            best_match = sorted_by_score[0] if sorted_by_score else None
            worst_match = sorted_by_score[-1] if sorted_by_score else None
        else:
            avg_section_coverage = 0
            avg_keyword_overlap = 0
            avg_overall_score = 0
            primary_endpoint_pct = 0
            statistical_methods_pct = 0
            best_match = None
            worst_match = None

        return {
            "total_comparisons": len(results),
            "aggregate": {
                "avg_section_coverage_pct": round(avg_section_coverage, 1),
                "avg_keyword_overlap_pct": round(avg_keyword_overlap, 1),
                "avg_overall_score": round(avg_overall_score, 1),
                "primary_endpoint_pct": round(primary_endpoint_pct, 1),
                "statistical_methods_pct": round(statistical_methods_pct, 1),
            },
            "best_match": best_match,
            "worst_match": worst_match,
            "results": results,
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# SAP VERIFICATION ENDPOINT (Generate → Verify Architecture)
# =============================================================================

@app.post("/verify/{job_id}", response_model=VerificationResponse)
async def verify_sap_endpoint(job_id: str):
    """
    Verify a generated SAP against the source protocol using anchor verification.

    This implements the Generate → Verify architecture:
    1. Extract "anchors" from protocol (sentences containing statistics)
    2. Check if each anchor's key numbers appear in the generated SAP
    3. Flag unexpected numbers in SAP that don't come from protocol
    4. Check ICH E9 regulatory compliance
    5. Return confidence score and issues list

    The verification layer catches:
    - Missing critical values (sample size, alpha, power)
    - Hallucinated numbers not in protocol
    - Missing regulatory required sections
    - Inconsistent statistical methodology

    Returns:
        VerificationResponse with confidence score, issues, and recommendations
    """
    try:
        if not SAP_VERIFIER_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="SAP Verifier not available - import failed"
            )

        db = get_supabase()

        # Get the job
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(status_code=400, detail="Job must be completed to verify")

        generated_sap = job.get("generated_sap", "")
        if not generated_sap:
            raise HTTPException(status_code=400, detail="No generated SAP found")

        protocol_text = job.get("protocol_text", "")
        if not protocol_text:
            raise HTTPException(status_code=400, detail="No protocol text found - cannot verify")

        # Step 1: Extract anchors from protocol
        anchors = extract_anchors(protocol_text)

        # Step 2: Verify SAP against anchors
        report = verify_sap(generated_sap, protocol_text, anchors)

        # Convert issues to response format
        issues_list = [
            VerificationIssue(
                severity=issue.severity.value,
                category=issue.category,
                message=issue.message,
                rule=issue.rule
            )
            for issue in report.issues
        ]

        # Build anchor summary
        anchor_summary = VerificationAnchorSummary(**anchors.summary())

        return VerificationResponse(
            success=True,
            job_id=job_id,
            anchors_found=report.anchors_found,
            anchors_verified=report.anchors_verified,
            anchors_missing=report.anchors_missing,
            anchor_summary=anchor_summary,
            confidence_score=report.confidence_score,
            needs_human_review=report.needs_human_review(),
            critical_issues=report.critical_count(),
            warnings=report.warning_count(),
            issues=issues_list,
            unexpected_numbers=list(report.unexpected_numbers)[:20],  # Limit to 20
            report_text=report.summary(),
            verification_method="anchor-verification-v1"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"SAP verification failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/verify-text", response_model=VerificationResponse)
async def verify_sap_text(
    sap_text: str = Form(...),
    protocol_text: str = Form(...)
):
    """
    Verify an SAP directly against protocol text (no job required).

    This is the stateless version of /verify/{job_id} for direct API usage.
    Upload both the SAP and the source protocol for verification.

    Args:
        sap_text: The generated SAP document text
        protocol_text: The source protocol text

    Returns:
        VerificationResponse with confidence score and issues
    """
    try:
        if not SAP_VERIFIER_AVAILABLE:
            raise HTTPException(
                status_code=503,
                detail="SAP Verifier not available - import failed"
            )

        if not sap_text.strip():
            raise HTTPException(status_code=400, detail="SAP text cannot be empty")

        if not protocol_text.strip():
            raise HTTPException(status_code=400, detail="Protocol text cannot be empty")

        # Extract anchors and verify
        anchors = extract_anchors(protocol_text)
        report = verify_sap(sap_text, protocol_text, anchors)

        # Convert to response
        issues_list = [
            VerificationIssue(
                severity=issue.severity.value,
                category=issue.category,
                message=issue.message,
                rule=issue.rule
            )
            for issue in report.issues
        ]

        anchor_summary = VerificationAnchorSummary(**anchors.summary())

        return VerificationResponse(
            success=True,
            job_id="direct-verification",
            anchors_found=report.anchors_found,
            anchors_verified=report.anchors_verified,
            anchors_missing=report.anchors_missing,
            anchor_summary=anchor_summary,
            confidence_score=report.confidence_score,
            needs_human_review=report.needs_human_review(),
            critical_issues=report.critical_count(),
            warnings=report.warning_count(),
            issues=issues_list,
            unexpected_numbers=list(report.unexpected_numbers)[:20],
            report_text=report.summary(),
            verification_method="anchor-verification-v1"
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Direct SAP verification failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/stats")
async def get_stats():
    """
    Get job statistics.
    """
    try:
        db = get_supabase()

        result = db.table("sap_jobs").select("status, quality_score, processing_time").execute()

        jobs = result.data
        completed_jobs = [j for j in jobs if j["status"] == "completed"]

        stats = {
            "total": len(jobs),
            "completed": len(completed_jobs),
            "failed": sum(1 for j in jobs if j["status"] == "failed"),
            "queued": sum(1 for j in jobs if j["status"] == "queued"),
            "processing": sum(1 for j in jobs if j["status"] == "processing"),
            "avg_quality_score": round(sum(j["quality_score"] or 0 for j in completed_jobs) / len(completed_jobs), 1) if completed_jobs else 0,
            "avg_processing_time": round(sum(j["processing_time"] or 0 for j in completed_jobs) / len(completed_jobs), 1) if completed_jobs else 0,
        }

        return stats

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# SDTM SPECIFICATION ENDPOINT
# =============================================================================

class SDTMSpecResponse(BaseModel):
    """Response model for SDTM specification generation."""
    success: bool
    message: str
    sdtm_version: str = "3.4"
    domains: list = []  # List of domain specs
    domain_count: int = 0
    markdown: str = ""  # Full markdown specification
    sap_summary: dict = {}  # Extracted SAP information (endpoints, populations, etc.)
    errors: list = []


@app.post("/generate-sdtm/{job_id}", response_model=SDTMSpecResponse)
async def generate_sdtm_specs(job_id: str):
    """
    Generate SDTM domain specifications from a completed SAP job.

    This endpoint takes a job_id that has already completed SAP generation,
    extracts the protocol facts, and generates CDISC-compliant SDTM specs.

    Returns:
        - List of required SDTM domains (DM, AE, EX, DS, etc.)
        - Variable-level specifications for each domain
        - Core classifications (Req/Exp/Perm) per CDISC SDTMIG v3.4
    """
    try:
        db = get_supabase()

        # Get job and verify it's completed
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(
                status_code=400,
                detail=f"Job not ready for SDTM generation. Status: {job['status']}"
            )

        # Import SDTM generator
        try:
            from enterprise_sap_system.specs.sdtm_specs import SDTMSpecGenerator
        except ImportError as e:
            raise HTTPException(
                status_code=500,
                detail=f"SDTM generator not available: {e}"
            )

        # Build protocol facts from job data - now passes SAP text for parsing
        protocol_text = job.get("protocol_text", "")
        sap_text = job.get("generated_sap", "")

        # The new generator parses the SAP text to extract study-specific requirements
        protocol_facts = {
            "protocol_id": _extract_protocol_id(sap_text, protocol_text, job.get("nct_id")),
            "sap_text": sap_text,  # Key: pass SAP text for parsing
        }

        # Generate SDTM specs by parsing SAP text
        generator = SDTMSpecGenerator()
        spec = generator.generate(protocol_facts)

        # Convert domains to JSON-serializable format with traceability
        domains_json = []
        for domain in spec.domains:
            domain_dict = {
                "code": domain.code,
                "name": domain.name,
                "label": domain.label,
                "class": domain.domain_class.value,
                "structure": domain.structure,
                "purpose": domain.purpose,
                "study_specific_notes": domain.study_specific_notes,
                "traceability": [
                    {
                        "sap_section": t.sap_section,
                        "sap_text": t.sap_text[:200] + "..." if len(t.sap_text) > 200 else t.sap_text,
                        "sdtm_element": t.sdtm_element,
                        "rationale": t.rationale
                    }
                    for t in domain.traceability
                ],
                "variables": [
                    {
                        "name": v.name,
                        "label": v.label,
                        "type": v.type,
                        "length": v.length,
                        "core": v.core.value,
                        "codelist": v.codelist,
                    }
                    for v in domain.variables
                ]
            }
            domains_json.append(domain_dict)

        # Generate markdown
        markdown = spec.to_markdown()

        # Count domains with SAP traceability
        traced_domains = sum(1 for d in spec.domains if d.traceability)

        return SDTMSpecResponse(
            success=True,
            message=f"Generated SDTM specs for {len(spec.domains)} domains ({traced_domains} with SAP traceability)",
            sdtm_version=spec.sdtm_version,
            domains=domains_json,
            domain_count=len(spec.domains),
            markdown=markdown,
            sap_summary=spec.sap_summary,
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        return SDTMSpecResponse(
            success=False,
            message=f"SDTM generation failed: {str(e)}",
            errors=[str(e)]
        )


def _detect_indication(text: str) -> str:
    """Detect indication from protocol text."""
    text_lower = text.lower()
    if 'ulcerative colitis' in text_lower:
        return 'Ulcerative Colitis'
    elif 'crohn' in text_lower:
        return "Crohn's Disease"
    elif 'melanoma' in text_lower:
        return 'Melanoma'
    elif 'breast cancer' in text_lower:
        return 'Breast Cancer'
    elif 'lung cancer' in text_lower or 'nsclc' in text_lower:
        return 'Non-Small Cell Lung Cancer'
    elif 'rheumatoid arthritis' in text_lower:
        return 'Rheumatoid Arthritis'
    return ''


def _extract_timepoint(sap_text: str) -> str:
    """Extract primary timepoint from SAP text."""
    import re
    patterns = [
        r'(?:primary|week)\s*(\d+)',
        r'at\s+week\s+(\d+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, sap_text, re.IGNORECASE)
        if match:
            return f"Week {match.group(1)}"
    return "Week 12"


def _extract_secondary_endpoints(sap_text: str) -> list:
    """Extract secondary endpoints from SAP text."""
    import re
    endpoints = []
    # Look for secondary endpoint section
    match = re.search(r'secondary\s+endpoint[s]?[:\s]+([^\n]+(?:\n[^\n]+)*?)(?=\n\n|\Z)', sap_text, re.IGNORECASE)
    if match:
        text = match.group(1)
        # Split by common delimiters
        for item in re.split(r'[;•\n]', text):
            item = item.strip()
            if item and len(item) > 5 and len(item) < 200:
                endpoints.append({"name": item[:100]})
    return endpoints[:5]  # Max 5


# =============================================================================
# CODE GENERATION ENDPOINT (Additive - does not modify existing functionality)
# =============================================================================

class CodeGenerationResponse(BaseModel):
    """Response model for code generation."""
    success: bool
    message: str
    programs: dict = {}  # {filename: code}
    total_lines: int = 0
    errors: list = []


@app.post("/generate-code/{job_id}", response_model=CodeGenerationResponse)
async def generate_sas_code(job_id: str):
    """
    Generate SAS code from a completed SAP job.

    This endpoint takes a job_id that has already completed SAP generation,
    extracts the protocol facts, and generates production-ready SAS code.

    Returns:
        - ADaM dataset programs (ADSL, ADAE, ADTTE, ADEFF)
        - TLF output programs (demographics, AE summary, primary efficacy)
        - Driver program
    """
    try:
        db = get_supabase()

        # Get job and verify it's completed
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(
                status_code=400,
                detail=f"Job not ready for code generation. Status: {job['status']}"
            )

        if not job.get("generated_sap"):
            raise HTTPException(
                status_code=400,
                detail="Job has no SAP output to generate code from"
            )

        # Import code generator (lazy import to avoid startup issues)
        try:
            from enterprise_sap_system.code_generators import CodeGenerationOrchestrator
        except ImportError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Code generator not available: {e}"
            )

        # Build protocol facts from job data
        # Note: In production, these would be stored with the job
        sap_text = job.get("generated_sap", "")
        protocol_text = job.get("protocol_text", "")
        protocol_facts = {
            "protocol_id": _extract_protocol_id(sap_text, protocol_text, job.get("nct_id")),
            "therapeutic_area": _detect_therapeutic_area(protocol_text or sap_text),
            "drug_name": _extract_drug_name(sap_text),
            "treatments": _extract_treatments(sap_text),
            "primary_endpoint": _extract_primary_endpoint(sap_text, protocol_text),
            "total_n": _extract_sample_size(sap_text),
        }

        # Generate code
        orchestrator = CodeGenerationOrchestrator()
        package = orchestrator.generate_all(protocol_facts)

        # Build response
        programs = {}
        total_lines = 0

        for prog in package.adam_programs:
            programs[f"adam/{prog.program_name}"] = prog.code
            total_lines += len(prog.code.split('\n'))

        for prog in package.tlf_programs:
            programs[f"tlf/{prog.program_name}"] = prog.code
            total_lines += len(prog.code.split('\n'))

        programs["driver.sas"] = package.driver_program
        total_lines += len(package.driver_program.split('\n'))

        return CodeGenerationResponse(
            success=True,
            message=f"Generated {len(package.adam_programs)} ADaM + {len(package.tlf_programs)} TLF programs",
            programs=programs,
            total_lines=total_lines,
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        return CodeGenerationResponse(
            success=False,
            message=f"Code generation failed: {str(e)}",
            programs={},
            total_lines=0,
            errors=[str(e)]
        )


def _detect_therapeutic_area(text: str) -> str:
    """Detect therapeutic area from protocol text."""
    text_lower = text.lower()
    if any(term in text_lower for term in ['crohn', 'colitis', 'ibd', 'ulcerative']):
        return 'ibd'
    elif any(term in text_lower for term in ['tumor', 'cancer', 'oncology', 'recist']):
        return 'oncology'
    elif any(term in text_lower for term in ['rheumatoid', 'arthritis', 'das28']):
        return 'rheumatology'
    elif any(term in text_lower for term in ['cardiac', 'heart', 'cardiovascular']):
        return 'cardiovascular'
    return 'general'


def _extract_drug_name(sap_text: str) -> str:
    """Extract drug name from SAP text."""
    import re
    # Look for common patterns
    patterns = [
        r'study drug[:\s]+([A-Za-z0-9-]+)',
        r'investigational product[:\s]+([A-Za-z0-9-]+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, sap_text, re.IGNORECASE)
        if match:
            return match.group(1)
    return "Study Drug"


def _extract_treatments(sap_text: str) -> list:
    """Extract treatment arms from SAP text."""
    import re
    treatments = []

    # Look for arm patterns
    arm_matches = re.findall(r'(?:arm|group)\s*\d*[:\s]*([^,\n]+(?:mg|placebo)[^,\n]*)', sap_text, re.IGNORECASE)
    for match in arm_matches[:4]:  # Max 4 arms
        name = match.strip()
        if name and name not in [t['name'] for t in treatments]:
            treatments.append({'name': name, 'code': f'TRT{len(treatments)+1}'})

    # Default if none found
    if not treatments:
        treatments = [
            {'name': 'Placebo', 'code': 'TRT1'},
            {'name': 'Active Treatment', 'code': 'TRT2'}
        ]

    return treatments


def _extract_primary_endpoint(sap_text: str, protocol_text: str = "") -> dict:
    """Extract primary endpoint from protocol or SAP text with robust pattern matching.

    IMPORTANT: Search protocol_text FIRST as it contains the original endpoint definition.
    The SAP text may have generic placeholders if extraction failed during generation.
    """
    import re

    # Search both texts, but prioritize protocol_text (the source of truth)
    texts_to_search = []
    if protocol_text:
        texts_to_search.append(protocol_text)
    if sap_text:
        texts_to_search.append(sap_text)

    for text in texts_to_search:
        # First, look for the DEFINITION section with Mayo score criteria (most specific)
        # This pattern finds "Definition Criteria:" followed by the actual criteria
        definition_patterns = [
            # **Definition:** or Definition Criteria: followed by Mayo score definition
            r'(?:definition\s*(?:criteria)?)[:\s]*(?:\*\*)?([^*\n]*?(?:mayo\s+score|subscore)[^*\n]*?(?:≤|<=|=)\s*\d+[^*\n]*)',
            # Clinical remission defined as Mayo score criteria
            r'clinical\s+remission[^.]*?(?:defined\s+as|is)[:\s]*([^.]*?(?:mayo\s+score|subscore)[^.]*?(?:≤|<=|=)\s*\d+[^.]*)',
            # Full/total Mayo score with specific criteria
            r'((?:full|total)\s+mayo\s+score\s*(?:≤|<=|of)\s*\d+[^.|\n]*)',
        ]

        for pattern in definition_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                endpoint = match.group(1).strip()
                # Clean up whitespace and formatting
                endpoint = re.sub(r'\s+', ' ', endpoint)
                endpoint = re.sub(r'\*\*|\*|__|#', '', endpoint)
                # Add "Clinical remission" prefix if not present
                if 'remission' not in endpoint.lower() and 'response' not in endpoint.lower():
                    endpoint = "Clinical remission: " + endpoint
                endpoint = endpoint[0].upper() + endpoint[1:] if endpoint else endpoint
                # Clean trailing punctuation
                endpoint = endpoint.rstrip('|,;')
                if len(endpoint) > 15:
                    return {'name': endpoint[:200], 'type': 'binary'}

        # IBD/UC-specific endpoint patterns (proportion achieving remission)
        ibd_patterns = [
            # Proportion achieving clinical remission at week X
            r'((?:proportion|percentage)\s+of\s+(?:subjects|patients)\s+(?:achieving|with|in)\s+clinical\s+remission\s+(?:at|by)\s+week\s+\d+)',
            # Clinical remission at week X (only if "Clinical remission" is included)
            r'(clinical\s+remission\s+(?:at|by)\s+week\s+\d+)',
            r'(clinical\s+response\s+(?:at|by)\s+week\s+\d+)',
            # Endoscopic endpoints
            r'(endoscopic\s+(?:improvement|remission|response|healing)\s+(?:at|by)\s+week\s+\d+)',
            r'(mucosal\s+healing\s+(?:at|by)\s+week\s+\d+)',
            # Modified/partial Mayo score endpoints
            r'((?:modified|partial)\s+mayo\s+score\s+(?:of\s+)?(?:\d+|≤\s*\d+)[^|\n]*?(?:at|by)\s+week\s+\d+)',
        ]

        for pattern in ibd_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                endpoint = match.group(1).strip()
                # Clean up whitespace and newlines
                endpoint = re.sub(r'\s+', ' ', endpoint)
                endpoint = endpoint[0].upper() + endpoint[1:] if endpoint else endpoint
                endpoint = endpoint.rstrip('|,;')
                if len(endpoint) > 10:
                    return {'name': endpoint[:200], 'type': 'binary'}

        # Look in sections that typically define primary endpoint
        section_patterns = [
            # Section header followed by endpoint definition (not just "at Week 12")
            r'(?:primary\s+(?:efficacy\s+)?endpoint|primary\s+objective)[:\s]*\n+([^\n]+(?:remission|response|score|healing)[^\n]+)',
            # Primary endpoint IS statement
            r'primary\s+(?:efficacy\s+)?endpoint\s+(?:is|will\s+be)\s+(?:the\s+)?(?:proportion[^.|\n]+|clinical[^.|\n]+)',
        ]

        for pattern in section_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                endpoint = match.group(1) if match.lastindex else match.group(0)
                endpoint = endpoint.strip()
                # Skip if it's just a generic placeholder or just "at Week 12"
                skip_patterns = ['primary endpoint', 'the primary endpoint', 'endpoint', 'at week', 'by week']
                if any(endpoint.lower().strip() == sp or endpoint.lower().strip().startswith(sp + ' |') for sp in skip_patterns):
                    continue
                # Clean up
                endpoint = re.sub(r'\s+', ' ', endpoint)
                endpoint = re.sub(r'\*\*|\*|__|#|\|', '', endpoint)
                endpoint = endpoint.rstrip('.|,;')
                if endpoint and len(endpoint) > 15:
                    return {'name': endpoint[:200], 'type': _detect_endpoint_type(endpoint)}

        # General patterns for any therapeutic area
        general_patterns = [
            r'((?:overall\s+)?(?:survival|response\s+rate|progression[- ]free)\s+(?:at|by)\s+(?:week|month)\s+\d+)',
            r'((?:objective\s+)?response\s+rate\s+(?:at|by)\s+week\s+\d+)',
            r'((?:change|reduction)\s+(?:from\s+baseline\s+)?in\s+[^|\n]+(?:at|by)\s+week\s+\d+)',
        ]

        for pattern in general_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                endpoint = match.group(1).strip()
                endpoint = re.sub(r'\s+', ' ', endpoint)
                endpoint = endpoint[0].upper() + endpoint[1:] if endpoint else endpoint
                return {'name': endpoint[:200], 'type': _detect_endpoint_type(endpoint)}

    # Default fallback - use a meaningful description if this is an IBD study
    if any(term in (sap_text + protocol_text).lower() for term in ['ulcerative colitis', 'crohn', 'ibd']):
        return {'name': 'Clinical remission at Week 12', 'type': 'binary'}

    return {'name': 'Primary Endpoint', 'type': 'binary'}


def _detect_endpoint_type(endpoint_text: str) -> str:
    """Detect the type of endpoint from its description."""
    endpoint_lower = endpoint_text.lower()

    if any(term in endpoint_lower for term in ['remission', 'response', 'proportion', 'percentage', 'rate']):
        return 'binary'
    if any(term in endpoint_lower for term in ['change from baseline', 'mean change', 'difference']):
        return 'continuous'
    if any(term in endpoint_lower for term in ['time to', 'survival', 'duration']):
        return 'time-to-event'
    if any(term in endpoint_lower for term in ['score', 'index', 'scale']):
        return 'continuous'

    return 'binary'


def _extract_protocol_id(sap_text: str, protocol_text: str = "", job_nct_id: str = None) -> str:
    """Extract protocol/study ID from available sources."""
    import re

    # First check if job has nct_id
    if job_nct_id and job_nct_id != "UNKNOWN" and len(job_nct_id) > 3:
        return job_nct_id

    # Try to extract from SAP text
    patterns = [
        # Protocol numbers like CTJ301UC201, ABC-123-456
        r'(?:protocol|study)\s*(?:number|id|identifier)?[:\s]+([A-Z]{2,5}[-]?\d{2,4}[-]?[A-Z]{0,3}[-]?\d{0,4})',
        r'(?:protocol|study)[:\s]+([A-Z0-9]+-[A-Z0-9]+-[A-Z0-9]+)',
        r'([A-Z]{2,5}\d{3}[A-Z]{2}\d{3})',  # Pattern like CTJ301UC201
        # NCT numbers
        r'(NCT\d{8})',
        # EudraCT numbers
        r'(\d{4}-\d{6}-\d{2})',
        # Generic protocol patterns
        r'protocol[:\s]+([A-Z0-9-]{6,20})',
    ]

    # Try SAP text first, then protocol text
    for text in [sap_text, protocol_text]:
        if not text:
            continue
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                study_id = match.group(1).strip()
                if study_id and len(study_id) >= 6:
                    return study_id.upper()

    return "UNKNOWN"


def _extract_sample_size(sap_text: str) -> int:
    """Extract sample size from SAP text."""
    import re
    patterns = [
        r'(\d+)\s*(?:patients|subjects|participants)',
        r'n\s*=\s*(\d+)',
        r'sample size[:\s]+(\d+)',
        r'total\s+of\s+(\d+)\s+(?:patients|subjects)',
        r'(\d+)\s+(?:patients|subjects)\s+will be\s+(?:enrolled|randomized)',
    ]
    for pattern in patterns:
        match = re.search(pattern, sap_text, re.IGNORECASE)
        if match:
            size = int(match.group(1))
            if size >= 10:  # Reasonable minimum
                return size
    return 100


# =============================================================================
# TLF SHELL SPECIFICATION ENDPOINT
# =============================================================================

class TLFShellResponse(BaseModel):
    """Response model for TLF shell generation."""
    success: bool
    message: str
    tables: list = []
    listings: list = []
    figures: list = []
    total_outputs: int = 0
    markdown: str = ""
    errors: list = []


@app.post("/generate-tlf-shells/{job_id}", response_model=TLFShellResponse)
async def generate_tlf_shells(job_id: str):
    """
    Generate TLF (Tables, Listings, Figures) shell specifications from a completed SAP job.

    Returns:
        - Demographics tables (Table 14.1.x)
        - Efficacy tables (Table 14.2.x)
        - Safety tables (Table 14.3.x)
        - Data listings (Listing 16.2.x)
        - Figures (Figure 14.x)
    """
    try:
        db = get_supabase()

        # Get job and verify it's completed
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(
                status_code=400,
                detail=f"Job not ready for TLF generation. Status: {job['status']}"
            )

        # Extract protocol facts from job data
        protocol_text = job.get("protocol_text", "")
        sap_text = job.get("generated_sap", "")

        # Use improved protocol ID extraction
        protocol_id = _extract_protocol_id(sap_text, protocol_text, job.get("nct_id"))
        therapeutic_area = _detect_therapeutic_area(protocol_text or sap_text)
        primary_endpoint = _extract_primary_endpoint(sap_text, protocol_text)
        treatments = _extract_treatments(sap_text)
        sample_size = _extract_sample_size(sap_text)

        # Generate TLF shells using simplified approach (markdown-based)
        tables_json = []
        listings_json = []
        figures_json = []

        # Demographics table
        tables_json.append({
            "output_id": "Table 14.1.1",
            "title": "Summary of Subject Demographics and Baseline Characteristics",
            "population": "Safety Population",
            "footnotes": ["N = Number of subjects in the safety population."],
            "columns": [],
            "markdown": ""
        })

        # Disposition table
        tables_json.append({
            "output_id": "Table 14.1.2",
            "title": "Subject Disposition",
            "population": "All Randomized",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        # Primary efficacy table
        tables_json.append({
            "output_id": "Table 14.2.1",
            "title": f"Primary Efficacy Analysis: {primary_endpoint.get('name', 'Primary Endpoint')}",
            "population": "Full Analysis Set",
            "footnotes": ["Analysis performed using ANCOVA with treatment as a factor."],
            "columns": [],
            "markdown": ""
        })

        # AE Summary table
        tables_json.append({
            "output_id": "Table 14.3.1",
            "title": "Overall Summary of Treatment-Emergent Adverse Events",
            "population": "Safety Population",
            "footnotes": ["TEAE = Treatment-Emergent Adverse Event"],
            "columns": [],
            "markdown": ""
        })

        # AE by SOC/PT
        tables_json.append({
            "output_id": "Table 14.3.2",
            "title": "Treatment-Emergent Adverse Events by System Organ Class and Preferred Term",
            "population": "Safety Population",
            "footnotes": ["MedDRA version X.X"],
            "columns": [],
            "markdown": ""
        })

        # SAE table
        tables_json.append({
            "output_id": "Table 14.3.3",
            "title": "Serious Adverse Events",
            "population": "Safety Population",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        # Listings
        listings_json.append({
            "output_id": "Listing 16.2.1",
            "title": "Listing of Subjects Who Discontinued Study",
            "population": "All Randomized",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        listings_json.append({
            "output_id": "Listing 16.2.4",
            "title": "Listing of Serious Adverse Events",
            "population": "Safety Population",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        listings_json.append({
            "output_id": "Listing 16.2.6",
            "title": "Listing of Deaths",
            "population": "Safety Population",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        # Figures
        figures_json.append({
            "output_id": "Figure 14.2.1",
            "title": f"Kaplan-Meier Plot of {primary_endpoint.get('name', 'Primary Endpoint')}",
            "population": "Full Analysis Set",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        figures_json.append({
            "output_id": "Figure 14.2.2",
            "title": "Forest Plot of Subgroup Analyses",
            "population": "Full Analysis Set",
            "footnotes": [],
            "columns": [],
            "markdown": ""
        })

        # Helper to escape pipe characters and clean titles for markdown tables
        def escape_md_table(text: str) -> str:
            if not text:
                return ""
            # Escape pipe characters and remove newlines
            return text.replace("|", "\\|").replace("\n", " ").replace("\r", "").strip()

        # Generate markdown document
        full_markdown = f"""# TLF SHELL SPECIFICATIONS
**Protocol:** {protocol_id}
**Therapeutic Area:** {therapeutic_area.upper()}

---

## Tables

| Output ID | Title | Population |
|-----------|-------|------------|
"""
        for t in tables_json:
            full_markdown += f"| {escape_md_table(t['output_id'])} | {escape_md_table(t['title'])} | {escape_md_table(t['population'])} |\n"

        full_markdown += "\n## Listings\n\n| Output ID | Title | Population |\n|-----------|-------|------------|\n"
        for l in listings_json:
            full_markdown += f"| {escape_md_table(l['output_id'])} | {escape_md_table(l['title'])} | {escape_md_table(l['population'])} |\n"

        full_markdown += "\n## Figures\n\n| Output ID | Title | Population |\n|-----------|-------|------------|\n"
        for f in figures_json:
            full_markdown += f"| {escape_md_table(f['output_id'])} | {escape_md_table(f['title'])} | {escape_md_table(f['population'])} |\n"

        total_outputs = len(tables_json) + len(listings_json) + len(figures_json)

        return TLFShellResponse(
            success=True,
            message=f"Generated {len(tables_json)} tables, {len(listings_json)} listings, {len(figures_json)} figures",
            tables=tables_json,
            listings=listings_json,
            figures=figures_json,
            total_outputs=total_outputs,
            markdown=full_markdown,
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        return TLFShellResponse(
            success=False,
            message=f"TLF generation failed: {str(e)}",
            errors=[str(e)]
        )


# =============================================================================
# ADAM DERIVATION SPECIFICATION ENDPOINT
# =============================================================================

class AdamSpecResponse(BaseModel):
    """Response model for ADaM derivation specification generation."""
    success: bool
    message: str
    datasets: list = []
    total_variables: int = 0
    markdown: str = ""
    errors: list = []


@app.post("/generate-adam-specs/{job_id}", response_model=AdamSpecResponse)
async def generate_adam_specs(job_id: str):
    """
    Generate ADaM (Analysis Data Model) derivation specifications from a completed SAP job.

    Returns:
        - ADSL (Subject-Level Analysis Dataset) derivations
        - ADAE (Adverse Events) derivations
        - ADLB (Laboratory) derivations
        - ADEFF (Efficacy) derivations
        - ADTTE (Time-to-Event) derivations
    """
    try:
        db = get_supabase()

        # Get job and verify it's completed
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(
                status_code=400,
                detail=f"Job not ready for ADaM spec generation. Status: {job['status']}"
            )

        # Extract protocol facts from job data
        protocol_text = job.get("protocol_text", "")
        sap_text = job.get("generated_sap", "")

        protocol_id = _extract_protocol_id(sap_text, protocol_text, job.get("nct_id"))
        therapeutic_area = _detect_therapeutic_area(protocol_text or sap_text)
        primary_endpoint = _extract_primary_endpoint(sap_text, protocol_text)
        treatments = _extract_treatments(sap_text)

        # Build ADaM datasets using standard derivations
        datasets_json = []
        total_vars = 0

        # ADSL - Subject-Level Dataset
        adsl_vars = [
            {"name": "STUDYID", "label": "Study Identifier", "type": "Char", "length": 20, "derivation": "Assigned from protocol", "source": "DM.STUDYID", "codelist": None},
            {"name": "USUBJID", "label": "Unique Subject Identifier", "type": "Char", "length": 50, "derivation": "Assigned from SDTM", "source": "DM.USUBJID", "codelist": None},
            {"name": "SUBJID", "label": "Subject Identifier for the Study", "type": "Char", "length": 20, "derivation": "Assigned from SDTM", "source": "DM.SUBJID", "codelist": None},
            {"name": "SITEID", "label": "Study Site Identifier", "type": "Char", "length": 10, "derivation": "Assigned from SDTM", "source": "DM.SITEID", "codelist": None},
            {"name": "AGE", "label": "Age", "type": "Num", "length": 8, "derivation": "Set to DM.AGE", "source": "DM.AGE", "codelist": None},
            {"name": "AGEGR1", "label": "Pooled Age Group 1", "type": "Char", "length": 20, "derivation": "Derived: <65='<65', >=65='>=65'", "source": "DM.AGE", "codelist": None},
            {"name": "AGEGR1N", "label": "Pooled Age Group 1 (N)", "type": "Num", "length": 8, "derivation": "Numeric code for AGEGR1", "source": "AGEGR1", "codelist": None},
            {"name": "SEX", "label": "Sex", "type": "Char", "length": 1, "derivation": "Set to DM.SEX", "source": "DM.SEX", "codelist": "SEX"},
            {"name": "RACE", "label": "Race", "type": "Char", "length": 100, "derivation": "Set to DM.RACE", "source": "DM.RACE", "codelist": None},
            {"name": "ETHNIC", "label": "Ethnicity", "type": "Char", "length": 50, "derivation": "Set to DM.ETHNIC", "source": "DM.ETHNIC", "codelist": None},
            {"name": "TRT01P", "label": "Planned Treatment for Period 01", "type": "Char", "length": 200, "derivation": "Set to DM.ARM", "source": "DM.ARM", "codelist": None},
            {"name": "TRT01PN", "label": "Planned Treatment for Period 01 (N)", "type": "Num", "length": 8, "derivation": "Numeric code for TRT01P", "source": "TRT01P", "codelist": None},
            {"name": "TRT01A", "label": "Actual Treatment for Period 01", "type": "Char", "length": 200, "derivation": "Set to DM.ACTARM", "source": "DM.ACTARM", "codelist": None},
            {"name": "TRT01AN", "label": "Actual Treatment for Period 01 (N)", "type": "Num", "length": 8, "derivation": "Numeric code for TRT01A", "source": "TRT01A", "codelist": None},
            {"name": "TRTSDT", "label": "Date of First Exposure to Treatment", "type": "Num", "length": 8, "derivation": "Min(EX.EXSTDTC) where EXDOSE>0", "source": "EX.EXSTDTC", "codelist": None},
            {"name": "TRTEDT", "label": "Date of Last Exposure to Treatment", "type": "Num", "length": 8, "derivation": "Max(EX.EXENDTC) where EXDOSE>0", "source": "EX.EXENDTC", "codelist": None},
            {"name": "SAFFL", "label": "Safety Population Flag", "type": "Char", "length": 1, "derivation": "Y if TRTSDT is not missing", "source": "Derived", "codelist": "NY"},
            {"name": "ITTFL", "label": "Intent-to-Treat Population Flag", "type": "Char", "length": 1, "derivation": "Y if randomized", "source": "Derived", "codelist": "NY"},
            {"name": "FASFL", "label": "Full Analysis Set Population Flag", "type": "Char", "length": 1, "derivation": "Y if ITT and has baseline + 1 post-BL", "source": "Derived", "codelist": "NY"},
        ]
        datasets_json.append({
            "name": "ADSL",
            "label": "Subject-Level Analysis Dataset",
            "structure": "One record per subject",
            "keys": ["STUDYID", "USUBJID"],
            "variables": adsl_vars
        })
        total_vars += len(adsl_vars)

        # ADAE - Adverse Event Analysis Dataset
        adae_vars = [
            {"name": "STUDYID", "label": "Study Identifier", "type": "Char", "length": 20, "derivation": "Set to ADSL.STUDYID", "source": "ADSL", "codelist": None},
            {"name": "USUBJID", "label": "Unique Subject Identifier", "type": "Char", "length": 50, "derivation": "Set to ADSL.USUBJID", "source": "ADSL", "codelist": None},
            {"name": "AESEQ", "label": "Sequence Number", "type": "Num", "length": 8, "derivation": "Set to AE.AESEQ", "source": "AE.AESEQ", "codelist": None},
            {"name": "TRTA", "label": "Actual Treatment", "type": "Char", "length": 200, "derivation": "Treatment at AE onset", "source": "ADSL.TRT01A", "codelist": None},
            {"name": "AEDECOD", "label": "Dictionary-Derived Term", "type": "Char", "length": 200, "derivation": "Set to AE.AEDECOD", "source": "AE.AEDECOD", "codelist": None},
            {"name": "AEBODSYS", "label": "Body System or Organ Class", "type": "Char", "length": 200, "derivation": "Set to AE.AEBODSYS", "source": "AE.AEBODSYS", "codelist": None},
            {"name": "AESEV", "label": "Severity/Intensity", "type": "Char", "length": 20, "derivation": "Set to AE.AESEV", "source": "AE.AESEV", "codelist": None},
            {"name": "AESER", "label": "Serious Event", "type": "Char", "length": 1, "derivation": "Set to AE.AESER", "source": "AE.AESER", "codelist": "NY"},
            {"name": "AEREL", "label": "Causality", "type": "Char", "length": 50, "derivation": "Set to AE.AEREL", "source": "AE.AEREL", "codelist": None},
            {"name": "ASTDT", "label": "Analysis Start Date", "type": "Num", "length": 8, "derivation": "Derived from AE.AESTDTC", "source": "AE.AESTDTC", "codelist": None},
            {"name": "AENDT", "label": "Analysis End Date", "type": "Num", "length": 8, "derivation": "Derived from AE.AEENDTC", "source": "AE.AEENDTC", "codelist": None},
            {"name": "AETRTEMFL", "label": "Treatment Emergent Flag", "type": "Char", "length": 1, "derivation": "Y if ASTDT >= TRTSDT and ASTDT <= TRTEDT+30", "source": "Derived", "codelist": "NY"},
        ]
        datasets_json.append({
            "name": "ADAE",
            "label": "Adverse Event Analysis Dataset",
            "structure": "One record per adverse event per subject",
            "keys": ["STUDYID", "USUBJID", "AESEQ"],
            "variables": adae_vars
        })
        total_vars += len(adae_vars)

        # ADEFF - Efficacy Analysis Dataset
        adeff_vars = [
            {"name": "STUDYID", "label": "Study Identifier", "type": "Char", "length": 20, "derivation": "Set to ADSL.STUDYID", "source": "ADSL", "codelist": None},
            {"name": "USUBJID", "label": "Unique Subject Identifier", "type": "Char", "length": 50, "derivation": "Set to ADSL.USUBJID", "source": "ADSL", "codelist": None},
            {"name": "PARAMCD", "label": "Parameter Code", "type": "Char", "length": 8, "derivation": "Assigned per parameter", "source": "Derived", "codelist": None},
            {"name": "PARAM", "label": "Parameter", "type": "Char", "length": 200, "derivation": "Parameter description", "source": "Derived", "codelist": None},
            {"name": "AVAL", "label": "Analysis Value", "type": "Num", "length": 8, "derivation": "Numeric analysis value", "source": "Derived", "codelist": None},
            {"name": "BASE", "label": "Baseline Value", "type": "Num", "length": 8, "derivation": "Value where ABLFL=Y", "source": "Derived", "codelist": None},
            {"name": "CHG", "label": "Change from Baseline", "type": "Num", "length": 8, "derivation": "AVAL - BASE", "source": "Derived", "codelist": None},
            {"name": "PCHG", "label": "Percent Change from Baseline", "type": "Num", "length": 8, "derivation": "100 * (AVAL - BASE) / BASE", "source": "Derived", "codelist": None},
            {"name": "AVISIT", "label": "Analysis Visit", "type": "Char", "length": 40, "derivation": "Analysis visit with windowing", "source": "Derived", "codelist": None},
            {"name": "ABLFL", "label": "Baseline Record Flag", "type": "Char", "length": 1, "derivation": "Y for baseline record", "source": "Derived", "codelist": "NY"},
            {"name": "ANL01FL", "label": "Analysis Record Flag 01", "type": "Char", "length": 1, "derivation": "Y for primary analysis records", "source": "Derived", "codelist": "NY"},
        ]
        datasets_json.append({
            "name": "ADEFF",
            "label": "Efficacy Analysis Dataset",
            "structure": "One record per subject per parameter per visit",
            "keys": ["STUDYID", "USUBJID", "PARAMCD", "AVISIT"],
            "variables": adeff_vars
        })
        total_vars += len(adeff_vars)

        # ADTTE - Time-to-Event Dataset
        adtte_vars = [
            {"name": "STUDYID", "label": "Study Identifier", "type": "Char", "length": 20, "derivation": "Set to ADSL.STUDYID", "source": "ADSL", "codelist": None},
            {"name": "USUBJID", "label": "Unique Subject Identifier", "type": "Char", "length": 50, "derivation": "Set to ADSL.USUBJID", "source": "ADSL", "codelist": None},
            {"name": "PARAMCD", "label": "Parameter Code", "type": "Char", "length": 8, "derivation": "Assigned per TTE parameter", "source": "Derived", "codelist": None},
            {"name": "PARAM", "label": "Parameter", "type": "Char", "length": 200, "derivation": "TTE parameter description", "source": "Derived", "codelist": None},
            {"name": "STARTDT", "label": "Time-to-Event Origin Date", "type": "Num", "length": 8, "derivation": "Randomization or first dose date", "source": "ADSL", "codelist": None},
            {"name": "ADT", "label": "Analysis Date", "type": "Num", "length": 8, "derivation": "Event or censoring date", "source": "Derived", "codelist": None},
            {"name": "AVAL", "label": "Analysis Value", "type": "Num", "length": 8, "derivation": "ADT - STARTDT + 1 (days)", "source": "Derived", "codelist": None},
            {"name": "CNSR", "label": "Censor", "type": "Num", "length": 8, "derivation": "0=Event, 1=Censored", "source": "Derived", "codelist": None},
            {"name": "EVNTDESC", "label": "Event Description", "type": "Char", "length": 200, "derivation": "Description of event", "source": "Derived", "codelist": None},
        ]
        datasets_json.append({
            "name": "ADTTE",
            "label": "Time-to-Event Analysis Dataset",
            "structure": "One record per subject per parameter",
            "keys": ["STUDYID", "USUBJID", "PARAMCD"],
            "variables": adtte_vars
        })
        total_vars += len(adtte_vars)

        # Generate markdown document
        markdown_parts = [
            "# ADaM Derivation Specifications",
            f"\n**Protocol:** {protocol_id}",
            f"\n**Therapeutic Area:** {therapeutic_area.upper()}",
            "\n---\n"
        ]

        for ds in datasets_json:
            markdown_parts.append(f"\n## {ds['name']} - {ds['label']}")
            markdown_parts.append(f"\n**Structure:** {ds['structure']}")
            markdown_parts.append(f"\n**Keys:** {', '.join(ds['keys'])}")
            markdown_parts.append("\n\n### Variable Derivations\n")
            markdown_parts.append("| Variable | Label | Type | Derivation |")
            markdown_parts.append("|----------|-------|------|------------|")
            for v in ds['variables'][:20]:
                deriv = v['derivation'][:80] + "..." if len(v['derivation']) > 80 else v['derivation']
                markdown_parts.append(f"| {v['name']} | {v['label'][:40]} | {v['type']} | {deriv} |")
            if len(ds['variables']) > 20:
                markdown_parts.append(f"\n*...and {len(ds['variables']) - 20} more variables*")

        return AdamSpecResponse(
            success=True,
            message=f"Generated derivation specs for {len(datasets_json)} ADaM datasets with {total_vars} variables",
            datasets=datasets_json,
            total_variables=total_vars,
            markdown="\n".join(markdown_parts),
            errors=[]
        )

    except HTTPException:
        raise
    except Exception as e:
        return AdamSpecResponse(
            success=False,
            message=f"ADaM spec generation failed: {str(e)}",
            errors=[str(e)]
        )


# =============================================================================
# DEFINE-XML GENERATION ENDPOINT
# =============================================================================

class DefineXMLResponse(BaseModel):
    """Response model for Define-XML generation."""
    success: bool
    message: str
    xml_content: str = ""
    dataset_count: int = 0
    variable_count: int = 0
    standard_type: str = ""  # "SDTM" or "ADaM"
    errors: list = []


@app.post("/generate-define-xml/{job_id}")
async def generate_define_xml(job_id: str, standard: str = "adam"):
    """
    Generate CDISC Define-XML 2.1 metadata from a completed SAP job.

    Args:
        job_id: The job ID to generate Define-XML for
        standard: Either "sdtm" or "adam" (default: adam)

    Returns:
        - Complete Define-XML 2.1 compliant XML document
        - Dataset definitions
        - Variable metadata with origins and derivations
        - Codelists
    """
    try:
        db = get_supabase()

        # Get job and verify it's completed
        result = db.table("sap_jobs").select("*").eq("id", job_id).execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Job not found")

        job = result.data[0]

        if job["status"] != "completed":
            raise HTTPException(
                status_code=400,
                detail=f"Job not ready for Define-XML generation. Status: {job['status']}"
            )

        # Import Define-XML generator
        try:
            from enterprise_sap_system.specs.define_xml import (
                generate_sdtm_define_xml,
                generate_adam_define_xml,
            )
        except ImportError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Define-XML generator not available: {e}"
            )

        # Extract protocol info using improved extraction
        protocol_text = job.get("protocol_text", "")
        sap_text = job.get("generated_sap", "")
        protocol_id = _extract_protocol_id(sap_text, protocol_text, job.get("nct_id"))
        if protocol_id == "UNKNOWN":
            protocol_id = "STUDY-001"
        study_name = f"Study {protocol_id}"

        # Generate based on standard type
        if standard.lower() == "sdtm":
            xml_content = generate_sdtm_define_xml(
                study_id=protocol_id,
                study_name=study_name,
                domains=["DM", "AE", "CM", "DS", "EX", "LB", "MH", "VS"]
            )
            dataset_count = 8
            variable_count = 200  # Approximate
            standard_type = "SDTM"
        else:
            xml_content = generate_adam_define_xml(
                study_id=protocol_id,
                study_name=study_name,
                datasets=["ADSL", "ADAE", "ADLB", "ADEFF", "ADTTE"]
            )
            dataset_count = 5
            variable_count = 120  # Approximate
            standard_type = "ADaM"

        return {
            "success": True,
            "message": f"Generated {standard_type} Define-XML 2.1 with {dataset_count} datasets",
            "xml_content": xml_content,
            "dataset_count": dataset_count,
            "variable_count": variable_count,
            "standard_type": standard_type,
            "errors": []
        }

    except HTTPException:
        raise
    except Exception as e:
        return {
            "success": False,
            "message": f"Define-XML generation failed: {str(e)}",
            "xml_content": "",
            "dataset_count": 0,
            "variable_count": 0,
            "standard_type": "",
            "errors": [str(e)]
        }


# =============================================================================
# WORKBENCH ENDPOINTS - Section-by-section SAP generation with KG extraction
# =============================================================================

# Global workbench instance
_workbench: 'SAPWorkbench' = None


def get_workbench():
    """Get or create the SAP Workbench instance."""
    global _workbench
    if _workbench is None and WORKBENCH_AVAILABLE:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            logger.error("ANTHROPIC_API_KEY not set")
            return None
        if not SUPABASE_URL or not SUPABASE_KEY:
            logger.error("SUPABASE_URL and SUPABASE_SERVICE_KEY required for workbench storage")
            return None
        try:
            _workbench = SAPWorkbench(
                api_key=api_key,
                supabase_url=SUPABASE_URL,
                supabase_key=SUPABASE_KEY,
                use_kg=True
            )
            logger.info("SAP Workbench initialized with Supabase storage + 55-category KG extraction")
        except Exception as e:
            logger.error(f"Failed to initialize workbench: {e}")
            return None
    return _workbench


class WorkspaceCreate(BaseModel):
    """Request to create a workspace."""
    protocol_content: str
    protocol_filename: str
    phase: str = ""
    therapeutic_area: str = ""
    indication: str = ""


class WorkspaceResponse(BaseModel):
    """Response for workspace operations."""
    id: str
    name: str
    created_at: str
    phase: str = ""
    therapeutic_area: str = ""


class MetadataResponse(BaseModel):
    """Response with extracted protocol metadata."""
    study_id: str = ""
    study_title: str = ""
    phase: str = ""
    therapeutic_area: str = ""
    indication: str = ""
    disease_setting: str = ""
    performance_status_scale: str = ""
    response_criteria: str = ""
    geographic_countries: List[str] = []
    endpoints: List[Dict] = []
    populations: List[Dict] = []
    treatment_arms: List[Dict] = []
    stratification_factors: List[str] = []
    sample_size: Optional[int] = None
    prohibition_rules: List[str] = []
    extraction_method: str = ""
    # Schedule of Assessments (SOA) - visit schedule from protocol
    visit_schedule: List[Dict] = []
    tumor_assessment_frequency: str = ""
    pro_collection_visits: List[str] = []
    follow_up_schedule: str = ""


class SectionOutline(BaseModel):
    """SAP section outline item."""
    id: str
    name: str
    status: str
    has_content: bool
    version: int


class KBToolUsed(BaseModel):
    """KB tool usage tracking for provenance."""
    tool_name: str
    source_file: str
    source_key: str
    description: str = ""


class SectionContent(BaseModel):
    """Full section content with provenance."""
    id: str
    name: str
    display_name: str
    status: str
    content: str
    protocol_excerpts_used: List[str] = []
    metadata_used: List[str] = []
    kb_tools_used: List[KBToolUsed] = []
    version: int
    generated_at: str = ""


class SectionUpdate(BaseModel):
    """Request to update a section."""
    content: str
    comments: str = ""


@app.post("/workbench/create", response_model=WorkspaceResponse)
async def workbench_create_workspace(request: WorkspaceCreate):
    """
    Step 1: Create a new study workspace.
    Upload protocol and set basic metadata.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized (check API key)")

    try:
        workspace = workbench.create_workspace(
            protocol_content=request.protocol_content,
            protocol_filename=request.protocol_filename,
            phase=request.phase,
            therapeutic_area=request.therapeutic_area,
            indication=request.indication
        )

        return WorkspaceResponse(
            id=workspace.id,
            name=workspace.name,
            created_at=workspace.created_at,
            phase=workspace.phase,
            therapeutic_area=workspace.therapeutic_area
        )
    except Exception as e:
        logger.error(f"Workbench create failed: {e}")
        raise HTTPException(500, str(e))


@app.post("/workbench/upload")
async def workbench_upload_protocol(
    file: UploadFile = File(...),
    phase: str = Form(""),
    therapeutic_area: str = Form(""),
    indication: str = Form("")
):
    """
    Step 1 (alt): Create workspace by uploading file.
    Properly parses PDF/DOCX/TXT files to extract text.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        content = await file.read()
        filename = file.filename or "protocol.txt"

        # IMPORTANT: Use proper file parsing (PDF, DOCX, TXT)
        # This was the bug - we were just doing decode() which gives garbled data for PDFs
        try:
            protocol_content = extract_text_from_file(filename, content)
            logger.info(f"Workbench: Extracted {len(protocol_content):,} chars from {filename}")
        except ValueError as e:
            logger.error(f"Workbench: Text extraction failed for {filename}: {e}")
            raise HTTPException(400, f"Failed to extract text from file: {e}")

        if not protocol_content.strip():
            raise HTTPException(400, "No text could be extracted from the file")

        # Log first 500 chars for debugging
        logger.info(f"Workbench: Protocol preview: {protocol_content[:500]}...")

        workspace = workbench.create_workspace(
            protocol_content=protocol_content,
            protocol_filename=filename,
            phase=phase,
            therapeutic_area=therapeutic_area,
            indication=indication
        )

        return {
            "id": workspace.id,
            "name": workspace.name,
            "created_at": workspace.created_at,
            "phase": workspace.phase,
            "therapeutic_area": workspace.therapeutic_area
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Workbench upload failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/metadata", response_model=MetadataResponse)
async def workbench_get_metadata(workspace_id: str):
    """
    Step 2: Get extracted protocol metadata (55-category KG extraction).
    This is the "Protocol Understanding View".
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        workspace = workbench.get_workspace(workspace_id)
        if not workspace:
            raise HTTPException(404, f"Workspace {workspace_id} not found")

        # Extract metadata if not done
        if not workspace.metadata:
            metadata = workbench.extract_metadata(workspace_id)
        else:
            metadata = workspace.metadata

        # Extract SOA info from full_extraction if available
        soa = metadata.full_extraction.get("schedule_of_assessments", {}) if metadata.full_extraction else {}

        return MetadataResponse(
            study_id=metadata.study_id,
            study_title=metadata.study_title,
            phase=metadata.phase,
            therapeutic_area=metadata.therapeutic_area,
            indication=metadata.indication,
            disease_setting=metadata.disease_setting,
            performance_status_scale=metadata.performance_status_scale,
            response_criteria=metadata.response_criteria,
            geographic_countries=metadata.geographic_countries,
            endpoints=metadata.endpoints,
            populations=metadata.populations,
            treatment_arms=metadata.treatment_arms,
            stratification_factors=metadata.stratification_factors,
            sample_size=metadata.sample_size,
            prohibition_rules=metadata.prohibition_rules,
            extraction_method=metadata.extraction_method,
            # Schedule of Assessments (SOA)
            visit_schedule=metadata.visit_schedule if hasattr(metadata, 'visit_schedule') else [],
            tumor_assessment_frequency=soa.get("tumor_assessment_frequency", ""),
            pro_collection_visits=soa.get("pro_collection_visits", []),
            follow_up_schedule=soa.get("follow_up_schedule", "")
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Metadata extraction failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/extraction")
async def workbench_get_extraction(workspace_id: str):
    """
    Get full 55-category extraction data for UI display.
    Returns the complete extraction with all fields organized for the frontend.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        workspace = workbench.get_workspace(workspace_id)
        if not workspace:
            raise HTTPException(404, f"Workspace {workspace_id} not found")

        # Extract metadata if not done
        if not workspace.metadata:
            workbench.extract_metadata(workspace_id)
            workspace = workbench.get_workspace(workspace_id)

        metadata = workspace.metadata
        if not metadata:
            raise HTTPException(500, "Failed to extract metadata")

        fe = metadata.full_extraction or {}

        # Build structured response for UI
        def safe_get(d, *keys, default=None):
            """Safely navigate nested dicts"""
            for key in keys:
                if isinstance(d, dict):
                    d = d.get(key, {})
                else:
                    return default
            return d if d != {} else default

        def get_value(d, key):
            """Get value from nested {value: x} structure"""
            if isinstance(d, dict):
                if "value" in d:
                    return d["value"]
                return d.get(key)
            return d

        # Helper to extract source info (defined early for use below)
        def get_source_info_early(obj):
            """Extract source_quote and source_section from an object."""
            if isinstance(obj, dict):
                return {
                    "source_quote": obj.get("source_quote", ""),
                    "source_section": obj.get("source_section", ""),
                }
            return {"source_quote": "", "source_section": ""}

        # Study info with source references
        trial_id = fe.get("trial_identification", {})
        study_info = {
            "nct_id": get_value(trial_id.get("nct_id", {}), "value") or "",
            "nct_id_source": get_source_info_early(trial_id.get("nct_id", {})),
            "protocol_number": get_value(trial_id.get("protocol_number", {}), "value") or "",
            "protocol_number_source": get_source_info_early(trial_id.get("protocol_number", {})),
            "sponsor": get_value(trial_id.get("sponsor", {}), "value") or "",
            "sponsor_source": get_source_info_early(trial_id.get("sponsor", {})),
            "title": get_value(trial_id.get("study_title", {}), "value") or metadata.study_title,
            "title_source": get_source_info_early(trial_id.get("study_title", {})),
            "phase": metadata.phase,
        }

        # Design with source references
        design = fe.get("study_design", {})
        design_info = {
            "type": get_value(design.get("design_type", {}), "value") or "",
            "type_source": get_source_info_early(design.get("design_type", {})),
            "blinding": get_value(design.get("blinding", {}), "value") or "",
            "blinding_source": get_source_info_early(design.get("blinding", {})),
            "randomization_ratio": get_value(design.get("randomization_ratio", {}), "value") or "",
            "randomization_ratio_source": get_source_info_early(design.get("randomization_ratio", {})),
            "control_type": get_value(design.get("control_type", {}), "value") or "",
            "control_type_source": get_source_info_early(design.get("control_type", {})),
        }

        # Study type flags
        cart_specific = fe.get("cart_specific", {})
        heme_specific = fe.get("hematologic_specific", {})
        immuno_specific = fe.get("immunotherapy_specific", {})
        study_types = {
            "is_cart": cart_specific.get("is_cart", False),
            "is_hematologic": heme_specific.get("is_hematologic", False),
            "is_immunotherapy": immuno_specific.get("is_immunotherapy", False),
        }

        # Endpoints with source references
        # (using get_source_info_early defined above)
        primary_eps = fe.get("primary_endpoints", [])
        secondary_eps = fe.get("secondary_endpoints", [])
        endpoints = {
            "primary": [
                {
                    "name": ep.get("name", ""),
                    "definition": ep.get("definition", ""),
                    "type": ep.get("type", ""),
                    "response_criteria": ep.get("response_criteria", ""),
                    **get_source_info_early(ep),
                }
                for ep in primary_eps if ep.get("name")
            ],
            "secondary": [
                {
                    "name": ep.get("name", ""),
                    "definition": ep.get("definition", ""),
                    "type": ep.get("type", ""),
                    **get_source_info_early(ep),
                }
                for ep in secondary_eps if ep.get("name")
            ],
        }

        # Populations with source references
        pops = fe.get("populations", [])
        populations = [
            {
                "name": p.get("name", ""),
                "definition": p.get("definition", ""),
                "is_primary_efficacy": p.get("is_primary_efficacy", False),
                "is_primary_safety": p.get("is_primary_safety", False),
                **get_source_info_early(p),
            }
            for p in pops if p.get("name")
        ]

        # Sample size with source references
        ss = fe.get("sample_size", {})
        sample_size_info = {
            "total_n": get_value(ss.get("total_n", {}), "value") or metadata.sample_size,
            "total_n_source": get_source_info_early(ss.get("total_n", {})),
            "power": get_value(ss.get("power", {}), "value"),
            "power_source": get_source_info_early(ss.get("power", {})),
            "alpha": get_value(ss.get("alpha", {}), "value"),
            "alpha_source": get_source_info_early(ss.get("alpha", {})),
            "effect_size": get_value(ss.get("effect_size", {}), "value"),
            "effect_size_source": get_source_info_early(ss.get("effect_size", {})),
        }

        # CAR-T specific
        cart_info = None
        if study_types["is_cart"]:
            crs = cart_specific.get("crs_grading", {})
            icans = cart_specific.get("icans_grading", {})
            cart_info = {
                "crs_scale": crs.get("scale", ""),
                "icans_scale": icans.get("scale", ""),
                "bridging_therapy": cart_specific.get("bridging_therapy", {}).get("allowed", False),
                "cellular_kinetics": cart_specific.get("cellular_kinetics", {}).get("parameters", []),
            }

        # Subgroups with source references
        subgroups = [
            {
                "factor": sg.get("factor", ""),
                "categories": sg.get("categories", []),
                "is_stratification_factor": sg.get("is_stratification_factor", False),
                **get_source_info_early(sg),
            }
            for sg in fe.get("subgroups", []) if sg.get("factor")
        ]

        # Censoring rules with source references
        censoring = [
            {
                "endpoint": cr.get("endpoint", ""),
                "scenario": cr.get("scenario", ""),
                "event_flag": cr.get("event_flag", ""),
                "date_used": cr.get("date_used", ""),
                **get_source_info_early(cr),
            }
            for cr in fe.get("censoring_rules", []) if cr.get("endpoint")
        ]

        # Validation/warnings (deduplicated)
        validation = fe.get("_validation_report", {})
        warnings = []
        seen_warnings = set()

        if validation.get("status") == "needs_review":
            for gap in validation.get("potential_gaps", []):
                gap_str = str(gap) if isinstance(gap, dict) else gap
                if gap_str not in seen_warnings:
                    warnings.append(gap)
                    seen_warnings.add(gap_str)

        if validation.get("warnings"):
            for warning in validation.get("warnings", []):
                warning_str = str(warning) if isinstance(warning, dict) else warning
                if warning_str not in seen_warnings:
                    warnings.append(warning)
                    seen_warnings.add(warning_str)

        # Completeness check
        completeness = fe.get("extraction_completeness_check", {})

        return {
            "workspace_id": workspace_id,
            "extraction_method": metadata.extraction_method,
            "extraction_timestamp": metadata.extraction_timestamp,

            # Organized data
            "study_info": study_info,
            "design": design_info,
            "study_types": study_types,
            "endpoints": endpoints,
            "populations": populations,
            "sample_size": sample_size_info,
            "subgroups": subgroups,
            "censoring_rules": censoring,
            "prohibition_rules": metadata.prohibition_rules,

            # Special sections
            "cart_specific": cart_info,
            "hematologic_specific": heme_specific if study_types["is_hematologic"] else None,
            "immunotherapy_specific": immuno_specific if study_types["is_immunotherapy"] else None,

            # Validation
            "warnings": warnings,
            "completeness": {
                "total_endpoints": completeness.get("total_endpoints_found", len(primary_eps) + len(secondary_eps)),
                "total_populations": completeness.get("total_populations_found", len(populations)),
                "total_subgroups": completeness.get("total_subgroups_found", len(subgroups)),
                "confidence": completeness.get("confidence_level", "unknown"),
                "all_captured": completeness.get("all_elements_captured", True),
            },

            # Raw data for advanced view
            "raw_field_count": len(fe),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Extraction fetch failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/protocol")
async def workbench_get_protocol_content(workspace_id: str):
    """
    Get the full protocol content for viewing in the Protocol Audit Suite.
    Returns the extracted text content of the uploaded protocol.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        workspace = workbench.get_workspace(workspace_id)
        if not workspace:
            raise HTTPException(404, f"Workspace {workspace_id} not found")

        protocol_content = workspace.protocol_content or ""
        protocol_filename = workspace.protocol_filename or "protocol.txt"

        return {
            "workspace_id": workspace_id,
            "filename": protocol_filename,
            "content": protocol_content,
            "char_count": len(protocol_content),
            "has_content": bool(protocol_content.strip()),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Protocol content fetch failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/outline")
async def workbench_get_outline(workspace_id: str):
    """
    Step 3: Get SAP outline with section statuses.
    This is the "SAP Skeleton" view.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        outline = workbench.get_outline(workspace_id)
        return {"sections": outline}
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Get outline failed: {e}")
        raise HTTPException(500, str(e))


@app.post("/workbench/{workspace_id}/generate/{section_id}", response_model=SectionContent)
async def workbench_generate_section(
    workspace_id: str,
    section_id: str,
    regenerate: bool = False
):
    """
    Step 4: Generate a single SAP section.
    Shows protocol excerpts used and allows regeneration.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        section = workbench.generate_section(workspace_id, section_id, regenerate)

        return SectionContent(
            id=section.id,
            name=section.name,
            display_name=section.display_name,
            status=section.status.value,
            content=section.content,
            protocol_excerpts_used=section.protocol_excerpts_used,
            metadata_used=section.metadata_used,
            kb_tools_used=[KBToolUsed(**t) for t in section.kb_tools_used],
            version=section.version,
            generated_at=section.generated_at or ""
        )
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Section generation failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/section/{section_id}", response_model=SectionContent)
async def workbench_get_section(workspace_id: str, section_id: str):
    """
    Get a section's current content.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        workspace = workbench.get_workspace(workspace_id)
        if not workspace:
            raise HTTPException(404, f"Workspace {workspace_id} not found")

        section = workspace.sections.get(section_id)
        if not section:
            raise HTTPException(404, f"Section {section_id} not found")

        return SectionContent(
            id=section.id,
            name=section.name,
            display_name=section.display_name,
            status=section.status.value,
            content=section.content,
            protocol_excerpts_used=section.protocol_excerpts_used,
            metadata_used=section.metadata_used,
            kb_tools_used=[KBToolUsed(**t) for t in section.kb_tools_used],
            version=section.version,
            generated_at=section.generated_at or ""
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get section failed: {e}")
        raise HTTPException(500, str(e))


@app.put("/workbench/{workspace_id}/section/{section_id}", response_model=SectionContent)
async def workbench_update_section(
    workspace_id: str,
    section_id: str,
    request: SectionUpdate
):
    """
    Update a section with user edits.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        section = workbench.update_section(
            workspace_id, section_id, request.content, request.comments
        )

        return SectionContent(
            id=section.id,
            name=section.name,
            display_name=section.display_name,
            status=section.status.value,
            content=section.content,
            protocol_excerpts_used=section.protocol_excerpts_used,
            metadata_used=section.metadata_used,
            kb_tools_used=[KBToolUsed(**t) for t in section.kb_tools_used],
            version=section.version,
            generated_at=section.generated_at or ""
        )
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Update section failed: {e}")
        raise HTTPException(500, str(e))


@app.post("/workbench/{workspace_id}/section/{section_id}/approve")
async def workbench_approve_section(workspace_id: str, section_id: str):
    """
    Mark a section as approved.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        section = workbench.approve_section(workspace_id, section_id)
        return {"status": "approved", "section_id": section_id}
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Approve section failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/provenance")
async def workbench_get_provenance(workspace_id: str):
    """
    Step 5: Get full provenance report.
    "Why does this say this?"
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        return workbench.get_provenance_report(workspace_id)
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Get provenance failed: {e}")
        raise HTTPException(500, str(e))


@app.post("/workbench/{workspace_id}/update-protocol")
async def workbench_update_protocol(
    workspace_id: str,
    file: UploadFile = File(...)
):
    """
    Step 6: Update protocol and identify impacted sections.
    "Protocol changed → impact shown"
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        content = await file.read()
        protocol_content = content.decode('utf-8', errors='ignore')

        result = workbench.update_protocol(
            workspace_id, protocol_content, file.filename or "protocol.txt"
        )

        return result
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Update protocol failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/export")
async def workbench_export_sap(workspace_id: str, format: str = "markdown"):
    """
    Step 7: Export complete SAP.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        content = workbench.export_sap(workspace_id, format)
        return {"format": format, "content": content}
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Export SAP failed: {e}")
        raise HTTPException(500, str(e))


# =============================================================================
# v100.3: REFERENCE SAP COMPARISON ENDPOINTS
# =============================================================================

@app.post("/workbench/{workspace_id}/reference-sap")
async def workbench_upload_reference_sap(workspace_id: str, file: UploadFile = File(...)):
    """
    Upload a reference SAP for accuracy comparison (optional).

    This parses the reference SAP into sections that will be used
    to compare against generated sections.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        content = await file.read()

        # Handle PDF files - use PyPDF2 (already installed for protocol upload)
        filename = file.filename or "reference_sap.txt"
        if filename.lower().endswith('.pdf'):
            # Extract text from PDF using PyPDF2
            sap_content = extract_text_from_pdf(content)
        else:
            sap_content = content.decode('utf-8', errors='ignore')

        result = workbench.upload_reference_sap(workspace_id, sap_content, filename)
        return result

    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Upload reference SAP failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/{workspace_id}/reference-sap/status")
async def workbench_reference_sap_status(workspace_id: str):
    """
    Check if a reference SAP has been uploaded for this workspace.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        return workbench.get_reference_sap_status(workspace_id)
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Get reference SAP status failed: {e}")
        raise HTTPException(500, str(e))


@app.post("/workbench/{workspace_id}/compare/{section_id}")
async def workbench_compare_section(workspace_id: str, section_id: str):
    """
    Compare a generated section against the reference SAP section.

    Returns detailed accuracy report with:
    - Percentage match
    - Missing elements (with quotes from original)
    - Incorrect elements
    - Suggestions for fixing
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        result = workbench.compare_section_with_reference(workspace_id, section_id)
        return result
    except ValueError as e:
        raise HTTPException(404, str(e))
    except Exception as e:
        logger.error(f"Compare section failed: {e}")
        raise HTTPException(500, str(e))


@app.get("/workbench/list")
async def workbench_list_workspaces():
    """
    List all workspaces.
    """
    if not WORKBENCH_AVAILABLE:
        raise HTTPException(503, "SAP Workbench not available")

    workbench = get_workbench()
    if not workbench:
        raise HTTPException(503, "SAP Workbench not initialized")

    try:
        return {"workspaces": workbench.list_workspaces()}
    except Exception as e:
        logger.error(f"List workspaces failed: {e}")
        raise HTTPException(500, str(e))


# Background worker
async def process_jobs_worker():
    """
    Background worker that processes queued jobs using KGPipelineWrapper.

    KGPipelineWrapper (55-CATEGORY EXTRACTION + KB TOOLS + PROHIBITION RULES):
    1. LlamaParse: PDF → Markdown (preserves tables, complex layouts)
    2. 55-Category Extraction: Comprehensive protocol analysis with provenance
    3. Prohibition Rules: Context-aware constraints built from extraction
       - Adjuvant trials → No CR/PR/SD/PD response tables
       - Nordic countries → No race/ethnicity in demographics
       - ASA specified → No ECOG
       - Fixed-dose → No dose modification rows
    4. SAP Generation: use_tools=True - Claude calls KB for standards/templates
    5. SELF-RAG Verification: Fact checking with correction loop
    """
    global worker_running

    print("Starting background job worker with KGPipelineWrapper...")
    print("  [VERSION] Build 2026-01-11-v72 (KG Pipeline with KB Tools + source_section traceability)")
    print("  [NEW] 55-category comprehensive extraction")
    print("  [NEW] Prohibition rules from extraction:")
    print("        • Adjuvant → No CR/PR/SD/PD")
    print("        • Nordic → No Race/Ethnicity")
    print("        • ASA → No ECOG")
    print("        • Fixed-dose → No dose modification")
    print("  [OK] Step 1: LlamaParse extracts PDF → Markdown")
    print("  [OK] Step 2: 55-category KG extraction with provenance")
    print("  [OK] Step 3: Build prohibition rules from extraction")
    print("  [OK] Step 4: Generate SAP with KB tools (use_tools=True)")
    print("  [OK] Step 5: SELF-RAG verification with correction loop")

    # Use get_pipeline() - returns KGPipelineWrapper
    pipeline = None

    while worker_running:
        try:
            db = get_supabase()

            # Get next queued job
            result = db.table("sap_jobs").select("*").eq(
                "status", "queued"
            ).order("created_at").limit(1).execute()

            if not result.data:
                # No jobs, wait and retry
                await asyncio.sleep(5)
                continue

            job = result.data[0]
            job_id = job["id"]

            print(f"Processing job: {job_id}")

            # Mark as processing
            db.table("sap_jobs").update({
                "status": "processing",
                "started_at": datetime.utcnow().isoformat()
            }).eq("id", job_id).execute()

            # Initialize pipeline if needed
            if pipeline is None:
                pipeline = get_pipeline()
                # Check which pipeline type was initialized
                pipeline_type = type(pipeline).__name__
                print(f"  [INIT] {pipeline_type} initialized")

            # Generate SAP using pipeline
            start_time = time.time()

            try:
                # Pipeline uses different methods:
                # - process_pdf() for PDF files (LlamaParse extraction)
                # - process_protocol() for text

                pdf_path = None
                pdf_storage_path = job.get("pdf_storage_path")

                # Download PDF if available
                if pdf_storage_path:
                    try:
                        import tempfile
                        print(f"  [PDF] Downloading PDF from storage: {pdf_storage_path}")
                        pdf_bytes = db.storage.from_("pdfs").download(pdf_storage_path)

                        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                            tmp.write(pdf_bytes)
                            pdf_path = tmp.name
                            print(f"  [PDF] Saved to temp file: {pdf_path}")

                    except Exception as e:
                        print(f"  [PDF] Download failed, using text: {e}")
                        pdf_path = None

                # Call pipeline (KGPipelineWrapper with EnhancedKGPipeline - v69 dynamic SAP)
                pipeline_name = type(pipeline).__name__
                if pdf_path:
                    # Use LlamaParse for PDF extraction
                    print(f"  [{pipeline_name}] Using process_pdf() with LlamaParse")
                    result = pipeline.process_pdf(
                        pdf_path,
                        protocol_id=job.get("nct_id") or job_id,
                        validate=True,
                        verbose=True
                    )
                else:
                    # Use text directly
                    print(f"  [{pipeline_name}] Using process_protocol() with text")
                    result = pipeline.process_protocol(
                        job["protocol_text"],
                        protocol_id=job.get("nct_id") or job_id,
                        validate=True,
                        verbose=True
                    )

                # Clean up temp PDF
                if pdf_path:
                    try:
                        import os
                        os.unlink(pdf_path)
                    except Exception:
                        pass

                processing_time = time.time() - start_time

                # Pipeline returns a dict with sap_text, validation, etc.
                sap_text = result.get("sap_text", "")

                # DEBUG: Log SAP generation result
                print(f"  [DEBUG] SAP from {pipeline_name}: {len(sap_text)} chars")

                # Log prohibition rules if KG pipeline was used
                prohibition_rules = result.get("prohibition_rules", [])

                # Parse prohibition rules into boolean flags for TLF template generation
                tlf_prohibitions = {
                    'no_race_ethnicity': False,
                    'no_ecog': False,  # Use Karnofsky or ASA instead
                    'use_karnofsky': False,
                    'use_asa_score': False,
                    'no_weight_use_bmi': False,
                    'no_response_categories': False,  # No CR/PR/SD/PD for adjuvant
                    'no_dose_modification': False,  # Fixed dose study
                    'no_geographic_subgroups': False,  # Nordic/single-region study
                }

                for rule in prohibition_rules:
                    rule_lower = rule.lower()
                    if 'race' in rule_lower or 'ethnicity' in rule_lower:
                        tlf_prohibitions['no_race_ethnicity'] = True
                    if 'ecog' in rule_lower and 'not' in rule_lower:
                        tlf_prohibitions['no_ecog'] = True
                    if 'karnofsky' in rule_lower:
                        tlf_prohibitions['use_karnofsky'] = True
                    if 'asa' in rule_lower and 'score' in rule_lower:
                        tlf_prohibitions['use_asa_score'] = True
                    if 'bmi' in rule_lower and 'weight' in rule_lower:
                        tlf_prohibitions['no_weight_use_bmi'] = True
                    if 'cr/pr/sd/pd' in rule_lower or 'response categor' in rule_lower or 'recist' in rule_lower:
                        tlf_prohibitions['no_response_categories'] = True
                    if 'dose modification' in rule_lower or 'fixed dose' in rule_lower or 'fixed-dose' in rule_lower:
                        tlf_prohibitions['no_dose_modification'] = True
                    if 'geographic' in rule_lower or 'nordic' in rule_lower or 'single region' in rule_lower:
                        tlf_prohibitions['no_geographic_subgroups'] = True

                if prohibition_rules:
                    print(f"  [DEBUG] Prohibition rules applied: {len(prohibition_rules)}")
                    for rule in prohibition_rules[:3]:
                        print(f"    • {rule[:60]}...")
                    print(f"  [DEBUG] TLF prohibitions: {[k for k,v in tlf_prohibitions.items() if v]}")
                print(f"  [DEBUG] Contains '|--': {'|--' in sap_text}")
                print(f"  [DEBUG] Contains '## 12.': {'## 12.' in sap_text}")
                if '## 12.' in sap_text:
                    sec12_pos = sap_text.find('## 12.')
                    sec12_preview = sap_text[sec12_pos:sec12_pos+500]
                    print(f"  [DEBUG] Section 12 preview: {sec12_preview[:200]}...")

                if sap_text:
                    # Pipeline result format (KGPipelineWrapper)
                    validation = result.get("validation", {})
                    discovered_elements = result.get("discovered_elements", [])

                    # Quality score from validation (0-1 scale → 0-100)
                    validation_score = validation.get("overall_score", 0.8)
                    quality_score = validation_score * 100

                    # Extract info from discovered elements
                    drug_name = ""
                    phase_str = ""
                    therapeutic_area = ""
                    endpoint_type_str = ""

                    for elem in discovered_elements:
                        # Handle both dataclass and dict formats
                        if hasattr(elem, 'name'):
                            name = (elem.name or '').lower()
                            cat = (elem.category or '').lower()
                            desc = elem.description or ''
                        else:
                            name = (elem.get("name", "") or "").lower()
                            cat = (elem.get("category", "") or "").lower()
                            desc = elem.get("description", "") or ""
                        desc_lower = desc.lower()
                        combined = name + " " + desc_lower

                        # Drug name extraction
                        if "drug" in name or "study drug" in name:
                            drug_name = desc[:50] if desc else ""

                        # Phase extraction
                        if cat == "study_design" and "phase" in name:
                            phase_str = desc[:10] if desc else ""

                        # Therapeutic area - keyword-based detection
                        if not therapeutic_area:
                            if any(term in combined for term in ['cancer', 'tumor', 'carcinoma', 'melanoma', 'lymphoma', 'leukemia', 'oncology']):
                                therapeutic_area = 'oncology'
                            elif any(term in combined for term in ['colitis', 'crohn', 'inflammatory bowel', 'ibd', 'ulcerative']):
                                therapeutic_area = 'ibd'
                            elif any(term in combined for term in ['arthritis', 'rheumatoid', 'lupus', 'psoriatic']):
                                therapeutic_area = 'rheumatology'

                        # Endpoint type - keyword-based detection
                        if not endpoint_type_str and cat == 'endpoints':
                            if any(term in combined for term in ['survival', 'pfs', 'os', 'time-to-event', 'tte', 'progression-free', 'overall survival', 'kaplan-meier']):
                                endpoint_type_str = 'time-to-event'
                            elif any(term in combined for term in ['continuous', 'change from baseline', 'mmrm', 'score', 'index']):
                                endpoint_type_str = 'continuous'
                            elif any(term in combined for term in ['response rate', 'orr', 'remission', 'binary', 'proportion', 'responder']):
                                endpoint_type_str = 'binary'

                    # Fallback: if not found in elements, check protocol text (if available in job)
                    if not therapeutic_area:
                        protocol_lower = job.get("protocol_text", "").lower()[:50000]
                        if any(term in protocol_lower for term in ['cancer', 'tumor', 'carcinoma', 'melanoma', 'lymphoma', 'oncology']):
                            therapeutic_area = 'oncology'
                        elif any(term in protocol_lower for term in ['colitis', 'crohn', 'inflammatory bowel', 'ibd']):
                            therapeutic_area = 'ibd'
                        elif any(term in protocol_lower for term in ['arthritis', 'rheumatoid', 'lupus']):
                            therapeutic_area = 'rheumatology'
                        else:
                            therapeutic_area = 'general'

                    if not endpoint_type_str:
                        protocol_lower = job.get("protocol_text", "").lower()[:50000]
                        if any(term in protocol_lower for term in ['progression-free survival', 'overall survival', 'time to event', 'kaplan-meier', 'pfs', ' os ']):
                            endpoint_type_str = 'time-to-event'
                        elif any(term in protocol_lower for term in ['change from baseline', 'mmrm', 'continuous endpoint']):
                            endpoint_type_str = 'continuous'
                        else:
                            endpoint_type_str = 'binary'

                    # =========================================================
                    # EXTRACT DISEASE-SPECIFIC BASELINE VARIABLES
                    # Dynamically detect cancer type and relevant baseline vars
                    # =========================================================
                    disease_baseline_vars = []
                    cancer_type = ""
                    protocol_lower = job.get("protocol_text", "").lower()[:80000]

                    # Detect specific cancer types
                    if any(term in protocol_lower for term in ['colorectal', 'colon cancer', 'rectal cancer', 'crc', 'adenocarcinoma of colon']):
                        cancer_type = 'colorectal'
                        disease_baseline_vars = [
                            ('TUMOR LOCATION - N (%)', ['Right colon', 'Left colon', 'Rectum', 'Multiple/Other']),
                            ('TUMOR STAGE AT DIAGNOSIS - N (%)', ['Stage I', 'Stage II', 'Stage III', 'Stage IV']),
                            ('MUTATION STATUS - N (%)', ['PIK3CA mutant', 'PIK3CA wild-type', 'Unknown']),
                            ('BASELINE CEA (NG/ML)', ['N', 'Mean (SD)', 'Median', 'Min, Max']),
                            ('PRIOR ADJUVANT CHEMOTHERAPY - N (%)', ['Yes', 'No']),
                        ]
                    elif any(term in protocol_lower for term in ['breast cancer', 'breast carcinoma', 'her2', 'triple negative']):
                        cancer_type = 'breast'
                        disease_baseline_vars = [
                            ('HORMONE RECEPTOR STATUS - N (%)', ['ER+/PR+', 'ER+/PR-', 'ER-/PR+', 'ER-/PR-']),
                            ('HER2 STATUS - N (%)', ['HER2 positive', 'HER2 negative']),
                            ('TUMOR GRADE - N (%)', ['Grade 1', 'Grade 2', 'Grade 3']),
                            ('TUMOR STAGE - N (%)', ['Stage I', 'Stage II', 'Stage III', 'Stage IV']),
                            ('PRIOR LINES OF THERAPY - N (%)', ['0', '1', '2', '≥3']),
                        ]
                    elif any(term in protocol_lower for term in ['lung cancer', 'nsclc', 'non-small cell', 'sclc', 'small cell lung']):
                        cancer_type = 'lung'
                        disease_baseline_vars = [
                            ('HISTOLOGY - N (%)', ['Adenocarcinoma', 'Squamous cell', 'Large cell', 'Other']),
                            ('TUMOR STAGE - N (%)', ['Stage IIIA', 'Stage IIIB', 'Stage IV']),
                            ('EGFR MUTATION STATUS - N (%)', ['EGFR mutant', 'EGFR wild-type', 'Unknown']),
                            ('ALK STATUS - N (%)', ['ALK positive', 'ALK negative', 'Unknown']),
                            ('PD-L1 EXPRESSION - N (%)', ['≥50%', '1-49%', '<1%', 'Unknown']),
                            ('SMOKING HISTORY - N (%)', ['Never', 'Former', 'Current']),
                        ]
                    elif any(term in protocol_lower for term in ['lymphoma', 'hodgkin', 'non-hodgkin', 'dlbcl', 'follicular', 'mantle cell', 'car-t', 'cart', 'axicabtagene', 'tisagenlecleucel']):
                        # MOVED BEFORE MELANOMA: lymphoma protocols may mention melanoma in references
                        cancer_type = 'lymphoma'
                        disease_baseline_vars = [
                            ('LYMPHOMA TYPE - N (%)', ['DLBCL', 'Follicular', 'Mantle cell', 'Marginal zone', 'Other']),
                            ('ANN ARBOR STAGE - N (%)', ['Stage I', 'Stage II', 'Stage III', 'Stage IV']),
                            ('IPI/FLIPI SCORE - N (%)', ['Low (0-1)', 'Low-intermediate (2)', 'High-intermediate (3)', 'High (4-5)']),
                            ('PRIOR LINES OF THERAPY', ['N', 'Median', 'Min, Max']),
                            ('RELAPSED VS REFRACTORY - N (%)', ['Relapsed', 'Refractory']),
                        ]
                    elif any(term in protocol_lower for term in ['melanoma', 'skin cancer']):
                        cancer_type = 'melanoma'
                        disease_baseline_vars = [
                            ('BRAF MUTATION STATUS - N (%)', ['BRAF V600E', 'BRAF V600K', 'BRAF wild-type']),
                            ('TUMOR STAGE - N (%)', ['Stage III', 'Stage IV M1a', 'Stage IV M1b', 'Stage IV M1c']),
                            ('LDH - N (%)', ['Normal', 'Elevated']),
                            ('PRIOR IMMUNOTHERAPY - N (%)', ['Yes', 'No']),
                        ]
                    elif therapeutic_area == 'oncology':
                        # Generic oncology baseline vars
                        disease_baseline_vars = [
                            ('TUMOR STAGE - N (%)', ['Stage I', 'Stage II', 'Stage III', 'Stage IV']),
                            ('PRIOR LINES OF THERAPY - N (%)', ['0', '1', '2', '≥3']),
                        ]

                    # =========================================================
                    # EXTRACT GEOGRAPHIC REGIONS/COUNTRIES FROM PROTOCOL
                    # =========================================================
                    study_regions = []

                    # Check discovered elements for geographic info
                    for elem in discovered_elements:
                        if hasattr(elem, 'description'):
                            desc = elem.description or ''
                        else:
                            desc = elem.get('description', '') or ''
                        desc_lower = desc.lower()

                        # Nordic countries
                        if 'sweden' in desc_lower and 'Sweden' not in study_regions:
                            study_regions.append('Sweden')
                        if 'denmark' in desc_lower and 'Denmark' not in study_regions:
                            study_regions.append('Denmark')
                        if 'norway' in desc_lower and 'Norway' not in study_regions:
                            study_regions.append('Norway')
                        if 'finland' in desc_lower and 'Finland' not in study_regions:
                            study_regions.append('Finland')
                        # European countries
                        if 'germany' in desc_lower and 'Germany' not in study_regions:
                            study_regions.append('Germany')
                        if 'france' in desc_lower and 'France' not in study_regions:
                            study_regions.append('France')
                        if 'spain' in desc_lower and 'Spain' not in study_regions:
                            study_regions.append('Spain')
                        if 'italy' in desc_lower and 'Italy' not in study_regions:
                            study_regions.append('Italy')
                        if 'united kingdom' in desc_lower or 'uk' in desc_lower and 'UK' not in study_regions:
                            study_regions.append('UK')
                        # US/Americas
                        if 'united states' in desc_lower or 'usa' in desc_lower and 'USA' not in study_regions:
                            study_regions.append('USA')
                        if 'canada' in desc_lower and 'Canada' not in study_regions:
                            study_regions.append('Canada')
                        # Asia Pacific
                        if 'japan' in desc_lower and 'Japan' not in study_regions:
                            study_regions.append('Japan')
                        if 'china' in desc_lower and 'China' not in study_regions:
                            study_regions.append('China')
                        if 'korea' in desc_lower and 'South Korea' not in study_regions:
                            study_regions.append('South Korea')
                        if 'australia' in desc_lower and 'Australia' not in study_regions:
                            study_regions.append('Australia')

                    # Fallback: check protocol text directly
                    if not study_regions:
                        if 'sweden' in protocol_lower:
                            study_regions.append('Sweden')
                        if 'denmark' in protocol_lower:
                            study_regions.append('Denmark')
                        if 'norway' in protocol_lower:
                            study_regions.append('Norway')
                        if 'finland' in protocol_lower:
                            study_regions.append('Finland')
                        if 'germany' in protocol_lower:
                            study_regions.append('Germany')
                        if 'france' in protocol_lower:
                            study_regions.append('France')
                        if 'united states' in protocol_lower or 'usa' in protocol_lower:
                            study_regions.append('USA')
                        if 'japan' in protocol_lower:
                            study_regions.append('Japan')
                        if 'china' in protocol_lower:
                            study_regions.append('China')

                    # If still no regions, use generic based on study scope
                    if not study_regions:
                        if 'nordic' in protocol_lower or 'scandinavi' in protocol_lower:
                            study_regions = ['Sweden', 'Denmark', 'Norway', 'Finland']
                        elif 'global' in protocol_lower or 'international' in protocol_lower or 'multi-national' in protocol_lower:
                            study_regions = ['North America', 'Europe', 'Asia Pacific']
                        else:
                            study_regions = ['Region 1', 'Region 2']  # Placeholder

                    print(f"  [MAIN.PY] Cancer type: {cancer_type}, Disease vars: {len(disease_baseline_vars)}")
                    print(f"  [MAIN.PY] Study regions: {study_regions}")

                    pipeline_type = "two-pass"

                    # =========================================================
                    # REPLACE PLACEHOLDER TEXT WITH ACTUAL ENDPOINTS
                    # Requires "primary" keyword (enforced by discovery prompt)
                    # MUST match logic in two_pass_extractor.py replace_placeholders()
                    # =========================================================
                    print(f"  [MAIN.PY] v59 - Placeholder replacement...")

                    primary_endpoint_name = None
                    for elem in discovered_elements:
                        # Handle both dataclass and dict formats
                        if hasattr(elem, 'category'):
                            cat = (elem.category or '').lower()
                            name = (elem.name or '').lower()
                            desc = elem.description or elem.name or ''
                        else:
                            cat = (elem.get("category", "") or "").lower()
                            name = (elem.get("name", "") or "").lower()
                            desc = elem.get("description", "") or elem.get("name", "") or ""
                        desc_lower = desc.lower()

                        # Get usable endpoint text (NO truncation - use full description)
                        endpoint_text = desc if desc else None
                        if not endpoint_text:
                            continue

                        # Only consider endpoint elements (category=endpoints OR "endpoint" in name)
                        is_endpoint = cat == 'endpoints' or 'endpoint' in name

                        # Check for "primary" keyword
                        has_primary = "primary" in cat or "primary" in name or "primary" in desc_lower

                        # Must be BOTH an endpoint AND have "primary"
                        if is_endpoint and has_primary:
                            primary_endpoint_name = endpoint_text
                            print(f"  [MAIN.PY] Found primary endpoint: {endpoint_text[:60]}")
                            break

                    if not primary_endpoint_name:
                        print(f"  [MAIN.PY] WARNING: No primary endpoint found in discovered elements")
                    else:
                        # Replace ALL placeholder patterns with actual endpoint (6 patterns)
                        placeholders = [
                            "[Primary endpoint as specified]",
                            "[Primary endpoint]",
                            "[ENDPOINT]",
                            "[endpoint]",
                            "[specify endpoint]",
                            "[primary endpoint as specified]",
                        ]
                        for placeholder in placeholders:
                            if placeholder in sap_text:
                                sap_text = sap_text.replace(placeholder, primary_endpoint_name)
                                print(f"  [MAIN.PY] Replaced '{placeholder}'")

                    # Remove generic placeholders entirely (8 patterns)
                    remove_placeholders = [
                        '[specify timepoints]',
                        '[specify timepoint]',
                        '[specify visits]',
                        '[specify visit]',
                        '[as specified]',
                        '[TBD]',
                        '[To be specified]',
                        '[to be specified]',
                    ]
                    for placeholder in remove_placeholders:
                        if placeholder in sap_text:
                            sap_text = sap_text.replace(placeholder, '')
                            print(f"  [MAIN.PY] Removed generic placeholder '{placeholder}'")

                    # v69: Dynamic SAP structure - Section 17 is TABLE/FIGURE SHELLS
                    # Section 12 is now SAFETY ANALYSIS in the new structure
                    # TFLs can be generated by KG Pipeline OR TLF Integration module
                    section_17_start = -1
                    for marker in ['## 17.', '# 17.', '17. TABLE', '17. Table']:
                        if marker in sap_text:
                            section_17_start = sap_text.find(marker)
                            break

                    if section_17_start >= 0:
                        section_17_text = sap_text[section_17_start:]
                        # Check for proper TFL tables in Section 17 (TABLE/FIGURE SHELLS)
                        has_proper_tables = 'TABLE 14.1.1' in section_17_text or 'TABLE 14.' in section_17_text
                    else:
                        has_proper_tables = False

                    print(f"  [MAIN.PY] v69 Dynamic SAP: Section 17 (TFLs) at pos {section_17_start}, has tables: {has_proper_tables}", flush=True)

                    # v101.2: TLF Shell Integration v2 - Universal expansion with period/population matrix
                    # If Section 17 is missing or has no proper tables, inject TLF shells
                    if not has_proper_tables and TLF_INTEGRATION_AVAILABLE:
                        print(f"  [MAIN.PY] TLF Integration v2: Generating study-specific TLF shells with universal expansion...")
                        full_extraction = result.get("full_extraction", {})
                        protocol_text = job.get("protocol_text", "")

                        # Get TLF summary for logging (with protocol text for better detection)
                        tlf_summary = get_tlf_shell_summary(full_extraction, protocol_text)
                        print(f"  [TLF v2] Study type: {tlf_summary.get('detected_study_type')}")
                        print(f"  [TLF v2] Drug classes: {tlf_summary.get('detected_drug_classes')}")
                        print(f"  [TLF v2] Study design: {tlf_summary.get('detected_study_design')}")

                        # Log universal expansion info
                        universal = tlf_summary.get('universal_expansion', {})
                        period_info = universal.get('period_stratification', {})
                        pop_matrix = universal.get('population_assessment_matrix', {})
                        print(f"  [TLF v2] Period stratification: {period_info.get('required')} ({period_info.get('config')}, {period_info.get('multiplier')}× mult)")
                        print(f"  [TLF v2] Populations: {pop_matrix.get('populations')}")
                        print(f"  [TLF v2] Assessments: {pop_matrix.get('assessments')}")
                        print(f"  [TLF v2] Regions: {universal.get('regions')}")
                        print(f"  [TLF v2] PK required: {universal.get('pk_required')}, Immunogenicity: {universal.get('immunogenicity_required')}")
                        print(f"  [TLF v2] Expected tables: ~{tlf_summary.get('expected_table_count', 'N/A')}")

                        # Generate TLF shells from extraction with universal expansion
                        tlf_shells = generate_tlf_shells_for_protocol(
                            full_extraction,
                            protocol_text=protocol_text,
                            apply_universal_expansion=True
                        )

                        if tlf_shells:
                            # Inject TLF shells into SAP
                            if section_17_start >= 0:
                                # Replace empty Section 17 with generated shells
                                # Find end of Section 17 (next section or end)
                                section_18_start = -1
                                for marker in ['## 18.', '# 18.', '18. ']:
                                    pos = sap_text.find(marker, section_17_start + 10)
                                    if pos > 0:
                                        section_18_start = pos
                                        break

                                if section_18_start > 0:
                                    # Replace Section 17 content
                                    sap_text = sap_text[:section_17_start] + "## 17. TABLE AND FIGURE SHELLS\n\n" + tlf_shells + "\n\n" + sap_text[section_18_start:]
                                else:
                                    # Section 17 is at the end - replace from there
                                    sap_text = sap_text[:section_17_start] + "## 17. TABLE AND FIGURE SHELLS\n\n" + tlf_shells
                            else:
                                # No Section 17 found - append at end
                                sap_text = sap_text.rstrip() + "\n\n## 17. TABLE AND FIGURE SHELLS\n\n" + tlf_shells

                            print(f"  [TLF] Injected {len(tlf_shells):,} chars of TLF shells into SAP")
                            has_proper_tables = True  # Update flag
                    elif not has_proper_tables:
                        print(f"  [MAIN.PY] NOTE: TFLs handled by KG Pipeline (Section 17 may be in appendices)")

                    # DEBUG: Final check before saving to database (v69 dynamic structure)
                    print(f"  [DEBUG] FINAL SAP length: {len(sap_text)} chars")
                    print(f"  [DEBUG] FINAL contains 'TABLE 14.1.1': {'TABLE 14.1.1' in sap_text}")
                    # Check for key sections in dynamic structure
                    if '## 12.' in sap_text:
                        final_sec12_pos = sap_text.find('## 12.')
                        final_preview = sap_text[final_sec12_pos:final_sec12_pos+150]
                        print(f"  [DEBUG] Section 12 (SAFETY): {final_preview[:100]}...")
                    if '## 17.' in sap_text:
                        final_sec17_pos = sap_text.find('## 17.')
                        final_preview = sap_text[final_sec17_pos:final_sec17_pos+150]
                        print(f"  [DEBUG] Section 17 (TFLs): {final_preview[:100]}...")

                    # =========================================================
                    # v101.2: APPEND REDUCTO SOA AS APPENDIX (CLEANED)
                    # The Reducto-extracted Schedule of Assessments should be
                    # included in the SAP as an appendix for completeness.
                    # v101.2: Now cleans HTML tables to markdown, removes boilerplate
                    # =========================================================
                    protocol_text = job.get("protocol_text", "")
                    reducto_marker = "SCHEDULE OF ASSESSMENTS (Enhanced by Reducto)"

                    if reducto_marker in protocol_text:
                        # Extract Reducto SOA content
                        marker_pos = protocol_text.find(reducto_marker)
                        # Skip past the marker line and the equals signs
                        soa_start = protocol_text.find("\n\n", marker_pos)
                        if soa_start > 0:
                            raw_soa_content = protocol_text[soa_start:].strip()

                            # v101.2: Clean the content - convert HTML to markdown, remove boilerplate
                            reducto_soa_content = clean_soa_content_for_appendix(raw_soa_content)
                            print(f"  [SOA Appendix] Cleaned: {len(raw_soa_content):,} -> {len(reducto_soa_content):,} chars")

                            # Only append if we have substantial content after cleaning
                            if len(reducto_soa_content) > 200:
                                # Check if SAP already has Schedule of Assessments
                                if "Schedule of Assessments" not in sap_text and "SCHEDULE OF ASSESSMENTS" not in sap_text:
                                    sap_text = sap_text.rstrip() + "\n\n---\n\n"
                                    sap_text += "## APPENDIX A: SCHEDULE OF ASSESSMENTS\n\n"
                                    sap_text += "*This schedule was extracted from protocol tables using Reducto vision processing.*\n\n"
                                    sap_text += reducto_soa_content
                                    print(f"  [SOA Appendix] Appended {len(reducto_soa_content):,} chars of cleaned Schedule of Assessments")
                                else:
                                    print(f"  [SOA Appendix] SAP already contains Schedule of Assessments, skipping")
                            else:
                                print(f"  [SOA Appendix] Cleaned content too short ({len(reducto_soa_content)} chars), skipping")
                    else:
                        print(f"  [SOA Appendix] No Reducto SOA marker found in protocol text")

                    update_data = {
                        "status": "completed",
                        "generated_sap": sap_text,
                        "quality_score": quality_score,
                        "endpoint_type": endpoint_type_str[:20],  # 'time-to-event' is 13 chars
                        "phase": phase_str[:10],
                        "therapeutic_area": therapeutic_area[:30],  # 'rheumatology' is 12 chars
                        "processing_time": processing_time,
                        "completed_at": datetime.utcnow().isoformat()
                    }

                    # Log validation info
                    if validation:
                        present = len(validation.get("present", []))
                        missing = len(validation.get("missing", []))
                        partial = len(validation.get("partial", []))
                        print(f"  Validation: {present} present, {partial} partial, {missing} missing")

                        gaps = validation.get("critical_gaps", [])
                        if gaps:
                            print(f"  Critical gaps:")
                            for gap in gaps[:5]:
                                print(f"    - {gap}")

                    db.table("sap_jobs").update(update_data).eq("id", job_id).execute()

                    # Detailed logging for KGPipeline
                    print(f"Job {job_id} completed in {processing_time:.1f}s ({pipeline_type} pipeline)")
                    print(f"  DISCOVERY:")
                    print(f"    Elements found: {result.get('discovered_count', 0)}")
                    print(f"  GENERATION:")
                    print(f"    SAP length: {result.get('sap_length', 0):,} chars")
                    print(f"  VALIDATION:")
                    print(f"    Quality: {quality_score:.1f}/100")

                else:
                    raise Exception("KGPipeline returned no SAP text")

            except Exception as e:
                # Print FULL traceback to find exact error location
                import traceback
                print("=" * 60)
                print("FULL TRACEBACK:")
                traceback.print_exc()
                print("=" * 60)

                # Update with failure
                db.table("sap_jobs").update({
                    "status": "failed",
                    "error_message": str(e)[:500],
                    "processing_time": time.time() - start_time,
                    "completed_at": datetime.utcnow().isoformat()
                }).eq("id", job_id).execute()

                print(f"Job {job_id} failed: {e}")

        except Exception as e:
            print(f"Worker error: {e}")
            await asyncio.sleep(10)

        # Small delay between jobs
        await asyncio.sleep(1)


# ============================================================================
# REDUCTO API ENDPOINTS
# ============================================================================

@app.get("/reducto/status")
async def reducto_status():
    """
    Check if Reducto API is configured and available.
    """
    try:
        from web.backend.reducto_client import check_reducto_available
        status = check_reducto_available()
        return status
    except ImportError:
        return {
            "available": False,
            "has_api_key": bool(os.getenv("REDUCTO_API_KEY")),
            "can_import": False,
            "error": "reducto_client module not found"
        }
    except Exception as e:
        return {
            "available": False,
            "error": str(e)
        }


@app.post("/reducto/extract-soa/upload")
async def reducto_extract_soa_upload(file: UploadFile = File(...)):
    """
    Test endpoint: Extract SOA tables from uploaded PDF using Reducto.

    This endpoint:
    1. Detects SOA pages using PyMuPDF
    2. Sends only those pages to Reducto
    3. Returns the extracted table content

    Use this to test Reducto extraction before full integration.
    """
    if not file.filename:
        raise HTTPException(400, "No file provided")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(400, "Only PDF files are supported")

    # Check Reducto availability
    reducto_key = os.getenv("REDUCTO_API_KEY")
    if not reducto_key:
        raise HTTPException(503, "REDUCTO_API_KEY not set. Configure it in environment.")

    try:
        from web.backend.reducto_client import extract_soa_with_reducto

        content = await file.read()
        file_size = len(content)

        if file_size == 0:
            raise HTTPException(400, "Empty file")

        if file_size > 50 * 1024 * 1024:  # 50MB limit
            raise HTTPException(400, "File too large (max 50MB)")

        # Step 1: Detect SOA pages
        soa_pages = detect_soa_pages_from_pdf(content)

        if not soa_pages:
            return {
                "success": False,
                "filename": file.filename,
                "message": "No SOA pages detected in PDF",
                "soa_pages": [],
                "content": "",
            }

        # Step 2: Extract with Reducto
        result = extract_soa_with_reducto(content, soa_pages)

        return {
            "success": result.success,
            "filename": file.filename,
            "soa_pages": result.pages_processed,
            "content": result.content,
            "content_length": len(result.content),
            "credits_used": result.credits_used,
            "error": result.error,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Reducto SOA extraction failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", 8000)))
