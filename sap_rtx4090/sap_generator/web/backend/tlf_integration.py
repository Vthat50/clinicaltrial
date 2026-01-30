"""
TLF Shell Integration Module v4
================================
Protocol-driven TLF shell generation for any therapeutic area.

Architecture:
  Step 1: Claude extracts structured FACTS from the protocol (returns JSON schema of facts)
  Step 2: Deterministic Python code builds the full TLF list from those facts (no API call)
  Step 3: YAML configs provide standard formatting + Markdown renderer produces shells

YAML configs used:
  - core/generation_rules.yaml    → TLF generation rules (what tables/figures/listings to create)
  - core/base_formatting.yaml     → universal formatting (populations, columns, footnotes, numbering)
  - core/analysis_templates.yaml  → statistical method → row structure mappings
  - therapeutic_areas/*.yaml      → disease-specific vocabularies (loaded for prompt context)
"""

import os
import json
import re
import traceback
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union
from io import BytesIO

import yaml

# =============================================================================
# YAML CONFIGURATION LOADING
# =============================================================================
YAML_DIR = Path(__file__).parent.parent.parent / "enterprise_sap_system" / "specs" / "tlf_configs"

_BASE_FMT = None
_ANALYSIS_TEMPLATES = None
_TA_CONFIGS = None
_GENERATION_RULES = None
_TA_AREA_RULES = None  # therapeutic_area_rules.yaml (from skills/INSTRUCTIONS.md)


def _load_yaml_configs():
    """Load YAML configs at first use. Returns (base_formatting, analysis_templates, ta_configs)."""
    global _BASE_FMT, _ANALYSIS_TEMPLATES, _TA_CONFIGS, _GENERATION_RULES
    if _BASE_FMT is not None:
        return _BASE_FMT, _ANALYSIS_TEMPLATES, _TA_CONFIGS

    core_dir = YAML_DIR / "core"
    ta_dir = YAML_DIR / "therapeutic_areas"

    # Base formatting (required)
    bf_path = core_dir / "base_formatting.yaml"
    if bf_path.exists():
        with open(bf_path) as f:
            _BASE_FMT = yaml.safe_load(f)
        print("[TLF Integration v4] Loaded base_formatting.yaml")
    else:
        print(f"[TLF Integration v4] WARNING: {bf_path} not found, using defaults")
        _BASE_FMT = {}

    # Analysis templates (required)
    at_path = core_dir / "analysis_templates.yaml"
    if at_path.exists():
        with open(at_path) as f:
            _ANALYSIS_TEMPLATES = yaml.safe_load(f)
        print("[TLF Integration v4] Loaded analysis_templates.yaml")
    else:
        print(f"[TLF Integration v4] WARNING: {at_path} not found")
        _ANALYSIS_TEMPLATES = {}

    # Therapeutic area configs (optional, loaded for prompt context)
    _TA_CONFIGS = {}
    if ta_dir.exists():
        for ta_file in sorted(ta_dir.glob("*.yaml")):
            with open(ta_file) as f:
                ta_data = yaml.safe_load(f)
            ta_name = ta_file.stem
            _TA_CONFIGS[ta_name] = ta_data
        print(f"[TLF Integration v4] Loaded {len(_TA_CONFIGS)} therapeutic area configs: {list(_TA_CONFIGS.keys())}")

    # Generation rules (optional — falls back to hardcoded if missing)
    gr_path = core_dir / "generation_rules.yaml"
    if gr_path.exists():
        with open(gr_path) as f:
            _GENERATION_RULES = yaml.safe_load(f)
        print("[TLF Integration v4] Loaded generation_rules.yaml")
    else:
        print(f"[TLF Integration v4] generation_rules.yaml not found, using hardcoded rules")
        _GENERATION_RULES = None

    # Therapeutic area rules (auto-generated from skills/INSTRUCTIONS.md)
    global _TA_AREA_RULES
    ta_rules_path = core_dir / "therapeutic_area_rules.yaml"
    if ta_rules_path.exists():
        with open(ta_rules_path) as f:
            _TA_AREA_RULES = yaml.safe_load(f)
        areas = list((_TA_AREA_RULES or {}).get("areas", {}).keys())
        print(f"[TLF Integration v4] Loaded therapeutic_area_rules.yaml ({len(areas)} areas: {areas})")
    else:
        _TA_AREA_RULES = None

    return _BASE_FMT, _ANALYSIS_TEMPLATES, _TA_CONFIGS


# =============================================================================
# ANALYSIS METHOD LOOKUP INDEX
# =============================================================================
# Maps flat method name → (category, method_key) for analysis_templates.yaml lookup
_METHOD_INDEX = None


def _build_method_index():
    """Build flat lookup: method_name → (category, method_key) from analysis_templates.yaml."""
    global _METHOD_INDEX
    if _METHOD_INDEX is not None:
        return _METHOD_INDEX

    _, templates, _ = _load_yaml_configs()
    _METHOD_INDEX = {}
    skip_keys = {"version", "config_type"}
    for category, methods in templates.items():
        if category in skip_keys or not isinstance(methods, dict):
            continue
        for method_key in methods:
            _METHOD_INDEX[method_key] = (category, method_key)

    return _METHOD_INDEX


# =============================================================================
# TABLE TYPE → ICH SECTION MAPPING
# =============================================================================
_TABLE_TYPE_TO_SECTION = {
    "disposition": "14.1",
    "demographics": "14.1",
    "medical_history": "14.1",
    "baseline": "14.1",
    "time_to_event": "14.2",
    "binary_response": "14.2",
    "continuous_endpoint": "14.2",
    "pro_qol": "14.2",
    "subgroup_forest": "14.2",
    "ae_overview": "14.3",
    "ae_by_soc_pt": "14.3",
    "ae_special": "14.3",
    "aesi": "14.3",
    "labs": "14.3",
    "lab_shift": "14.3",
    "vitals": "14.3",
    "ecg": "14.3",
    "exposure": "14.3",
    "concomitant_medications": "14.3",
    "pk_parameters": "14.4",
    "pk_concentration": "14.4",
    "immunogenicity": "14.4",
    "listing": "16.2",
}

# Table types that use descriptive/continuous stats (not n/%)
_DESCRIPTIVE_TYPES = {"labs", "vitals", "ecg", "pk_parameters", "pk_concentration", "continuous_endpoint"}

# Map analysis_templates category → table type for fallback
_TEMPLATE_CATEGORY_TO_TYPE = {
    "safety": {"ae_overview", "ae_by_soc_pt", "lab_shift", "exposure", "concomitant_medications"},
    "pharmacokinetics": {"pk_parameters", "pk_concentration", "immunogenicity"},
    "descriptive": {"demographics", "disposition", "medical_history"},
}


# =============================================================================
# STEP 1: CLAUDE EXTRACTS FACTS FROM THE PROTOCOL
# =============================================================================

# JSON schema description for the extraction prompt
_FACTS_SCHEMA = """{
  "study_design": {
    "phase": "Phase of the trial (e.g. Phase 3)",
    "type": "One of: superiority, equivalence, non_inferiority, biosimilar, single_arm",
    "blinding": "e.g. double-blind, open-label",
    "randomization": "e.g. 1:1, 2:1",
    "equivalence_margin": "If applicable, e.g. -12.5% to 12.5%. Empty string if not applicable."
  },
  "arms": [
    {"name": "Full arm name including backbone", "dose": "dose string", "route": "IV/SC/PO/etc"}
  ],
  "populations": [
    {"name": "Short name (ITT, Safety, PP, PK, etc.)", "definition": "One-sentence definition"}
  ],
  "endpoints": [
    {
      "name": "Endpoint name as stated in protocol",
      "type": "One of: binary, time_to_event, continuous, count",
      "primary": true,
      "analysis_method": "One of: logistic_regression, exact_binomial, clopper_pearson, cox_ph, kaplan_meier, ancova, mmrm, negative_binomial, descriptive",
      "covariates": ["List of covariates/stratification factors for this endpoint's model"],
      "populations": ["List of population short names this endpoint is analyzed in"],
      "reviews": ["List of review types if applicable, e.g. central, local. Empty list if N/A"],
      "response_criteria": "e.g. RECIST 1.1. Empty string if N/A",
      "landmark_timepoints": [6, 12, 24],
      "censoring_rules": "Brief description of censoring. Empty string if N/A",
      "extra_rows": ["Any additional row items, e.g. response categories CR/PR/SD/PD/NE"]
    }
  ],
  "study_periods": ["All study periods in order, e.g. Screening, Treatment, Follow-Up"],
  "treatment_periods": ["Only the treatment periods, e.g. Induction, Maintenance"],
  "assessments_collected": {
    "labs": true,
    "vitals": true,
    "ecg": true,
    "pk": true,
    "immunogenicity": true,
    "qol": ["List of QoL instrument names, or empty list if none"],
    "imaging": "Description of imaging or empty string",
    "physical_exam": true,
    "pregnancy_test": true,
    "ecog_ps": true,
    "viral_serology": true,
    "gene_screening": true
  },
  "aesis": [
    {"name": "AESI category name", "definition": "SMQ or PT list description"}
  ],
  "stratification_factors": ["List of randomization stratification factors"],
  "subgroups": ["List of pre-specified subgroup analyses, e.g. Age (<65, >=65), Sex"],
  "disease_specific_baseline": [
    "List of disease-specific baseline characteristics to include in demographics table"
  ],
  "backbone_therapies": ["Full description of each backbone therapy, e.g. Paclitaxel 200 mg/m2 IV"],
  "coding_dictionaries": {
    "ae": "MedDRA version string",
    "medications": "WHO Drug Dictionary version string"
  },
  "therapeutic_area": "e.g. oncology, immunology, cardiology"
}"""


def _build_fact_extraction_prompt(protocol_text: str) -> str:
    """Build the short extraction-only prompt for Claude. No table-building rules."""
    max_protocol_chars = 80000
    protocol_excerpt = protocol_text[:max_protocol_chars] if len(protocol_text) > max_protocol_chars else protocol_text

    prompt = f"""You are a biostatistician reading a clinical trial protocol. Extract the following information and return it as JSON.

Do NOT decide what tables to create. Only extract facts from the protocol.
Do NOT infer or guess. If the protocol does not explicitly state a value, use null for strings, empty list [] for arrays, or null for objects.

For example:
- If the protocol says "logistic regression model" → analysis_method: "logistic_regression"
- If the protocol says "ORR will be compared" but does NOT name a method → analysis_method: null
- If the protocol says "equivalence" or "biosimilar" → study_design.type: "biosimilar"
- If the protocol does not clearly state the design type → study_design.type: null
- If the protocol does not mention censoring rules → censoring_rules: null

For boolean fields in assessments_collected, use true ONLY if the protocol explicitly mentions collecting that assessment, false otherwise.
For list fields, include ALL items mentioned in the protocol — do not summarize or skip any.
For endpoints, include ALL primary AND secondary endpoints.
For AESIs, include ALL adverse events of special interest defined in the protocol.

Return ONLY valid JSON (no markdown fences, no explanation) with this schema:

{_FACTS_SCHEMA}

PROTOCOL TEXT:
{protocol_excerpt}"""

    return prompt


def _extract_protocol_facts(protocol_text: str, extraction: Dict) -> Dict:
    """Send protocol to Claude API. Returns parsed JSON of extracted facts."""
    try:
        from anthropic import Anthropic
    except ImportError:
        raise RuntimeError("anthropic package not installed")

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set")

    client = Anthropic()
    prompt = _build_fact_extraction_prompt(protocol_text)

    print("[TLF Integration v4] Asking Claude to extract protocol facts...")

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8192,
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = response.content[0].text.strip()

    # Strip markdown code fences if present
    if response_text.startswith("```"):
        lines = response_text.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        response_text = "\n".join(lines)

    try:
        facts = json.loads(response_text)
    except json.JSONDecodeError:
        # Retry with stricter instruction
        print("[TLF Integration v4] JSON parse failed, retrying with stricter prompt...")
        retry_response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            messages=[
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": response_text},
                {"role": "user", "content": "Your response was not valid JSON. Return ONLY the JSON object, no markdown fences, no explanation."}
            ]
        )
        retry_text = retry_response.content[0].text.strip()
        if retry_text.startswith("```"):
            lines = retry_text.split("\n")
            lines = [l for l in lines if not l.strip().startswith("```")]
            retry_text = "\n".join(lines)
        facts = json.loads(retry_text)

    # Validate and fill defaults for critical fields
    if "study_design" not in facts:
        facts["study_design"] = {"phase": "", "type": "superiority", "blinding": "", "randomization": "", "equivalence_margin": ""}
    if "arms" not in facts or not facts["arms"]:
        # Fall back to extraction dict
        arms_from_ext = extraction.get("treatment_arms", [])
        facts["arms"] = []
        for i, a in enumerate(arms_from_ext):
            if isinstance(a, str):
                facts["arms"].append({"name": a, "dose": "", "route": ""})
            elif isinstance(a, dict):
                facts["arms"].append({"name": a.get("arm_name", a.get("drug_name", f"Arm {i+1}")), "dose": a.get("dose", ""), "route": a.get("route", "")})
    if "populations" not in facts or not facts["populations"]:
        facts["populations"] = [
            {"name": "ITT", "definition": "All randomized subjects"},
            {"name": "Safety", "definition": "All subjects who received at least one dose"}
        ]
    if "endpoints" not in facts:
        facts["endpoints"] = []
    if "aesis" not in facts:
        facts["aesis"] = []
    if "assessments_collected" not in facts:
        facts["assessments_collected"] = {}
    if "study_periods" not in facts:
        facts["study_periods"] = []
    if "treatment_periods" not in facts:
        facts["treatment_periods"] = []
    if "stratification_factors" not in facts:
        facts["stratification_factors"] = []
    if "subgroups" not in facts:
        facts["subgroups"] = []
    if "disease_specific_baseline" not in facts:
        facts["disease_specific_baseline"] = []
    if "backbone_therapies" not in facts:
        facts["backbone_therapies"] = []
    if "coding_dictionaries" not in facts:
        facts["coding_dictionaries"] = {}

    # Ensure ITT and Safety are in populations
    pop_names = [p["name"] for p in facts["populations"]]
    if "ITT" not in pop_names:
        facts["populations"].append({"name": "ITT", "definition": "All randomized subjects"})
    if "Safety" not in pop_names:
        facts["populations"].append({"name": "Safety", "definition": "All subjects who received at least one dose"})

    print(f"[TLF Integration v4] Extracted facts: {len(facts.get('arms', []))} arms, "
          f"{len(facts.get('endpoints', []))} endpoints, {len(facts.get('aesis', []))} AESIs, "
          f"{len(facts.get('populations', []))} populations")
    print(f"[TLF Integration v4] Study design: {facts['study_design'].get('type', 'unknown')}, "
          f"Phase: {facts['study_design'].get('phase', 'unknown')}")
    print(f"[TLF Integration v4] Treatment periods: {facts.get('treatment_periods', [])}")

    return facts


# =============================================================================
# CONDITION EVALUATOR & PLACEHOLDER SUBSTITUTION (for generation_rules.yaml)
# =============================================================================
def _resolve_path(obj: Any, path: str) -> Any:
    """Resolve a dot-separated path against a data object.
    E.g. 'facts.assessments_collected.labs' resolves through nested dicts.
    """
    parts = path.split(".")
    current = obj
    for part in parts:
        if isinstance(current, dict):
            current = current.get(part)
        else:
            return None
        if current is None:
            return None
    return current


def _compare(left: Any, op: str, right: Any) -> bool:
    """Compare two values with the given operator."""
    if op == "==":
        if right == "true":
            return left is True or left == "true"
        if right == "false":
            return left is False or left == "false"
        return str(left) == str(right)
    elif op == "!=":
        return not _compare(left, "==", right)
    elif op == ">":
        try:
            return float(left) > float(right)
        except (TypeError, ValueError):
            return False
    elif op == "<":
        try:
            return float(left) < float(right)
        except (TypeError, ValueError):
            return False
    elif op == ">=":
        try:
            return float(left) >= float(right)
        except (TypeError, ValueError):
            return False
    elif op == "<=":
        try:
            return float(left) <= float(right)
        except (TypeError, ValueError):
            return False
    return False


def _evaluate_condition(condition: str, context: Dict) -> bool:
    """Evaluate a simple condition expression from YAML against a context dict.

    Supported patterns:
      path == value
      path in [a, b, c]
      path | length > N
      path | is_list_with_items
      path is not empty
      A or B
      A and B
    """
    if not condition or not condition.strip():
        return True

    condition = condition.strip()

    # Handle OR combinator
    if " or " in condition:
        parts = condition.split(" or ")
        return any(_evaluate_condition(p.strip(), context) for p in parts)

    # Handle AND combinator
    if " and " in condition:
        parts = condition.split(" and ")
        return all(_evaluate_condition(p.strip(), context) for p in parts)

    # Pattern: path | is_list_with_items
    if "| is_list_with_items" in condition:
        path = condition.split("|")[0].strip()
        val = _resolve_path(context, path)
        return isinstance(val, list) and len(val) > 0

    # Pattern: path | length > N
    m = re.match(r'^(.+?)\s*\|\s*length\s*(>|<|>=|<=|==|!=)\s*(\d+)$', condition)
    if m:
        path, op, num = m.group(1).strip(), m.group(2), m.group(3)
        val = _resolve_path(context, path)
        length = len(val) if isinstance(val, (list, dict, str)) else 0
        return _compare(length, op, num)

    # Pattern: path is not empty
    if " is not empty" in condition:
        path = condition.replace(" is not empty", "").strip()
        val = _resolve_path(context, path)
        if val is None:
            return False
        if isinstance(val, (str, list, dict)):
            return len(val) > 0
        return True

    # Pattern: path in [a, b, c]
    m = re.match(r'^(.+?)\s+in\s+\[(.+)\]$', condition)
    if m:
        path = m.group(1).strip()
        items_str = m.group(2)
        items = [i.strip().strip('"').strip("'") for i in items_str.split(",")]
        val = _resolve_path(context, path)
        return str(val) in items

    # Pattern: path == value / path != value / path > value etc.
    m = re.match(r'^(.+?)\s*(==|!=|>=|<=|>|<)\s*(.+)$', condition)
    if m:
        path, op, right = m.group(1).strip(), m.group(2), m.group(3).strip().strip('"').strip("'")
        val = _resolve_path(context, path)
        return _compare(val, op, right)

    return False


def _substitute_placeholders(template: str, context: Dict) -> str:
    """Replace {placeholder} tokens in a string using context dict.

    Supported placeholders:
      {endpoint.name}, {population}, {review}, {period},
      {aesi.name}, {aesi.definition}, {therapy.drug_name},
      {instrument}, {covariates}, {margin}, {landmarks},
      {censoring}, {criteria}
    """
    if "{" not in template:
        return template

    result = template

    # Direct dict lookups for simple top-level keys
    simple_keys = ["population", "review", "period", "instrument"]
    for key in simple_keys:
        placeholder = "{" + key + "}"
        if placeholder in result:
            result = result.replace(placeholder, str(context.get(key, "")))

    # Dotted-path placeholders
    dotted = re.findall(r'\{([a-zA-Z_][a-zA-Z0-9_.]+)\}', result)
    for path in dotted:
        placeholder = "{" + path + "}"
        if placeholder in result:
            val = _resolve_path(context, path)
            if val is not None:
                result = result.replace(placeholder, str(val))

    # Special computed placeholders
    if "{covariates}" in result:
        covariates = context.get("covariates", context.get("endpoint", {}).get("covariates", []))
        if isinstance(covariates, list):
            result = result.replace("{covariates}", ", ".join(str(c) for c in covariates))
        else:
            result = result.replace("{covariates}", str(covariates))

    if "{margin}" in result:
        margin = context.get("margin", context.get("facts", {}).get("study_design", {}).get("equivalence_margin", ""))
        result = result.replace("{margin}", str(margin))

    if "{landmarks}" in result:
        landmarks = context.get("landmarks", context.get("endpoint", {}).get("landmark_timepoints", []))
        if isinstance(landmarks, list):
            result = result.replace("{landmarks}", ", ".join(f"{t}-month" for t in landmarks))
        else:
            result = result.replace("{landmarks}", str(landmarks))

    if "{censoring}" in result:
        censoring = context.get("censoring", context.get("endpoint", {}).get("censoring_rules", ""))
        result = result.replace("{censoring}", str(censoring))

    if "{criteria}" in result:
        criteria = context.get("criteria", context.get("endpoint", {}).get("response_criteria", ""))
        result = result.replace("{criteria}", str(criteria))

    return result


# =============================================================================
# STEP 2: DETERMINISTIC TABLE BUILDER (no API call)
# =============================================================================
def _build_endpoint_footnotes(ep: Dict, facts: Dict) -> List[str]:
    """Generate study-specific footnotes — reads from YAML if available, else hardcoded."""
    global _GENERATION_RULES
    _load_yaml_configs()
    if _GENERATION_RULES and "endpoint_footnote_templates" in _GENERATION_RULES:
        return _build_endpoint_footnotes_yaml(ep, facts)
    return _build_endpoint_footnotes_hardcoded(ep, facts)


def _build_endpoint_footnotes_yaml(ep: Dict, facts: Dict) -> List[str]:
    """Generate endpoint footnotes from YAML templates."""
    templates = _GENERATION_RULES["endpoint_footnote_templates"]
    footnotes = []
    method = ep.get("analysis_method", "")
    covariates = ep.get("covariates", [])
    design_type = facts.get("study_design", {}).get("type", "")
    margin = facts.get("study_design", {}).get("equivalence_margin", "")

    # Build context for condition evaluation and placeholder substitution
    ctx = {
        "facts": facts,
        "endpoint": ep,
        "covariates": covariates,
        "margin": margin,
    }

    # Pre-compute comma-joined covariates for substitution
    covariates_str = ", ".join(covariates) if isinstance(covariates, list) else str(covariates)
    ctx["covariates"] = covariates_str

    # Method-specific templates
    method_templates = templates.get(method, [])
    for tmpl in method_templates:
        if not isinstance(tmpl, dict):
            continue
        # Check condition
        cond = tmpl.get("condition", "")
        if cond and not _evaluate_condition(cond, ctx):
            continue
        text = tmpl.get("text", "")
        # Append conditional suffixes
        if covariates and "append_if_covariates" in tmpl:
            text += _substitute_placeholders(tmpl["append_if_covariates"], ctx)
        if margin and "append_if_margin" in tmpl:
            text += _substitute_placeholders(tmpl["append_if_margin"], ctx)
        if "suffix" in tmpl:
            text += tmpl["suffix"]
        if text:
            footnotes.append(text)

    # Common footnotes (applied for any method)
    common_templates = templates.get("_common", [])
    for tmpl in common_templates:
        if not isinstance(tmpl, dict):
            continue
        cond = tmpl.get("condition", "")
        if cond and not _evaluate_condition(cond, ctx):
            continue
        text = _substitute_placeholders(tmpl.get("text", ""), ctx)
        if text:
            footnotes.append(text)

    return footnotes


def _build_endpoint_footnotes_hardcoded(ep: Dict, facts: Dict) -> List[str]:
    """Generate study-specific footnotes from extracted facts for an endpoint (hardcoded fallback)."""
    footnotes = []
    design_type = facts.get("study_design", {}).get("type", "")
    margin = facts.get("study_design", {}).get("equivalence_margin", "")
    coding = facts.get("coding_dictionaries", {})

    method = ep.get("analysis_method", "")
    covariates = ep.get("covariates", [])
    censoring = ep.get("censoring_rules", "")
    criteria = ep.get("response_criteria", "")
    landmarks = ep.get("landmark_timepoints", [])

    # Analysis method footnote
    if method == "logistic_regression":
        fn = "Logistic regression model"
        if covariates:
            fn += f" adjusted for {', '.join(covariates)}"
        fn += "."
        footnotes.append(fn)
        if design_type in ("biosimilar", "equivalence"):
            ci_fn = "Risk difference and 95% CI estimated by Miettinen-Nurminen method."
            if margin:
                ci_fn += f" Equivalence margin: {margin}."
            footnotes.append(ci_fn)
    elif method == "cox_ph":
        fn = "Hazard ratio from Cox proportional hazards model"
        if covariates:
            fn += f" stratified by {', '.join(covariates)}"
        fn += "."
        footnotes.append(fn)
    elif method == "kaplan_meier":
        footnotes.append("Kaplan-Meier estimates; 95% CI by Brookmeyer-Crowley method.")
    elif method in ("exact_binomial", "clopper_pearson"):
        footnotes.append("95% CI calculated using Clopper-Pearson exact method.")
    elif method == "ancova":
        fn = "LS means and p-value from ANCOVA model"
        if covariates:
            fn += f" with covariates: {', '.join(covariates)}"
        fn += "."
        footnotes.append(fn)
    elif method == "mmrm":
        fn = "LS means from mixed model for repeated measures (MMRM)"
        if covariates:
            fn += f" with covariates: {', '.join(covariates)}"
        fn += "."
        footnotes.append(fn)
    elif method == "negative_binomial":
        fn = "Rate ratio from negative binomial regression"
        if covariates:
            fn += f" adjusted for {', '.join(covariates)}"
        fn += "."
        footnotes.append(fn)

    # Censoring footnote
    if censoring:
        footnotes.append(f"Censoring: {censoring}.")

    # Response criteria footnote
    if criteria:
        footnotes.append(f"Response assessed per {criteria}.")

    # Landmark timepoints footnote
    if landmarks:
        timepoint_strs = [f"{t}-month" for t in landmarks]
        footnotes.append(f"Landmark timepoints: {', '.join(timepoint_strs)}.")

    return footnotes


def _default_analysis_method(endpoint_type: Optional[str], facts: Dict) -> str:
    """Assign a default analysis method — reads from YAML if available, else hardcoded."""
    global _GENERATION_RULES
    _load_yaml_configs()
    if _GENERATION_RULES and "default_analysis_methods" in _GENERATION_RULES:
        return _default_analysis_method_yaml(endpoint_type, facts)
    return _default_analysis_method_hardcoded(endpoint_type, facts)


def _default_analysis_method_yaml(endpoint_type: Optional[str], facts: Dict) -> str:
    """Assign default analysis method from YAML rules (first match wins)."""
    rules = _GENERATION_RULES["default_analysis_methods"]
    design_type = facts.get("study_design", {}).get("type") or ""

    for rule in rules:
        rule_ep_type = rule.get("endpoint_type", "")
        # _default matches anything
        if rule_ep_type != "_default" and rule_ep_type != endpoint_type:
            continue
        # Check optional design_type_in filter
        design_filter = rule.get("design_type_in")
        if design_filter and design_type not in design_filter:
            continue
        return rule.get("method", "descriptive")

    return "descriptive"


def _default_analysis_method_hardcoded(endpoint_type: Optional[str], facts: Dict) -> str:
    """Assign default analysis method (hardcoded fallback)."""
    design_type = facts.get("study_design", {}).get("type") or ""

    if endpoint_type == "binary":
        if design_type in ("biosimilar", "equivalence"):
            return "logistic_regression"
        return "exact_binomial"
    elif endpoint_type == "time_to_event":
        return "cox_ph"
    elif endpoint_type == "continuous":
        return "ancova"
    elif endpoint_type == "count":
        return "negative_binomial"
    return "descriptive"


def _build_tlf_list(facts: Dict) -> Dict:
    """Build the complete TLF list — reads from YAML if available, else hardcoded."""
    global _GENERATION_RULES
    _load_yaml_configs()
    if _GENERATION_RULES:
        return _build_tlf_list_yaml(facts)
    return _build_tlf_list_hardcoded(facts)


def _build_tlf_list_yaml(facts: Dict) -> Dict:
    """
    Build the complete TLF list from generation_rules.yaml.
    Generic rules engine — iterates YAML sections, evaluates conditions,
    substitutes placeholders.
    """
    rules = _GENERATION_RULES
    tables = []
    figures = []
    listings = []

    arms = facts.get("arms", [])
    endpoints = facts.get("endpoints", [])
    aesis = facts.get("aesis", [])
    periods = facts.get("treatment_periods", [])
    assessments = facts.get("assessments_collected", {})
    populations = facts.get("populations", [])
    subgroups = facts.get("subgroups", [])

    # Context for condition evaluation
    ctx = {"facts": facts}

    # ── MANDATORY TABLES ──
    for tbl in rules.get("mandatory_tables", []):
        entry = {"type": tbl["type"], "title": tbl["title"], "population": tbl["population"]}
        extra_from = tbl.get("extra_rows_from")
        if extra_from:
            entry["extra_rows"] = _resolve_path(ctx, extra_from) or []
        tables.append(entry)

    # ── ENDPOINT TABLES ──
    ep_config = rules.get("endpoint_tables", {})
    type_mapping = ep_config.get("type_mapping", {})
    title_pattern = ep_config.get("title_pattern", "{endpoint.name} ({population} Population)")
    title_review_pattern = ep_config.get("title_with_review_pattern", "{endpoint.name} ({review} Review) ({population} Population)")

    for ep in endpoints:
        if not ep.get("analysis_method"):
            ep["analysis_method"] = _default_analysis_method(ep.get("type"), facts)

        mapped_type = type_mapping.get(ep.get("type", ""), "continuous_endpoint")

        for pop in ep.get("populations", ["ITT"]) or ["ITT"]:
            reviews = ep.get("reviews") or [None]
            for review in reviews:
                ep_ctx = {"endpoint": ep, "population": pop, "review": review.title() if review else None}
                if review:
                    title = _substitute_placeholders(title_review_pattern, ep_ctx)
                else:
                    title = _substitute_placeholders(title_pattern, ep_ctx)

                tables.append({
                    "type": mapped_type,
                    "title": title,
                    "population": pop,
                    "endpoint": ep["name"],
                    "analysis_method": ep.get("analysis_method"),
                    "footnotes": _build_endpoint_footnotes(ep, facts),
                    "extra_rows": ep.get("extra_rows", [])
                })

    # ── AESI TABLES ──
    aesi_config = rules.get("aesi_tables", {})
    if aesi_config:
        for aesi in aesis:
            aesi_ctx = {"aesi": aesi}
            title = _substitute_placeholders(aesi_config.get("title_pattern", "AESI - {aesi.name}"), aesi_ctx)
            entry = {
                "type": aesi_config.get("type", "ae_special"),
                "title": title,
                "population": aesi_config.get("population", "Safety"),
            }
            extra_from = aesi_config.get("extra_rows_from")
            if extra_from:
                val = _resolve_path(aesi_ctx, extra_from)
                entry["extra_rows"] = [val] if isinstance(val, str) else (val or [])
            tables.append(entry)

    # ── ASSESSMENT-CONDITIONAL TABLES ──
    for group in rules.get("assessment_tables", []):
        cond = group.get("condition", "")
        if cond and not _evaluate_condition(cond, ctx):
            continue

        for_each = group.get("for_each")
        if for_each:
            items = _resolve_path(ctx, for_each) or []
            if not isinstance(items, list):
                items = [items]
            for item in items:
                for tbl in group.get("tables", []):
                    item_ctx = {"instrument": item, "facts": facts}
                    title = _substitute_placeholders(tbl["title"], item_ctx)
                    tables.append({"type": tbl["type"], "title": title, "population": tbl["population"]})
        else:
            for tbl in group.get("tables", []):
                tables.append({"type": tbl["type"], "title": tbl["title"], "population": tbl["population"]})

    # ── BACKBONE THERAPY TABLES ──
    bb_config = rules.get("backbone_therapy_tables", {})
    if bb_config:
        for therapy in facts.get("backbone_therapies", []):
            drug_name = therapy.split(" ")[0]
            therapy_ctx = {"therapy": {"drug_name": drug_name}}
            title = _substitute_placeholders(bb_config.get("title_pattern", "{therapy.drug_name} Exposure"), therapy_ctx)
            tables.append({
                "type": bb_config.get("type", "exposure"),
                "title": title,
                "population": bb_config.get("population", "Safety"),
            })

    # ── ADDITIONAL SAFETY TABLES ──
    for tbl in rules.get("additional_safety_tables", []):
        tables.append({"type": tbl["type"], "title": tbl["title"], "population": tbl["population"]})

    # ── DESIGN-CONDITIONAL SAFETY TABLES ──
    for group in rules.get("design_conditional_safety_tables", []):
        cond = group.get("condition", "")
        if cond and not _evaluate_condition(cond, ctx):
            continue
        for tbl in group.get("tables", []):
            tables.append({"type": tbl["type"], "title": tbl["title"], "population": tbl["population"]})

    # ── SUBGROUP TABLES ──
    sg_config = rules.get("subgroup_tables", {})
    if sg_config:
        global_cond = sg_config.get("global_condition", "")
        if not global_cond or _evaluate_condition(global_cond, ctx):
            for ep in endpoints:
                ep_ctx = {"endpoint": ep, "facts": facts}
                item_cond = sg_config.get("condition", "")
                if item_cond and not _evaluate_condition(item_cond, ep_ctx):
                    continue
                title = _substitute_placeholders(sg_config.get("title_pattern", "Subgroup Analysis of {endpoint.name}"), ep_ctx)
                tables.append({
                    "type": sg_config.get("type", "subgroup_forest"),
                    "title": title,
                    "population": sg_config.get("population", "ITT"),
                    "endpoint": ep.get("name", ""),
                })

    # ── PERIOD-SPLIT TABLES ──
    ps_config = rules.get("period_split_tables", {})
    if ps_config:
        cond = ps_config.get("condition", "")
        if not cond or _evaluate_condition(cond, ctx):
            for period in periods:
                period_ctx = {"period": period}
                for tbl in ps_config.get("tables", []):
                    # Per-table condition (e.g., biosimilar-only period-split tables)
                    tbl_cond = tbl.get("condition", "")
                    if tbl_cond and not _evaluate_condition(tbl_cond, ctx):
                        continue
                    title = _substitute_placeholders(tbl.get("title_pattern", tbl.get("title", "")), period_ctx)
                    tables.append({"type": tbl["type"], "title": title, "population": tbl["population"]})

    # ── FIGURES ──
    for fig_rule in rules.get("figure_rules", []):
        # Check global condition
        global_cond = fig_rule.get("global_condition", "")
        if global_cond and not _evaluate_condition(global_cond, ctx):
            continue

        for_each = fig_rule.get("for_each")
        if for_each:
            items = _resolve_path(ctx, for_each) or []
            for item in items:
                item_ctx = {"endpoint": item, "facts": facts}
                # Check per-item condition
                item_cond = fig_rule.get("condition", "")
                if item_cond and not _evaluate_condition(item_cond, item_ctx):
                    continue
                title = _substitute_placeholders(fig_rule.get("title_pattern", ""), item_ctx)
                ep_field = fig_rule.get("endpoint_field", "endpoint.name")
                figures.append({
                    "type": fig_rule.get("type", "figure"),
                    "title": title,
                    "endpoint": _resolve_path(item_ctx, ep_field) or "",
                })
        else:
            # Non-iterable figure (e.g., CONSORT, tipping point, lab trend)
            title = _substitute_placeholders(fig_rule.get("title_pattern", ""), ctx)
            figures.append({
                "type": fig_rule.get("type", "figure"),
                "title": title,
            })

    # ── MANDATORY LISTINGS ──
    for lst in rules.get("mandatory_listings", []):
        listings.append({
            "type": "listing",
            "title": f"Listing of {lst['title']}",
            "population": lst["population"],
            "variables": lst.get("columns", []),
        })

    # ── ASSESSMENT-CONDITIONAL LISTINGS ──
    for group in rules.get("assessment_listings", []):
        cond = group.get("condition", "")
        if cond and not _evaluate_condition(cond, ctx):
            continue

        for_each = group.get("for_each")
        if for_each:
            items = _resolve_path(ctx, for_each) or []
            if not isinstance(items, list):
                items = [items]
            for item in items:
                for lst in group.get("listings", []):
                    item_ctx = {"instrument": item, "facts": facts}
                    title = _substitute_placeholders(lst["title"], item_ctx)
                    listings.append({
                        "type": "listing",
                        "title": f"Listing of {title}",
                        "population": lst["population"],
                        "variables": lst.get("columns", []),
                    })
        else:
            for lst in group.get("listings", []):
                listings.append({
                    "type": "listing",
                    "title": f"Listing of {lst['title']}",
                    "population": lst["population"],
                    "variables": lst.get("columns", []),
                })

    # ── BACKBONE THERAPY LISTINGS ──
    bb_lst_config = rules.get("backbone_therapy_listings", {})
    if bb_lst_config:
        for therapy in facts.get("backbone_therapies", []):
            drug_name = therapy.split(" ")[0]
            therapy_ctx = {"therapy": {"drug_name": drug_name}}
            for lst in bb_lst_config.get("listings", []):
                title = _substitute_placeholders(lst["title"], therapy_ctx)
                listings.append({
                    "type": "listing",
                    "title": f"Listing of {title}",
                    "population": lst["population"],
                    "variables": lst.get("columns", []),
                })

    # ── BIOSIMILAR-CONDITIONAL LISTINGS ──
    bs_lst_config = rules.get("biosimilar_listings", {})
    if bs_lst_config:
        cond = bs_lst_config.get("condition", "")
        if not cond or _evaluate_condition(cond, ctx):
            for lst in bs_lst_config.get("listings", []):
                listings.append({
                    "type": "listing",
                    "title": f"Listing of {lst['title']}",
                    "population": lst["population"],
                    "variables": lst.get("columns", []),
                })

    # ── THERAPEUTIC AREA-SPECIFIC RULES ──
    # Apply area-specific mandatory/common tables from therapeutic_area_rules.yaml
    # These supplement (not replace) the core rules above.
    if _TA_AREA_RULES:
        ta = facts.get("therapeutic_area", "").lower().replace(" ", "_")
        area_rules = _TA_AREA_RULES.get("areas", {}).get(ta, {})
        if area_rules:
            existing_types = {t.get("type") for t in tables}
            added = 0

            # Mandatory area types — always add if not already present
            for rule in area_rules.get("mandatory_table_types", []):
                ttype = rule["type"]
                if ttype not in existing_types:
                    cond = rule.get("condition", "")
                    if cond and not _evaluate_condition(cond, ctx):
                        continue
                    tables.append({"type": ttype, "title": "", "population": "Safety", "_from_area_rules": True})
                    existing_types.add(ttype)
                    added += 1

            # Common area types — add if not present and condition met
            for rule in area_rules.get("common_table_types", []):
                ttype = rule["type"]
                if ttype not in existing_types:
                    cond = rule.get("condition", "")
                    if cond and not _evaluate_condition(cond, ctx):
                        continue
                    tables.append({"type": ttype, "title": "", "population": "Safety", "_from_area_rules": True})
                    existing_types.add(ttype)
                    added += 1

            if added:
                print(f"[TLF Integration v4] Added {added} tables from {ta} area rules")

            # Area-specific notes are available in area_rules["notes"] for LLM context
            # but are not used by the deterministic engine — they're for the hybrid path.

    # ── BUILD STUDY_INFO ──
    study_info = {
        "arm_names": [a["name"] for a in arms],
        "populations": [p["name"] for p in populations],
        "study_periods": facts.get("study_periods", []),
        "stratification_factors": facts.get("stratification_factors", []),
        "therapeutic_area": facts.get("therapeutic_area", ""),
        "design_type": facts.get("study_design", {}).get("type", "")
    }

    print(f"[TLF Integration v4] Built TLF list (YAML engine): {len(tables)} tables, {len(figures)} figures, {len(listings)} listings")

    return {
        "study_info": study_info,
        "tables": tables,
        "figures": figures,
        "listings": listings
    }


def _build_tlf_list_hardcoded(facts: Dict) -> Dict:
    """
    Build the complete TLF list deterministically from extracted facts (hardcoded fallback).
    Returns a dict with study_info, tables, figures, listings — same structure
    as the old Claude response, ready for _format_table() processing.
    """
    tables = []
    figures = []
    listings = []

    arms = facts.get("arms", [])
    endpoints = facts.get("endpoints", [])
    aesis = facts.get("aesis", [])
    periods = facts.get("treatment_periods", [])
    assessments = facts.get("assessments_collected", {})
    populations = facts.get("populations", [])
    subgroups = facts.get("subgroups", [])

    # ── MANDATORY TABLES (every SAP) ──
    tables.append({"type": "disposition", "title": "Subject Disposition", "population": "All Screened"})
    tables.append({"type": "demographics", "title": "Demographics and Baseline Characteristics",
                   "population": "ITT", "extra_rows": facts.get("disease_specific_baseline", [])})
    tables.append({"type": "medical_history", "title": "Medical History", "population": "ITT"})
    tables.append({"type": "exposure", "title": "Study Drug Exposure", "population": "Safety"})
    tables.append({"type": "concomitant_medications", "title": "Concomitant Medications", "population": "Safety"})
    tables.append({"type": "ae_overview", "title": "Overview of TEAEs", "population": "Safety"})
    tables.append({"type": "ae_by_soc_pt", "title": "TEAEs by SOC and PT", "population": "Safety"})

    # ── ONE TABLE PER ENDPOINT x POPULATION x REVIEW ──
    for ep in endpoints:
        # Fill in analysis_method default if protocol didn't specify one
        if not ep.get("analysis_method"):
            ep["analysis_method"] = _default_analysis_method(ep.get("type"), facts)

        for pop in ep.get("populations", ["ITT"]) or ["ITT"]:
            reviews = ep.get("reviews") or [None]
            for review in reviews:
                title_parts = [ep["name"]]
                if review:
                    title_parts.append(f"({review.title()} Review)")
                title = " ".join(title_parts)

                tables.append({
                    "type": "time_to_event" if ep["type"] == "time_to_event" else "binary_response",
                    "title": f"{title} ({pop} Population)",
                    "population": pop,
                    "endpoint": ep["name"],
                    "analysis_method": ep.get("analysis_method"),
                    "footnotes": _build_endpoint_footnotes(ep, facts),
                    "extra_rows": ep.get("extra_rows", [])
                })

    # ── ONE TABLE PER AESI ──
    for aesi in aesis:
        tables.append({
            "type": "ae_special",
            "title": f"AESI - {aesi['name']}",
            "population": "Safety",
            "extra_rows": [aesi.get("definition", "")]
        })

    # ── ASSESSMENT-BASED TABLES ──
    if assessments.get("labs"):
        tables.append({"type": "labs", "title": "Laboratory Parameters - Summary Statistics", "population": "Safety"})
        tables.append({"type": "lab_shift", "title": "Laboratory Parameters - Shift Tables", "population": "Safety"})
        tables.append({"type": "labs", "title": "Laboratory Parameters - CTCAE Grade Summary", "population": "Safety"})
        tables.append({"type": "lab_shift", "title": "Laboratory Parameters - CTCAE Grade Shift", "population": "Safety"})
    if assessments.get("vitals"):
        tables.append({"type": "vitals", "title": "Vital Signs", "population": "Safety"})
    if assessments.get("ecg"):
        tables.append({"type": "ecg", "title": "Electrocardiogram Parameters", "population": "Safety"})
    if assessments.get("pk"):
        tables.append({"type": "pk_concentration", "title": "PK - Trough Concentrations", "population": "PK"})
    if assessments.get("immunogenicity"):
        tables.append({"type": "immunogenicity", "title": "Immunogenicity Assessment", "population": "Safety"})
    qol_instruments = assessments.get("qol")
    if qol_instruments and isinstance(qol_instruments, list):
        for instrument in qol_instruments:
            tables.append({"type": "pro_qol", "title": f"Quality of Life - {instrument}", "population": "ITT"})
    if assessments.get("physical_exam"):
        tables.append({"type": "labs", "title": "Physical Examination Shift", "population": "Safety"})
    if assessments.get("ecog_ps"):
        tables.append({"type": "labs", "title": "ECOG Performance Status by Visit", "population": "Safety"})
    if assessments.get("pregnancy_test"):
        tables.append({"type": "labs", "title": "Pregnancy Test Summary", "population": "Safety"})
    if assessments.get("viral_serology"):
        tables.append({"type": "labs", "title": "Viral Serology", "population": "ITT"})
    if assessments.get("gene_screening"):
        tables.append({"type": "demographics", "title": "Gene Screening", "population": "ITT"})

    # ── BACKBONE THERAPY EXPOSURE ──
    for therapy in facts.get("backbone_therapies", []):
        drug_name = therapy.split(" ")[0]  # e.g. "Paclitaxel"
        tables.append({"type": "exposure", "title": f"{drug_name} Exposure", "population": "Safety"})

    # ── ADDITIONAL SAFETY TABLES ──
    tables.append({"type": "ae_by_soc_pt", "title": "TEAEs Grade \u22653", "population": "Safety"})
    tables.append({"type": "ae_by_soc_pt", "title": "Serious TEAEs", "population": "Safety"})
    tables.append({"type": "ae_by_soc_pt", "title": "TEAEs Leading to Discontinuation", "population": "Safety"})
    tables.append({"type": "ae_by_soc_pt", "title": "TEAEs Leading to Death", "population": "Safety"})
    tables.append({"type": "ae_by_soc_pt", "title": "TEAEs with PT \u22655% in Either Group", "population": "Safety"})

    # ── PERIOD SPLITS ──
    if len(periods) > 1:
        for period in periods:
            tables.append({"type": "ae_by_soc_pt", "title": f"TEAEs by SOC and PT - {period}", "population": "Safety"})
        for period in periods:
            tables.append({"type": "concomitant_medications", "title": f"Concomitant Medications - {period}", "population": "Safety"})

    # ── FIGURES ──
    for ep in endpoints:
        if ep["type"] == "time_to_event":
            figures.append({"type": "km_plot", "title": f"KM Plot of {ep['name']}", "endpoint": ep["name"]})
    if subgroups:
        for ep in endpoints:
            if ep.get("primary") or ep["type"] == "time_to_event":
                figures.append({"type": "forest_plot", "title": f"Forest Plot of {ep['name']} by Subgroup", "endpoint": ep["name"]})

    # ── STANDARD LISTINGS (every SAP has these) ──
    standard_listings = [
        ("Deaths", "Safety", ["Subject ID", "Age", "Sex", "Treatment", "Date of Death", "Cause of Death"]),
        ("Serious Adverse Events", "Safety", ["Subject ID", "Age", "Sex", "Treatment", "SOC", "PT", "Severity", "Relationship", "Outcome", "Start Date", "End Date"]),
        ("TEAEs Leading to Discontinuation", "Safety", ["Subject ID", "Treatment", "SOC", "PT", "Severity", "Relationship", "Action Taken"]),
        ("TEAEs Leading to Death", "Safety", ["Subject ID", "Treatment", "SOC", "PT", "Relationship", "Date of Death"]),
        ("Adverse Events of Special Interest", "Safety", ["Subject ID", "Treatment", "AESI Category", "PT", "Severity", "Relationship", "Outcome"]),
        ("Protocol Deviations", "ITT", ["Subject ID", "Treatment", "Category", "Description", "Date", "Major/Minor"]),
        ("Screening Failures", "All Screened", ["Subject ID", "Reason for Screen Failure"]),
        ("Subject Disposition", "ITT", ["Subject ID", "Treatment", "Status", "Reason for Discontinuation"]),
        ("Demographics", "ITT", ["Subject ID", "Age", "Sex", "Race", "Ethnicity", "Weight", "Height"]),
        ("Medical History", "ITT", ["Subject ID", "SOC", "PT", "Ongoing"]),
        ("Disease Characteristics", "ITT", ["Subject ID", "Treatment", "Disease Status", "Stage", "Histology"]),
        ("Prior Medications", "Safety", ["Subject ID", "Treatment", "ATC Class", "Drug Name", "Start Date", "Stop Date"]),
        ("Concomitant Medications", "Safety", ["Subject ID", "Treatment", "ATC Class", "Drug Name", "Start Date", "Stop Date"]),
        ("Study Drug Exposure", "Safety", ["Subject ID", "Treatment", "Cycle", "Dose", "Date"]),
        ("All Adverse Events", "Safety", ["Subject ID", "Treatment", "SOC", "PT", "Severity", "Relationship", "Serious", "Start Date", "End Date", "Outcome"]),
        ("Efficacy Data", "ITT", ["Subject ID", "Treatment", "Visit", "Assessment", "Response"]),
    ]
    if assessments.get("labs"):
        standard_listings += [
            ("Clinical Chemistry", "Safety", ["Subject ID", "Treatment", "Parameter", "Visit", "Result", "Flag", "CTCAE Grade"]),
            ("Hematology", "Safety", ["Subject ID", "Treatment", "Parameter", "Visit", "Result", "Flag", "CTCAE Grade"]),
            ("Urinalysis", "Safety", ["Subject ID", "Treatment", "Parameter", "Visit", "Result", "Flag"]),
        ]
    if assessments.get("vitals"):
        standard_listings.append(("Vital Signs", "Safety", ["Subject ID", "Treatment", "Parameter", "Visit", "Result"]))
    if assessments.get("ecg"):
        standard_listings.append(("ECG", "Safety", ["Subject ID", "Treatment", "Visit", "Interpretation"]))
    if assessments.get("pk"):
        standard_listings += [
            ("PK Concentrations", "PK", ["Subject ID", "Treatment", "Visit", "Timepoint", "Concentration"]),
            ("PK Parameters", "PK", ["Subject ID", "Treatment", "Parameter", "Value"]),
        ]
    if assessments.get("immunogenicity"):
        standard_listings.append(("Immunogenicity", "Safety", ["Subject ID", "Treatment", "Visit", "ADA Status", "Titer", "NAb Status"]))
    if qol_instruments and isinstance(qol_instruments, list):
        for instrument in qol_instruments:
            standard_listings.append((f"{instrument} Responses", "ITT", ["Subject ID", "Treatment", "Visit", "Scale", "Score"]))
    if assessments.get("physical_exam"):
        standard_listings.append(("Physical Examination", "Safety", ["Subject ID", "Treatment", "Visit", "Body System", "Finding"]))
    if assessments.get("pregnancy_test"):
        standard_listings.append(("Pregnancy Test", "Safety", ["Subject ID", "Treatment", "Visit", "Type", "Result"]))
    if assessments.get("viral_serology"):
        standard_listings.append(("Viral Serology", "ITT", ["Subject ID", "Treatment", "Visit", "Parameter", "Result"]))

    # Backbone therapy listings
    for therapy in facts.get("backbone_therapies", []):
        drug_name = therapy.split(" ")[0]
        standard_listings.append((f"{drug_name} Exposure", "Safety", ["Subject ID", "Treatment", "Cycle", "Dose", "Date"]))

    for title, pop, variables in standard_listings:
        listings.append({"type": "listing", "title": f"Listing of {title}", "population": pop, "variables": variables})

    # ── BUILD STUDY_INFO ──
    study_info = {
        "arm_names": [a["name"] for a in arms],
        "populations": [p["name"] for p in populations],
        "study_periods": facts.get("study_periods", []),
        "stratification_factors": facts.get("stratification_factors", []),
        "therapeutic_area": facts.get("therapeutic_area", ""),
        "design_type": facts.get("study_design", {}).get("type", "")
    }

    print(f"[TLF Integration v4] Built TLF list: {len(tables)} tables, {len(figures)} figures, {len(listings)} listings")

    return {
        "study_info": study_info,
        "tables": tables,
        "figures": figures,
        "listings": listings
    }


# =============================================================================
# STEP 3: YAML FORMATS EACH TABLE (unchanged from v3)
# =============================================================================
def _get_placeholder(fmt: str) -> str:
    """Convert a format name to its xxx placeholder string."""
    base_fmt, _, _ = _load_yaml_configs()
    placeholders = base_fmt.get("placeholder_formats", {})
    return placeholders.get(fmt, "xxx")


def _format_table(table_spec: Dict, study_info: Dict) -> Dict:
    """
    Merge a single table spec from the builder with YAML formatting config.
    Returns a dict ready for rendering.
    """
    base_fmt, analysis_templates, _ = _load_yaml_configs()
    method_index = _build_method_index()

    table_type = table_spec.get("type", "")
    analysis_method = table_spec.get("analysis_method", "")
    population = table_spec.get("population", "ITT")
    title = table_spec.get("title", table_type.replace("_", " ").title())
    endpoint = table_spec.get("endpoint", "")
    extra_rows = table_spec.get("extra_rows", [])

    # ── Determine rows ──────────────────────────────────────────
    rows = []

    # Try analysis_templates first (detailed row structure with formatting)
    if analysis_method and analysis_method in method_index:
        category, method_key = method_index[analysis_method]
        template = analysis_templates.get(category, {}).get(method_key, {})
        template_rows = template.get("rows", [])
        for r in template_rows:
            if isinstance(r, dict):
                rows.append({
                    "label": r.get("label", ""),
                    "format": r.get("format", ""),
                    "indent": r.get("indent", 0),
                    "type": r.get("type", "data"),
                    "bold": r.get("bold", False),
                })
            elif isinstance(r, str):
                rows.append({"label": r, "format": "", "indent": 0, "type": "data", "bold": False})

    # Fall back to base_formatting standard_rows if no analysis template rows
    if not rows:
        # Check if this table_type has an analysis_templates entry (e.g., ae_overview → safety.ae_overview)
        for cat, type_set in _TEMPLATE_CATEGORY_TO_TYPE.items():
            if table_type in type_set and cat in analysis_templates:
                template = analysis_templates[cat].get(table_type, {})
                template_rows = template.get("rows", [])
                for r in template_rows:
                    if isinstance(r, dict):
                        rows.append({
                            "label": r.get("label", ""),
                            "format": r.get("format", ""),
                            "indent": r.get("indent", 0),
                            "type": r.get("type", "data"),
                            "bold": r.get("bold", False),
                        })
                if rows:
                    break

    # Final fallback: base_formatting standard_rows (simple string lists)
    if not rows:
        standard = base_fmt.get("standard_rows", {}).get(table_type, [])
        for item in standard:
            if isinstance(item, str):
                indent = len(item) - len(item.lstrip())
                indent_level = indent // 2
                rows.append({
                    "label": item.strip(),
                    "format": "count_pct" if indent_level > 0 else "",
                    "indent": indent_level,
                    "type": "data",
                    "bold": False,
                })
            elif isinstance(item, dict):
                rows.append({
                    "label": item.get("variable", item.get("label", "")),
                    "format": item.get("format", ""),
                    "indent": 0,
                    "type": item.get("type", "data"),
                    "bold": False,
                })

    # Append extra rows from the builder
    for extra in extra_rows:
        if isinstance(extra, str):
            rows.append({"label": extra, "format": "count_pct", "indent": 0, "type": "data", "bold": False})
        elif isinstance(extra, dict):
            rows.append({
                "label": extra.get("label", ""),
                "format": extra.get("format", "count_pct"),
                "indent": extra.get("indent", 0),
                "type": "data",
                "bold": False,
            })

    # ── Determine columns ────────────────────────────────────────
    arm_names = study_info.get("arm_names", ["Treatment", "Control"])
    columns = _build_columns(arm_names, table_type)

    # ── Determine footnotes ──────────────────────────────────────
    footnotes = _get_footnotes(table_type, analysis_method, population, arm_names)

    # Add study-specific footnotes from the builder (endpoint footnotes)
    for fn in table_spec.get("footnotes", []):
        fn_text = fn if isinstance(fn, str) else str(fn)
        if fn_text not in footnotes:
            footnotes.append(fn_text)

    # Add any method-specific footnotes from analysis_templates
    if analysis_method and analysis_method in method_index:
        category, method_key = method_index[analysis_method]
        template = analysis_templates.get(category, {}).get(method_key, {})
        for fn in template.get("footnotes", []):
            fn_text = fn if isinstance(fn, str) else str(fn)
            # Substitute arm name placeholders
            if arm_names:
                fn_text = fn_text.replace("{arm1}", arm_names[0])
                if len(arm_names) > 1:
                    fn_text = fn_text.replace("{arm2}", arm_names[1])
            if fn_text not in footnotes:
                footnotes.append(fn_text)

    # ── Determine source dataset ─────────────────────────────────
    source = ""
    # From analysis_templates
    if analysis_method and analysis_method in method_index:
        category, method_key = method_index[analysis_method]
        template = analysis_templates.get(category, {}).get(method_key, {})
        source = template.get("adam_dataset", "")
    # From base_formatting adam_datasets
    if not source:
        for ds_key, ds_info in base_fmt.get("adam_datasets", {}).items():
            if table_type in ds_info.get("used_for", []):
                source = ds_info.get("name", ds_key.upper())
                break
    if not source:
        source = "ADSL"

    # ── Orientation ──────────────────────────────────────────────
    orientation = base_fmt.get("orientation_defaults", {}).get(table_type, "PORTRAIT")

    # ── Section for numbering ────────────────────────────────────
    # Prefer YAML mapping, fall back to hardcoded Python dict
    yaml_section_map = (_GENERATION_RULES or {}).get("table_type_to_section", {})
    section = yaml_section_map.get(table_type) or _TABLE_TYPE_TO_SECTION.get(table_type, "14.3")

    return {
        "type": table_type,
        "section": section,
        "number": "",  # assigned later by _assign_table_numbers
        "title": title,
        "population": population,
        "endpoint": endpoint,
        "columns": columns,
        "rows": rows,
        "footnotes": footnotes,
        "source": source,
        "orientation": orientation,
    }


def _build_columns(arm_names: List[str], table_type: str) -> List[Dict]:
    """Build column definitions from arm names + base_formatting column_templates."""
    base_fmt, _, _ = _load_yaml_configs()
    col_templates = base_fmt.get("column_templates", {})

    num_arms = len(arm_names)

    if table_type in ("lab_shift",):
        template_key = "shift_table"
    elif table_type in _DESCRIPTIVE_TYPES:
        template_key = "descriptive"
    elif num_arms == 0 or num_arms == 1:
        template_key = "single_arm"
    elif num_arms == 2:
        template_key = "two_arm_total"
    elif num_arms >= 3:
        template_key = "three_arm_total"
    else:
        template_key = "two_arm_total"

    template = col_templates.get(template_key, [])

    # Build columns, substituting arm names
    columns = []
    arm_idx = 0
    for col in template:
        header = col.get("header", "")
        # Replace {ARM1}, {ARM2}, {ARM3} placeholders
        for i in range(min(num_arms, 4)):
            placeholder = f"{{ARM{i+1}}}"
            if placeholder in header and arm_idx < num_arms:
                header = header.replace(placeholder, arm_names[i])

        columns.append({
            "header": header,
            "width": col.get("width", 1.5),
            "align": col.get("alignment", "C"),
            "format": col.get("format", ""),
        })

    # If template didn't match, build simple columns
    if not columns:
        columns.append({"header": "Parameter", "width": 3.0, "align": "L", "format": ""})
        for arm in arm_names:
            columns.append({"header": f"{arm}\n(N=xxx)", "width": 1.5, "align": "C", "format": ""})
        if num_arms >= 2:
            columns.append({"header": "Total\n(N=xxx)", "width": 1.5, "align": "C", "format": ""})

    return columns


def _get_footnotes(table_type: str, analysis_method: str,
                   population: str, arm_names: List[str]) -> List[str]:
    """Assemble footnotes from base_formatting configs."""
    base_fmt, _, _ = _load_yaml_configs()
    footnotes = []

    # Population footnote
    populations = base_fmt.get("populations", {})
    pop_key = population.lower().replace(" ", "_").replace("-", "_")
    # Try exact match, then fuzzy
    pop_config = populations.get(pop_key)
    if not pop_config:
        for pk, pv in populations.items():
            if population.lower() in pk.lower() or pk.lower() in population.lower():
                pop_config = pv
                break
    if pop_config:
        footnotes.append(pop_config.get("footnote", f"{population} Population."))
    else:
        footnotes.append(f"{population} Population.")

    # Percentage denomination footnote
    footnotes.append(base_fmt.get("footnote_templates", {}).get("general", {}).get("percentage_denom", "Percentages based on N in column header."))

    # AE-specific footnotes
    if table_type in ("ae_overview", "ae_by_soc_pt", "ae_special", "aesi"):
        coding = base_fmt.get("footnote_templates", {}).get("coding", {})
        if "meddra" in coding:
            footnotes.append(coding["meddra"].replace("{VERSION}", "XX.X"))
        defs = base_fmt.get("footnote_templates", {}).get("definitions", {})
        if "teae" in defs:
            footnotes.append(defs["teae"])

    # Subject count footnote for safety tables
    if table_type in ("ae_overview", "ae_by_soc_pt", "ae_special", "aesi"):
        general = base_fmt.get("footnote_templates", {}).get("general", {})
        if "subject_count" in general:
            footnotes.append(general["subject_count"])

    return footnotes


def _assign_table_numbers(tables: List[Dict]) -> List[Dict]:
    """Assign sequential table numbers within each ICH section."""
    section_counters = {}
    for table in tables:
        section = table.get("section", "14.3")
        if section not in section_counters:
            section_counters[section] = 0
        section_counters[section] += 1
        table["number"] = f"{section}.{section_counters[section]}"
    return tables


# =============================================================================
# RENDER TO MARKDOWN (unchanged from v3)
# =============================================================================
def _render_shells(formatted_tables: List[Dict], figures: List[Dict],
                   listings: List[Dict], study_info: Dict) -> str:
    """Convert formatted table/figure/listing dicts into markdown string."""
    sections = []

    # Header
    sections.append("## TABLE AND FIGURE SHELLS\n")

    # Tables
    if formatted_tables:
        for table in formatted_tables:
            sections.append(_render_single_table(table))

    # Figures
    if figures:
        sections.append("\n---\n\n### FIGURE SHELLS\n")
        fig_counter = 0
        for fig in figures:
            fig_counter += 1
            sections.append(_render_figure(fig, fig_counter))

    # Listings
    if listings:
        sections.append("\n---\n\n### LISTING SHELLS\n")
        listing_counter = 0
        for lst in listings:
            listing_counter += 1
            sections.append(_render_listing(lst, listing_counter, study_info))

    # Summary
    sections.append(_render_summary(formatted_tables, figures, listings, study_info))

    return "\n\n".join(sections)


def _placeholder_for_format(fmt: str) -> str:
    """Return the xxx-style placeholder for a given format type."""
    mapping = {
        "count": "xxx",
        "count_pct": "xxx (xx.x)",
        "percentage": "xx.x",
        "mean": "xx.x",
        "mean_sd": "xx.x (xx.xx)",
        "median": "xx.x",
        "median_ci": "xx.x [xx.x, xx.x]",
        "min_max": "xx, xx",
        "ci_95": "[xx.x, xx.x]",
        "hazard_ratio": "x.xx [xx.x, xx.x]",
        "p_value": "x.xxxx",
        "rate_ratio": "x.xx [xx.x, xx.x]",
        "or_ci": "x.xx [xx.x, xx.x]",
        "events_rate": "xxx (xx.x)",
        "n_pct": "xxx (xx.x)",
    }
    return mapping.get(fmt, "xxx")


def _render_single_table(table: Dict) -> str:
    """Render one formatted table dict as markdown."""
    number = table.get("number", "")
    title = table.get("title", "")
    population = table.get("population", "")
    source = table.get("source", "")
    orientation = table.get("orientation", "PORTRAIT")
    columns = table.get("columns", [])
    rows = table.get("rows", [])
    footnotes = table.get("footnotes", [])

    lines = []

    # Title
    pop_suffix = f" ({population} Population)" if population and population.lower() not in title.lower() else ""
    lines.append(f"### Table {number}: {title}{pop_suffix}")
    lines.append("")

    # Metadata
    lines.append(f"**Source Dataset:** {source} | **Population:** {population} | **Orientation:** {orientation}")
    lines.append("")

    # Table header
    if columns:
        headers = [c.get("header", "").replace("\n", " ") for c in columns]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "|".join(["---" for _ in columns]) + "|")
    else:
        lines.append("| Parameter | Value |")
        lines.append("|-----------|-------|")

    # Table rows
    num_data_cols = max(len(columns) - 1, 1)  # exclude first "Parameter" column
    for row in rows:
        label = row.get("label", "")
        row_type = row.get("type", "data")
        fmt = row.get("format", "")
        indent = row.get("indent", 0)
        bold = row.get("bold", False)

        if row_type == "spacer" or not label:
            # Empty spacer row
            lines.append("| " + " | ".join(["" for _ in columns]) + " |")
            continue

        # Apply indent
        display_label = ("  " * indent) + label
        if bold:
            display_label = f"**{display_label}**"

        if row_type == "header":
            # Section header row — label spans, data cells empty
            cells = [display_label] + ["" for _ in range(num_data_cols)]
        else:
            # Data row with placeholders
            placeholder = _placeholder_for_format(fmt)
            cells = [display_label] + [placeholder for _ in range(num_data_cols)]

        lines.append("| " + " | ".join(cells) + " |")

    lines.append("")

    # Footnotes
    if footnotes:
        lines.append("**Footnotes:**")
        for i, fn in enumerate(footnotes, 1):
            lines.append(f"{i}. {fn}")
        lines.append("")

    return "\n".join(lines)


def _render_figure(fig: Dict, counter: int) -> str:
    """Render a figure shell."""
    fig_type = fig.get("type", "figure")
    title = fig.get("title", f"Figure {counter}")
    endpoint = fig.get("endpoint", "")

    lines = [
        f"### Figure {counter}: {title}",
        "",
        f"**Type:** {fig_type.replace('_', ' ').title()}",
    ]
    if endpoint:
        lines.append(f"**Endpoint:** {endpoint}")

    if "km" in fig_type.lower():
        lines.extend([
            "",
            "```",
            "  1.0 |--*--------",
            "      |   \\",
            "  0.8 |    *------",
            "      |          \\",
            "  0.6 |           *---",
            "      |",
            "  0.4 |",
            "      |",
            "  0.2 |",
            "      |",
            "  0.0 |________________",
            "      0   6  12  18  24  (months)",
            "",
            "  --- Arm 1    ─── Arm 2",
            "```",
            "",
            "Number at risk table included below figure.",
        ])
    elif "forest" in fig_type.lower():
        lines.extend([
            "",
            "```",
            "  Subgroup          n    HR [95% CI]     Favors Arm1  Favors Arm2",
            "  ──────────────────────────────────────────────────────────────",
            "  Overall           xxx  x.xx [x.xx, x.xx]     |----*----|",
            "  Age <65           xxx  x.xx [x.xx, x.xx]   |---*---|",
            "  Age >=65          xxx  x.xx [x.xx, x.xx]     |------*------|",
            "  Male              xxx  x.xx [x.xx, x.xx]    |----*----|",
            "  Female            xxx  x.xx [x.xx, x.xx]   |---*---|",
            "                                           0.5    1.0    2.0",
            "```",
        ])
    else:
        lines.extend(["", "*[Figure placeholder — to be generated by statistical programming]*"])

    return "\n".join(lines)


def _render_listing(listing: Dict, counter: int, study_info: Dict) -> str:
    """Render a listing shell."""
    title = listing.get("title", f"Listing {counter}")
    population = listing.get("population", "Safety")
    variables = listing.get("variables", ["Subject ID", "Treatment", "Parameter", "Value"])

    lines = [
        f"### Listing 16.2.{counter}: {title} ({population} Population)",
        "",
        f"**Population:** {population} | **Orientation:** LANDSCAPE",
        "",
    ]

    # Render as table header
    lines.append("| " + " | ".join(variables) + " |")
    lines.append("|" + "|".join(["---" for _ in variables]) + "|")
    lines.append("| " + " | ".join(["xxx" for _ in variables]) + " |")
    lines.append("")

    return "\n".join(lines)


def _render_summary(tables: List[Dict], figures: List[Dict],
                    listings: List[Dict], study_info: Dict) -> str:
    """Render summary section with table counts."""
    lines = [
        "---",
        "",
        "### TLF Shell Summary",
        "",
        f"**Therapeutic Area:** {study_info.get('therapeutic_area', 'Not specified')}",
        f"**Treatment Arms:** {', '.join(study_info.get('arm_names', []))}",
        f"**Analysis Populations:** {', '.join(study_info.get('populations', []))}",
        "",
        "| Category | Count |",
        "|----------|-------|",
    ]

    # Count by ICH section
    section_counts = {}
    for t in tables:
        sec = t.get("section", "14.3")
        cat = {
            "14.1": "Disposition & Demographics",
            "14.2": "Efficacy",
            "14.3": "Safety",
            "14.4": "PK & Immunogenicity",
            "16.2": "Listings",
        }.get(sec, sec)
        section_counts[cat] = section_counts.get(cat, 0) + 1

    for cat, count in sorted(section_counts.items()):
        lines.append(f"| {cat} | {count} |")

    lines.append(f"| Figures | {len(figures)} |")
    lines.append(f"| Listings | {len(listings)} |")

    total = len(tables) + len(figures) + len(listings)
    lines.append(f"| **Total TLFs** | **{total}** |")
    lines.append("")

    return "\n".join(lines)


# =============================================================================
# RENDER TO DOCX (Word document shells)
# =============================================================================
def _render_shells_docx(formatted_tables: List[Dict], figures: List[Dict],
                        listings: List[Dict], study_info: Dict) -> bytes:
    """Convert formatted table/figure/listing dicts into a Word (.docx) document."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # --- Page setup: Letter, 1-inch margins ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    # --- Header with protocol ID + CONFIDENTIAL ---
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{study_info.get('protocol_id', 'PROTOCOL')}    CONFIDENTIAL"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.style.font.size = Pt(8)
    hp.style.font.color.rgb = RGBColor(128, 128, 128)

    # --- Helper: shade a cell ---
    def _shade_cell(cell, color_hex="1F4E79"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # --- Helper: set cell text with formatting ---
    def _set_cell(cell, text, bold=False, size=Pt(10), color=None, indent=0):
        cell.text = ""
        p = cell.paragraphs[0]
        prefix = "  " * indent
        run = p.add_run(prefix + str(text))
        run.font.name = "Arial"
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # --- Helper: add landscape section break ---
    def _add_landscape_section(doc):
        new_section = doc.add_section(2)  # WD_SECTION.NEW_PAGE = 2
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = Inches(11)
        new_section.page_height = Inches(8.5)
        new_section.top_margin = Inches(0.75)
        new_section.bottom_margin = Inches(0.75)
        new_section.left_margin = Inches(1)
        new_section.right_margin = Inches(1)
        return new_section

    # --- Helper: add portrait section break ---
    def _add_portrait_section(doc):
        new_section = doc.add_section(2)
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = Inches(8.5)
        new_section.page_height = Inches(11)
        new_section.top_margin = Inches(1)
        new_section.bottom_margin = Inches(0.75)
        new_section.left_margin = Inches(1)
        new_section.right_margin = Inches(1)
        return new_section

    HEADER_BG = "1F4E79"
    HEADER_TEXT = RGBColor(255, 255, 255)

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    doc.add_paragraph("")
    doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("TABLE, LISTING, AND FIGURE SHELLS")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Arial"

    doc.add_paragraph("")
    for label, key in [
        ("Protocol", "protocol_id"),
        ("Therapeutic Area", "therapeutic_area"),
        ("Treatment Arms", "arm_names"),
        ("Analysis Populations", "populations"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val = study_info.get(key, "")
        if isinstance(val, list):
            val = ", ".join(val)
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "Arial"
        run = p.add_run(str(val))
        run.font.size = Pt(12)
        run.font.name = "Arial"

    doc.add_paragraph("")
    summary_p = doc.add_paragraph()
    summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_p.add_run(
        f"{len(formatted_tables)} Tables  |  {len(figures)} Figures  |  {len(listings)} Listings  |  "
        f"{len(formatted_tables) + len(figures) + len(listings)} Total"
    )
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # =====================================================================
    # TABLES
    # =====================================================================
    for table_dict in formatted_tables:
        doc.add_page_break()

        number = table_dict.get("number", "")
        title = table_dict.get("title", "")
        population = table_dict.get("population", "")
        source = table_dict.get("source", "")
        orientation = table_dict.get("orientation", "PORTRAIT")
        columns = table_dict.get("columns", [])
        rows = table_dict.get("rows", [])
        footnotes = table_dict.get("footnotes", [])

        # Title
        pop_suffix = f" ({population} Population)" if population and population.lower() not in title.lower() else ""
        heading = doc.add_paragraph()
        run = heading.add_run(f"Table {number}: {title}{pop_suffix}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Arial"

        # Metadata line
        meta = doc.add_paragraph()
        run = meta.add_run(f"Source Dataset: {source}  |  Population: {population}  |  Orientation: {orientation}")
        run.font.size = Pt(8)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(100, 100, 100)

        # Build table
        if not columns:
            columns = [{"header": "Parameter"}, {"header": "Value"}]

        num_cols = len(columns)
        num_data_cols = max(num_cols - 1, 1)

        # Create docx table: header row + data rows
        tbl = doc.add_table(rows=1 + max(len(rows), 1), cols=num_cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr_row = tbl.rows[0]
        for i, col in enumerate(columns):
            cell = hdr_row.cells[i]
            _set_cell(cell, col.get("header", "").replace("\n", " "), bold=True, size=Pt(9), color=HEADER_TEXT)
            _shade_cell(cell, HEADER_BG)

        # Data rows
        for r_idx, row in enumerate(rows):
            if r_idx + 1 >= len(tbl.rows):
                break

            doc_row = tbl.rows[r_idx + 1]
            label = row.get("label", "")
            row_type = row.get("type", "data")
            fmt = row.get("format", "")
            indent = row.get("indent", 0)
            bold = row.get("bold", False)

            if row_type == "spacer" or not label:
                continue

            display_label = label
            _set_cell(doc_row.cells[0], display_label, bold=bold or row_type == "header", size=Pt(9), indent=indent)

            if row_type == "header":
                # Section header — empty data cells
                for ci in range(1, num_cols):
                    _set_cell(doc_row.cells[ci], "", size=Pt(9))
            else:
                # Data row with placeholders
                placeholder = _placeholder_for_format(fmt)
                for ci in range(1, num_cols):
                    _set_cell(doc_row.cells[ci], placeholder, size=Pt(9))

        # Footnotes
        if footnotes:
            fn_p = doc.add_paragraph()
            for i, fn in enumerate(footnotes, 1):
                run = fn_p.add_run(f"{i}. {fn}\n")
                run.font.size = Pt(8)
                run.font.name = "Arial"
                run.italic = True

    # =====================================================================
    # FIGURES
    # =====================================================================
    if figures:
        doc.add_page_break()
        heading = doc.add_paragraph()
        run = heading.add_run("FIGURE SHELLS")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"

        for fig_counter, fig in enumerate(figures, 1):
            doc.add_page_break()

            fig_type = fig.get("type", "figure")
            title = fig.get("title", f"Figure {fig_counter}")
            endpoint = fig.get("endpoint", "")

            # Title
            fp = doc.add_paragraph()
            run = fp.add_run(f"Figure {fig_counter}: {title}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"

            # Type
            tp = doc.add_paragraph()
            run = tp.add_run(f"Type: {fig_type.replace('_', ' ').title()}")
            run.font.size = Pt(9)
            run.font.name = "Arial"

            if endpoint:
                ep = doc.add_paragraph()
                run = ep.add_run(f"Endpoint: {endpoint}")
                run.font.size = Pt(9)
                run.font.name = "Arial"

            doc.add_paragraph("")

            if "km" in fig_type.lower():
                # KM plot placeholder with number-at-risk table
                box = doc.add_paragraph()
                box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = box.add_run("[Kaplan-Meier survival curve — to be generated by statistical programming]")
                run.font.size = Pt(10)
                run.italic = True

                doc.add_paragraph("")
                nar = doc.add_paragraph()
                run = nar.add_run("Number at Risk")
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"

                # Timepoint headers
                arms = study_info.get("arm_names", ["Arm 1", "Arm 2"])
                timepoints = ["0", "6", "12", "18", "24"]
                risk_tbl = doc.add_table(rows=1 + len(arms), cols=1 + len(timepoints))
                risk_tbl.style = "Table Grid"
                _set_cell(risk_tbl.rows[0].cells[0], "Month", bold=True, size=Pt(8))
                for ti, tp_val in enumerate(timepoints):
                    _set_cell(risk_tbl.rows[0].cells[ti + 1], tp_val, bold=True, size=Pt(8))
                    _shade_cell(risk_tbl.rows[0].cells[ti + 1], HEADER_BG)
                _shade_cell(risk_tbl.rows[0].cells[0], HEADER_BG)
                for ai, arm in enumerate(arms):
                    _set_cell(risk_tbl.rows[ai + 1].cells[0], arm, size=Pt(8))
                    for ti in range(len(timepoints)):
                        _set_cell(risk_tbl.rows[ai + 1].cells[ti + 1], "xxx", size=Pt(8))

            elif "forest" in fig_type.lower():
                box = doc.add_paragraph()
                box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = box.add_run("[Forest plot — to be generated by statistical programming]")
                run.font.size = Pt(10)
                run.italic = True

                doc.add_paragraph("")
                # Subgroup placeholder table
                sg_tbl = doc.add_table(rows=6, cols=4)
                sg_tbl.style = "Table Grid"
                for ci, h in enumerate(["Subgroup", "n", "HR [95% CI]", "Favors"]):
                    _set_cell(sg_tbl.rows[0].cells[ci], h, bold=True, size=Pt(8), color=HEADER_TEXT)
                    _shade_cell(sg_tbl.rows[0].cells[ci], HEADER_BG)
                for ri, sg in enumerate(["Overall", "Age <65", "Age >=65", "Male", "Female"], 1):
                    _set_cell(sg_tbl.rows[ri].cells[0], sg, size=Pt(8))
                    _set_cell(sg_tbl.rows[ri].cells[1], "xxx", size=Pt(8))
                    _set_cell(sg_tbl.rows[ri].cells[2], "x.xx [x.xx, x.xx]", size=Pt(8))
                    _set_cell(sg_tbl.rows[ri].cells[3], "---|---*---|---", size=Pt(8))

            elif "waterfall" in fig_type.lower():
                box = doc.add_paragraph()
                box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = box.add_run("[Waterfall plot of best percentage change from baseline — to be generated by statistical programming]")
                run.font.size = Pt(10)
                run.italic = True

            elif "swimmer" in fig_type.lower():
                box = doc.add_paragraph()
                box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = box.add_run("[Swimmer plot of individual subject response duration — to be generated by statistical programming]")
                run.font.size = Pt(10)
                run.italic = True

            else:
                box = doc.add_paragraph()
                box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = box.add_run("[Figure placeholder — to be generated by statistical programming]")
                run.font.size = Pt(10)
                run.italic = True

    # =====================================================================
    # LISTINGS (landscape)
    # =====================================================================
    if listings:
        _add_landscape_section(doc)

        heading = doc.add_paragraph()
        run = heading.add_run("LISTING SHELLS")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Arial"

        for lst_counter, lst in enumerate(listings, 1):
            if lst_counter > 1:
                doc.add_page_break()

            title = lst.get("title", f"Listing {lst_counter}")
            population = lst.get("population", "Safety")
            variables = lst.get("variables", ["Subject ID", "Treatment", "Parameter", "Value"])

            # Title
            lp = doc.add_paragraph()
            run = lp.add_run(f"Listing 16.2.{lst_counter}: {title} ({population} Population)")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"

            # Metadata
            mp = doc.add_paragraph()
            run = mp.add_run(f"Population: {population}  |  Orientation: LANDSCAPE")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.name = "Arial"

            # Column table: header + 1 placeholder row
            num_cols = len(variables)
            lst_tbl = doc.add_table(rows=2, cols=num_cols)
            lst_tbl.style = "Table Grid"
            lst_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ci, var in enumerate(variables):
                _set_cell(lst_tbl.rows[0].cells[ci], var, bold=True, size=Pt(8), color=HEADER_TEXT)
                _shade_cell(lst_tbl.rows[0].cells[ci], HEADER_BG)
                _set_cell(lst_tbl.rows[1].cells[ci], "xxx", size=Pt(8))

    # =====================================================================
    # SUMMARY PAGE
    # =====================================================================
    _add_portrait_section(doc)

    heading = doc.add_paragraph()
    run = heading.add_run("TLF Shell Summary")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    doc.add_paragraph("")

    for label, key in [
        ("Therapeutic Area", "therapeutic_area"),
        ("Treatment Arms", "arm_names"),
        ("Analysis Populations", "populations"),
    ]:
        p = doc.add_paragraph()
        val = study_info.get(key, "")
        if isinstance(val, list):
            val = ", ".join(val)
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.name = "Arial"
        p.add_run(str(val)).font.name = "Arial"

    doc.add_paragraph("")

    # Count by ICH section
    section_counts = {}
    for t in formatted_tables:
        sec = t.get("section", "14.3")
        cat = {
            "14.1": "Disposition & Demographics",
            "14.2": "Efficacy",
            "14.3": "Safety",
            "14.4": "PK & Immunogenicity",
        }.get(sec, sec)
        section_counts[cat] = section_counts.get(cat, 0) + 1

    summary_tbl = doc.add_table(rows=2 + len(section_counts) + 1, cols=2)
    summary_tbl.style = "Table Grid"
    _set_cell(summary_tbl.rows[0].cells[0], "Category", bold=True, size=Pt(9), color=HEADER_TEXT)
    _set_cell(summary_tbl.rows[0].cells[1], "Count", bold=True, size=Pt(9), color=HEADER_TEXT)
    _shade_cell(summary_tbl.rows[0].cells[0], HEADER_BG)
    _shade_cell(summary_tbl.rows[0].cells[1], HEADER_BG)

    row_idx = 1
    for cat, count in sorted(section_counts.items()):
        _set_cell(summary_tbl.rows[row_idx].cells[0], cat, size=Pt(9))
        _set_cell(summary_tbl.rows[row_idx].cells[1], str(count), size=Pt(9))
        row_idx += 1

    _set_cell(summary_tbl.rows[row_idx].cells[0], "Figures", size=Pt(9))
    _set_cell(summary_tbl.rows[row_idx].cells[1], str(len(figures)), size=Pt(9))
    row_idx += 1
    _set_cell(summary_tbl.rows[row_idx].cells[0], "Listings", size=Pt(9))
    _set_cell(summary_tbl.rows[row_idx].cells[1], str(len(listings)), size=Pt(9))

    # Total
    total = len(formatted_tables) + len(figures) + len(listings)
    total_row = summary_tbl.add_row()
    _set_cell(total_row.cells[0], "Total TLFs", bold=True, size=Pt(9))
    _set_cell(total_row.cells[1], str(total), bold=True, size=Pt(9))

    # --- Serialize to bytes ---
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# PUBLIC API — SAME SIGNATURES AS CURRENT (main.py compatible)
# =============================================================================
def generate_tlf_shells_for_protocol(
    extraction: Dict[str, Any],
    study_id: Optional[str] = None,
    priority: Optional[int] = None,
    protocol_text: str = "",
    apply_universal_expansion: bool = True,
    output_format: str = "markdown"
) -> Union[str, bytes]:
    """
    Generate TLF shells for a protocol.

    v4: Claude extracts facts → code builds TLF list → YAML formats → render.

    Args:
        extraction: Protocol extraction dict from KG pipeline
        study_id: Optional study ID (unused in v4, kept for API compat)
        priority: Optional priority level (unused in v4, kept for API compat)
        protocol_text: Full protocol text
        apply_universal_expansion: Unused in v4 (kept for API compat)
        output_format: "markdown" (default) or "docx"

    Returns:
        Markdown string (if output_format="markdown") or .docx bytes (if output_format="docx")
    """
    try:
        # Step 1: Extract facts from the protocol via Claude API
        facts = _extract_protocol_facts(protocol_text, extraction)

        # Step 2: Build the full TLF list deterministically from facts
        tlf_result = _build_tlf_list(facts)
        study_info = tlf_result["study_info"]
        table_specs = tlf_result.get("tables", [])
        figure_specs = tlf_result.get("figures", [])
        listing_specs = tlf_result.get("listings", [])

        # Step 3: Format each table using YAML configs
        formatted_tables = []
        for spec in table_specs:
            try:
                formatted_tables.append(_format_table(spec, study_info))
            except Exception as e:
                print(f"[TLF Integration v4] Warning: Failed to format table '{spec.get('title', spec.get('type'))}': {e}")

        # Assign sequential table numbers
        formatted_tables = _assign_table_numbers(formatted_tables)

        # Step 4: Render
        if output_format == "docx":
            docx_bytes = _render_shells_docx(formatted_tables, figure_specs, listing_specs, study_info)
            print(f"[TLF Integration v4] Generated .docx: {len(formatted_tables)} tables, {len(figure_specs)} figures, {len(listing_specs)} listings ({len(docx_bytes):,} bytes)")
            return docx_bytes
        else:
            shell_text = _render_shells(formatted_tables, figure_specs, listing_specs, study_info)
            print(f"[TLF Integration v4] Generated {len(formatted_tables)} tables, {len(figure_specs)} figures, {len(listing_specs)} listings ({len(shell_text):,} chars)")
            return shell_text

    except Exception as e:
        print(f"[TLF Integration v4] Error generating shells: {e}")
        traceback.print_exc()
        if output_format == "docx":
            return b""
        return _generate_fallback_shells(extraction)


def _generate_fallback_shells(extraction: Dict[str, Any]) -> str:
    """Generate basic TLF shells when Claude API is unavailable."""
    arms = extraction.get("treatment_arms", []) if extraction else []
    if len(arms) >= 2:
        arm1 = arms[0] if isinstance(arms[0], str) else arms[0].get("arm_name", "Treatment")
        arm2 = arms[1] if isinstance(arms[1], str) else arms[1].get("arm_name", "Control")
    else:
        arm1 = "Treatment"
        arm2 = "Control"

    return f"""## TABLE AND FIGURE SHELLS

### Table 14.1.1: Subject Disposition

| Category | {arm1} (N=xxx) | {arm2} (N=xxx) | Total (N=xxx) |
|----------|----------------|----------------|---------------|
| Screened | xxx | xxx | xxx |
| Screen Failures | xxx | xxx | xxx |
| Randomized | xxx | xxx | xxx |
| Treated | xxx | xxx | xxx |
| Completed Treatment | xxx | xxx | xxx |
| Discontinued | xxx | xxx | xxx |

### Table 14.2.1: Primary Efficacy Analysis

| Parameter | {arm1} (N=xxx) | {arm2} (N=xxx) |
|-----------|----------------|----------------|
| Responders, n (%) | xxx (xx.x) | xxx (xx.x) |
| Non-responders, n (%) | xxx (xx.x) | xxx (xx.x) |
| Difference [95% CI] | xx.x [xx.x, xx.x] | -- |
| P-value | x.xxxx | -- |

### Table 14.3.1: Overview of Adverse Events

| Category | {arm1} (N=xxx) n (%) | {arm2} (N=xxx) n (%) |
|----------|----------------------|----------------------|
| Any TEAE | xxx (xx.x) | xxx (xx.x) |
| Treatment-related TEAE | xxx (xx.x) | xxx (xx.x) |
| Grade >=3 TEAE | xxx (xx.x) | xxx (xx.x) |
| Serious AE | xxx (xx.x) | xxx (xx.x) |
| TEAE Leading to D/C | xxx (xx.x) | xxx (xx.x) |
| Deaths | xxx (xx.x) | xxx (xx.x) |

*Note: Full TLF shells require ANTHROPIC_API_KEY to be set.*
"""


# =============================================================================
# BACKWARDS-COMPATIBLE STUBS (used by main.py for logging)
# =============================================================================
def detect_study_type(extraction: Dict[str, Any]) -> str:
    """Stub for backwards compatibility. Returns 'protocol_driven'."""
    return "protocol_driven"


def get_tlf_shell_summary(extraction: Dict[str, Any], protocol_text: str = "") -> Dict[str, Any]:
    """
    Get a summary of expected TLF generation.
    Simplified from v2 — actual table decisions are made at generation time by code.
    """
    arms = extraction.get("treatment_arms", []) if extraction else []
    arm_names = [a.get("arm_name", f"Arm {i+1}") for i, a in enumerate(arms)]
    indication = extraction.get("disease_classification", {}).get("indication", "")

    return {
        "detected_study_type": "protocol_driven",
        "detected_drug_classes": [],
        "detected_study_design": extraction.get("study_design", {}).get("design_type", ""),
        "generator_available": True,
        "universal_generator_available": True,
        "universal_expansion": {
            "period_stratification": {
                "required": False,
                "config": "fact_driven",
                "multiplier": 1,
            },
            "population_assessment_matrix": {
                "populations": ["ITT", "Safety"],
                "assessments": ["fact_driven"],
                "multiplier": 1,
            },
            "regions": ["global"],
            "pk_required": False,
            "immunogenicity_required": False,
        },
        "expected_table_count": "~100 (determined at generation time from extracted facts)",
        "table_counts_by_category": {},
    }


def calculate_expected_table_count(extraction: Dict[str, Any], protocol_text: str = "") -> Dict[str, Any]:
    """Stub for backwards compatibility."""
    return {
        "total_tables": 100,
        "by_category": {},
        "multipliers": {"period": 1, "efficacy_population_assessment": 1},
    }


def build_universal_config(extraction: Dict[str, Any], protocol_text: str = "") -> Dict[str, Any]:
    """Stub for backwards compatibility."""
    return {
        "study_config": {
            "study_type": "protocol_driven",
            "require_period_stratification": False,
            "period_config": "fact_driven",
            "regions": ["global"],
        }
    }
