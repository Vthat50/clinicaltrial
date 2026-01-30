"""DOCX output writer for generated SAPs."""
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from typing import Optional

from rich.console import Console

console = Console()


@dataclass
class DocumentMetadata:
    """Metadata for the SAP document."""
    title: str
    nct_id: str
    version: str = "1.0"
    date: str = None
    sponsor: str = ""
    author: str = "Auto-generated"

    def __post_init__(self):
        if not self.date:
            self.date = datetime.now().strftime("%Y-%m-%d")


class SAPDocxWriter:
    """Generate formatted DOCX files for SAPs."""

    def __init__(self):
        self.docx_available = False
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            self.docx_available = True
        except ImportError:
            console.print("[yellow]python-docx not installed. Install with: pip install python-docx[/yellow]")

    def create_sap_document(
        self,
        metadata: DocumentMetadata,
        sections: dict[str, str],
        output_path: Path,
        validation_score: float = None
    ) -> bool:
        """Create a formatted SAP DOCX document.

        Args:
            metadata: Document metadata
            sections: Dict of section_type -> content
            output_path: Path to save DOCX file
            validation_score: Optional validation score to include

        Returns:
            True if successful
        """
        if not self.docx_available:
            console.print("[red]Cannot create DOCX - python-docx not installed[/red]")
            return False

        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title page
            self._add_title_page(doc, metadata)

            # Table of contents placeholder
            doc.add_heading("Table of Contents", level=1)
            doc.add_paragraph("[Update table of contents after final review]")
            doc.add_page_break()

            # Document history
            self._add_document_history(doc, metadata)
            doc.add_page_break()

            # Main sections
            section_order = [
                ("introduction", "1. Introduction"),
                ("objectives", "2. Study Objectives"),
                ("endpoints", "3. Study Endpoints"),
                ("study_design", "4. Study Design"),
                ("analysis_populations", "5. Analysis Populations"),
                ("sample_size", "6. Sample Size Determination"),
                ("statistical_methods", "7. General Statistical Methods"),
                ("efficacy_analyses", "8. Efficacy Analyses"),
                ("safety_analyses", "9. Safety Analyses"),
                ("missing_data", "10. Handling of Missing Data"),
                ("interim_analyses", "11. Interim Analyses"),
                ("subgroup_analyses", "12. Subgroup Analyses"),
            ]

            for section_key, section_title in section_order:
                if section_key in sections:
                    self._add_section(doc, section_title, sections[section_key])

            # Appendices
            doc.add_page_break()
            doc.add_heading("Appendices", level=1)
            doc.add_paragraph("A. List of Abbreviations")
            doc.add_paragraph("B. Shell Tables and Figures")
            doc.add_paragraph("C. Programming Specifications")

            # Validation score footer
            if validation_score is not None:
                doc.add_page_break()
                doc.add_heading("Document Quality Metrics", level=1)
                doc.add_paragraph(f"Validation Score: {validation_score}/100")
                doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")

            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            console.print(f"[green]SAP saved to {output_path}[/green]")
            return True

        except Exception as e:
            console.print(f"[red]Failed to create DOCX: {e}[/red]")
            return False

    def _add_title_page(self, doc, metadata: DocumentMetadata):
        """Add title page to document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title = doc.add_heading("Statistical Analysis Plan", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with study info
        doc.add_paragraph()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(metadata.title)
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Metadata table
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"NCT ID: {metadata.nct_id}\n").bold = True
        info_para.add_run(f"Version: {metadata.version}\n")
        info_para.add_run(f"Date: {metadata.date}\n")
        if metadata.sponsor:
            info_para.add_run(f"Sponsor: {metadata.sponsor}\n")

        doc.add_paragraph()
        doc.add_paragraph()

        # Confidentiality notice
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notice.add_run("CONFIDENTIAL").bold = True

        doc.add_page_break()

    def _add_document_history(self, doc, metadata: DocumentMetadata):
        """Add document history table."""
        doc.add_heading("Document History", level=1)

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        # Header row
        headers = ["Version", "Date", "Author", "Description"]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Data row
        table.rows[1].cells[0].text = metadata.version
        table.rows[1].cells[1].text = metadata.date
        table.rows[1].cells[2].text = metadata.author
        table.rows[1].cells[3].text = "Initial SAP generation"

    def _add_section(self, doc, title: str, content: str):
        """Add a section with title and content."""
        doc.add_heading(title, level=1)

        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Check for bullet points
                if para.strip().startswith('•') or para.strip().startswith('-'):
                    for line in para.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip().lstrip('•-').strip(), style='List Bullet')
                # Check for sub-headers
                elif para.strip().startswith('#'):
                    level = para.count('#', 0, 4)
                    text = para.lstrip('#').strip()
                    doc.add_heading(text, level=min(level + 1, 4))
                else:
                    doc.add_paragraph(para.strip())


def create_sap_docx(
    nct_id: str,
    title: str,
    sections: dict[str, str],
    output_dir: Path,
    validation_score: float = None
) -> Optional[Path]:
    """Convenience function to create SAP DOCX.

    Args:
        nct_id: Study NCT ID
        title: Study title
        sections: Dict of section_type -> content
        output_dir: Output directory
        validation_score: Optional validation score

    Returns:
        Path to created file, or None if failed
    """
    writer = SAPDocxWriter()

    metadata = DocumentMetadata(
        title=title,
        nct_id=nct_id,
    )

    output_path = output_dir / f"{nct_id}_SAP.docx"

    if writer.create_sap_document(metadata, sections, output_path, validation_score):
        return output_path
    return None
